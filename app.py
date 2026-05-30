from dotenv import load_dotenv
load_dotenv()

from flask import Flask, render_template, request, send_file, make_response, redirect, jsonify
import pandas as pd
from io import BytesIO
from flask import Response
from scraper_logic import scrape_jobs
import json
from datetime import datetime
import os
import glob
from apscheduler.schedulers.background import BackgroundScheduler
import atexit
import smtplib
from email.message import EmailMessage
from zipfile import ZipFile
from openpyxl import Workbook
import os
import psycopg2
from sqlalchemy import create_engine, text
import json
from werkzeug.security import generate_password_hash, check_password_hash
from flask import session
import logging
import sys
from flask_limiter import Limiter

# Configure logging to show in Railway
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Application Configuration
APP_DOMAIN = "findmeajob.xyz"  # Production domain
APP_URL = f"https://{APP_DOMAIN}"

# Database setup
_db_engine = None

def get_db_connection():
    global _db_engine
    if _db_engine is None:
        database_url = os.environ.get('DATABASE_URL')
        if database_url:
            _db_engine = create_engine(database_url)
    return _db_engine

def init_database():
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS saved_searches (
                    id SERIAL PRIMARY KEY,
                    name VARCHAR(255) NOT NULL,
                    timestamp VARCHAR(50) NOT NULL,
                    criteria JSONB NOT NULL,
                    schedule VARCHAR(50) DEFAULT 'none',
                    last_run_date VARCHAR(50) DEFAULT ''
                )
            """))
            conn.commit()

def init_users_table():
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS users (
                    id SERIAL PRIMARY KEY,
                    username VARCHAR(255) UNIQUE NOT NULL,
                    email VARCHAR(255) UNIQUE NOT NULL,
                    password_hash VARCHAR(255) NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            
            # Add email column to existing users table if it doesn't exist
            conn.execute(text("""
                ALTER TABLE users 
                ADD COLUMN IF NOT EXISTS email VARCHAR(255) UNIQUE
            """))
            
            # Add user_id columns to saved_searches table
            conn.execute(text("""
                ALTER TABLE saved_searches 
                ADD COLUMN IF NOT EXISTS user_id INTEGER REFERENCES users(id)
            """))
            conn.commit()

init_database()
init_users_table()

def init_files_table():
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS scheduled_files (
                    id SERIAL PRIMARY KEY,
                    search_name VARCHAR(255) NOT NULL,
                    user_id INTEGER REFERENCES users(id),
                    file_data BYTEA NOT NULL,
                    filename VARCHAR(255) NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(search_name, user_id)
                )
            """))
            conn.commit()

init_files_table()


def init_password_reset_table():
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS password_reset_tokens (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id),
                    token VARCHAR(255) UNIQUE NOT NULL,
                    expires_at TIMESTAMP NOT NULL,
                    used BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            conn.commit()

init_password_reset_table()

def init_user_activity_table():
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS user_activity (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id),
                    action_type VARCHAR(50) NOT NULL,
                    action_details TEXT,
                    ip_address VARCHAR(45),
                    user_agent TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            conn.commit()

init_user_activity_table()

def init_search_limits_table():
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS daily_search_limits (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id),
                    search_date DATE NOT NULL,
                    search_count INTEGER DEFAULT 0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(user_id, search_date)
                )
            """))
            conn.commit()

init_search_limits_table()


def cleanup_old_records():
    """Delete old records from tables that grow unbounded over time."""
    engine = get_db_connection()
    if not engine:
        return
    try:
        with engine.connect() as conn:
            # Keep only last 90 days of activity logs
            conn.execute(text("""
                DELETE FROM user_activity
                WHERE created_at < NOW() - INTERVAL '90 days'
            """))
            # Daily search limits older than 30 days are never needed again
            conn.execute(text("""
                DELETE FROM daily_search_limits
                WHERE search_date < CURRENT_DATE - INTERVAL '30 days'
            """))
            # Remove password reset tokens that are expired or used and older than 7 days
            conn.execute(text("""
                DELETE FROM password_reset_tokens
                WHERE (used = TRUE OR expires_at < NOW())
                AND created_at < NOW() - INTERVAL '7 days'
            """))
            conn.commit()
        logger.info("✅ Database cleanup completed successfully")
    except Exception as e:
        logger.error(f"❌ Database cleanup failed: {e}")

cleanup_old_records()

def init_cv_table():
    """Initialize table for storing user CVs"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS user_cvs (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    cv_name VARCHAR(255) NOT NULL,
                    file_data BYTEA NOT NULL,
                    file_type VARCHAR(10) NOT NULL,
                    file_size INTEGER NOT NULL,
                    extracted_text TEXT,
                    skills_detected TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(user_id, cv_name)
                )
            """))
            conn.commit()

init_cv_table()


def init_job_analyses_table():
    """Initialize table for storing job matching analyses"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS job_analyses (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    cv_id INTEGER REFERENCES user_cvs(id) ON DELETE CASCADE,
                    job_title VARCHAR(500) NOT NULL,
                    job_company VARCHAR(500),
                    job_description TEXT NOT NULL,
                    match_score INTEGER,
                    skills_match JSONB,
                    skills_missing JSONB,
                    recommendations TEXT,
                    full_analysis TEXT,
                    decision VARCHAR(50),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            conn.commit()

init_job_analyses_table()


def init_ai_usage_tracking_table():
    """Initialize table for tracking AI API usage and costs"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS ai_usage_tracking (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    feature_type VARCHAR(50) NOT NULL,
                    tokens_input INTEGER,
                    tokens_output INTEGER,
                    estimated_cost DECIMAL(10, 6),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            conn.commit()

init_ai_usage_tracking_table()


def init_prompt_templates_table():
    """Initialize table for storing system prompt templates"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS prompt_templates (
                    id SERIAL PRIMARY KEY,
                    name VARCHAR(100) NOT NULL,
                    description TEXT,
                    prompt_text TEXT NOT NULL,
                    target_profile VARCHAR(100),
                    is_default BOOLEAN DEFAULT FALSE,
                    version VARCHAR(20) DEFAULT 'v1',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            conn.commit()

init_prompt_templates_table()


def init_user_prompt_preferences_table():
    """Initialize table for storing user prompt preferences"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS user_prompt_preferences (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    template_id INTEGER REFERENCES prompt_templates(id),
                    custom_prompt_text TEXT,
                    is_active BOOLEAN DEFAULT TRUE,
                    ab_test_group VARCHAR(50),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))

            conn.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_user_prompt_active
                ON user_prompt_preferences(user_id, is_active)
            """))
            conn.commit()

init_user_prompt_preferences_table()


def init_master_templates_table():
    """Initialize table for storing user master resume templates"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS user_master_templates (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    template_text TEXT NOT NULL,
                    original_filename VARCHAR(255),
                    version INTEGER DEFAULT 1,
                    is_active BOOLEAN DEFAULT TRUE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))

            # Create partial unique index - only one active template per user
            conn.execute(text("""
                CREATE UNIQUE INDEX IF NOT EXISTS idx_user_master_template_active_unique
                ON user_master_templates(user_id)
                WHERE is_active = TRUE
            """))

            # Create regular index for lookups
            conn.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_user_master_template_active
                ON user_master_templates(user_id, is_active)
            """))
            conn.commit()

init_master_templates_table()


def init_cv_customization_sessions_table():
    """Initialize table for storing CV customization sessions"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS cv_customization_sessions (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    analysis_id INTEGER REFERENCES job_analyses(id) ON DELETE CASCADE,
                    job_title VARCHAR(500),
                    job_company VARCHAR(500),
                    selected_headline TEXT,
                    bullet_analysis JSONB,
                    approved_bullets JSONB,
                    new_bullets JSONB,
                    match_score_progression JSONB,
                    status VARCHAR(50) DEFAULT 'in_progress',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))

            # Add bullet_analysis column if it doesn't exist (for existing tables)
            conn.execute(text("""
                ALTER TABLE cv_customization_sessions
                ADD COLUMN IF NOT EXISTS bullet_analysis JSONB
            """))

            # Add approved_bullets column if it doesn't exist (for existing tables)
            conn.execute(text("""
                ALTER TABLE cv_customization_sessions
                ADD COLUMN IF NOT EXISTS approved_bullets JSONB
            """))

            conn.execute(text("""
                ALTER TABLE cv_customization_sessions
                ADD COLUMN IF NOT EXISTS selected_roles JSONB
            """))

            conn.execute(text("""
                ALTER TABLE cv_customization_sessions
                ADD COLUMN IF NOT EXISTS bullet_analysis_by_role JSONB
            """))

            conn.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_cv_sessions_user
                ON cv_customization_sessions(user_id, status)
            """))
            conn.commit()

init_cv_customization_sessions_table()


def init_interview_sessions_table():
    """Initialize table for storing interview practice sessions"""
    engine = get_db_connection()
    if engine:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS interview_sessions (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    cv_session_id INTEGER REFERENCES cv_customization_sessions(id) ON DELETE CASCADE,

                    -- Progress tracking
                    current_question INTEGER DEFAULT 1,
                    completed BOOLEAN DEFAULT FALSE,

                    -- All data in JSONB (flexible, fast, easy)
                    questions JSONB NOT NULL,
                    answers JSONB DEFAULT '{}',
                    evaluations JSONB DEFAULT '{}',

                    -- Summary scores (calculated when complete)
                    overall_score DECIMAL(3,1),

                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))

            conn.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_interview_user_cv
                ON interview_sessions(user_id, cv_session_id)
            """))

            conn.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_interview_completed
                ON interview_sessions(user_id, completed)
            """))

            conn.commit()

init_interview_sessions_table()


def seed_initial_prompt_templates():
    """Seed initial prompt templates if they don't exist"""
    engine = get_db_connection()
    if not engine:
        return

    with engine.connect() as conn:
        # Check if templates already exist
        result = conn.execute(text("SELECT COUNT(*) FROM prompt_templates"))
        count = result.fetchone()[0]

        if count > 0:
            return  # Templates already seeded

        # Template #1: Default - General Professional
        default_prompt = """You are an expert career coach and recruiter with 15+ years of experience evaluating candidates across various industries. Your goal is to provide honest, accurate assessments of how well a candidate's background matches a specific job opportunity.

**YOUR TASK:**
Analyze the candidate's CV against the job description and provide a comprehensive match assessment.

**EVALUATION CRITERIA:**
1. **Relevant Experience (40%):** How closely does their work history align with the role's requirements?
2. **Skills Match (30%):** Do they possess the technical and soft skills needed?
3. **Industry Fit (20%):** Does their industry background prepare them for this role?
4. **Seniority Match (10%):** Does their career level align with the role's expectations?

**SCORING GUIDELINES:**
- **90-100%:** Exceptional match - candidate exceeds most requirements
- **75-89%:** Strong match - candidate meets most key requirements with minor gaps
- **60-74%:** Good match - candidate meets many requirements but has notable gaps
- **50-59%:** Fair match - candidate has transferable skills but significant gaps
- **Below 50%:** Poor match - major gaps in experience or skills

**OUTPUT REQUIREMENTS:**
Respond with valid JSON containing:
- match_score (0-100)
- skills_match: Array of matched skills with evidence
- skills_missing: Array of skills needed but not demonstrated
- recommendations: Array of specific suggestions for the candidate"""

        # Template #2: Finance VP (Tier-1 IB) - Enhanced RECRUITER_SYSTEM_PROMPT
        finance_vp_prompt = """You are an experienced recruiter at a top-tier executive search firm with 20+ years of experience placing VP, SVP, and ED-level candidates at tier-1 investment banks (Goldman Sachs, JP Morgan, Morgan Stanley, Citi, etc.). You specialize in evaluating senior financial services talent for AI/ML governance, risk management, and product leadership roles.

**YOUR EXPERTISE:**
You have personally placed 200+ VPs and 50+ SVPs/EDs in financial services. You understand what separates a 92% match from an 85% match. You provide calibrated, consistent evaluations that help candidates understand exactly where they stand.

**CANDIDATE CONTEXT (Grace):**
- 15+ years in financial services at tier-1 institutions (JP Morgan, Citi, Morgan Stanley)
- Currently VP, AI Governance at JP Morgan
- Deep expertise in AI/ML risk management, model validation, regulatory compliance
- Led enterprise-wide AI governance frameworks
- Looking for VP/SVP/ED roles in AI governance, risk management, or product management at financial institutions

**YOUR EVALUATION FRAMEWORK (100-Point System):**

**1. DIRECT EXPERIENCE MATCH (40 points)**
- Exact role type match (AI Governance, Risk, Product)
- Institution type match (tier-1 IB vs fintech vs tech)
- Relevant domain expertise (AI/ML, regulatory, model risk)
- Years of experience at appropriate level

**2. TRANSFERABLE SKILLS (25 points)**
- Technical skills (AI/ML, data governance, model validation)
- Leadership skills (stakeholder management, cross-functional leadership)
- Regulatory expertise (Fed, OCC, GDPR, AI regulations)
- Strategic capabilities (framework design, policy development)

**3. INDUSTRY FIT (20 points)**
- Financial services background
- Regulatory environment familiarity
- Enterprise scale experience
- Cultural fit (bank vs tech vs consulting)

**4. SENIORITY MATCH (15 points)**
- Title level alignment (VP to VP, VP to SVP, etc.)
- Scope of responsibility match
- People management experience
- P&L or budget ownership

**CALIBRATION EXAMPLES (Learn these scoring patterns):**

**90-100% Matches:**
1. VP AI Governance at Goldman Sachs → VP AI Governance at Morgan Stanley (95%)
2. VP Model Risk at JP Morgan → VP AI Risk at Citi (92%)
3. SVP AI Strategy at Bank of America → SVP AI Governance at Wells Fargo (94%)

**75-89% Matches:**
4. VP AI Governance at JP Morgan → VP AI Governance at Stripe (fintech) (82%)
   - Reason: Strong role match but institution type shift (bank → fintech)
5. VP AI Governance at JP Morgan → VP Product Management (AI) at Capital One (78%)
   - Reason: Pivot from governance to product, different institution tier
6. VP Model Risk at JP Morgan → Director AI Strategy at McKinsey (consulting) (76%)
   - Reason: Title step down, industry shift to consulting

**60-74% Matches:**
7. VP AI Governance at JP Morgan → VP Product Management at tech company (68%)
   - Reason: Industry shift (finance → tech), role pivot (governance → product)
8. VP Model Risk at JP Morgan → VP Program Management (AI initiatives) at Citi (72%)
   - Reason: Shift from risk to program management, less specialized role
9. Senior Manager AI Governance at JP Morgan → VP AI Governance at regional bank (65%)
   - Reason: Title promotion needed, institution tier drop

**55-70% Matches:**
10. VP AI Governance at JP Morgan → VP Risk Management (no AI focus) at Morgan Stanley (65%)
    - Reason: Loses AI specialization, becomes generic risk role
11. VP Model Risk at JP Morgan → Senior Director Analytics at tech company (62%)
    - Reason: Title ambiguity, industry shift, role is more analytics than governance

**Below 55% Matches:**
12. VP AI Governance at JP Morgan → Head of Data Science at startup (48%)
    - Reason: Major industry shift, completely different role (governance → IC technical)
13. VP AI Governance at JP Morgan → Chief Risk Officer at small fintech (52%)
    - Reason: Title seems senior but scope much smaller, institution scale mismatch
14. VP Model Risk at JP Morgan → Product Manager (AI tools) at Series A startup (45%)
    - Reason: Title demotion, industry shift, completely different scope/scale

**CONSISTENCY RULES (Apply these strictly):**

1. **Exact Title + Institution Type Match = Minimum 85%**
   - VP AI Governance (bank) → VP AI Governance (bank) = 85-95%
   - Only deduct for scope, team size, or institution prestige differences

2. **Competitor Bank Match = Minimum 88%**
   - JP Morgan ↔ Goldman/Morgan Stanley/Citi = very high match
   - These are peer institutions with comparable complexity

3. **Years of Experience Scope Match:**
   - If candidate has "X years managing Y use cases" and JD requires "managing Z use cases," score should not drop below 80%
   - Example: Candidate managed 50 AI models, JD requires managing AI models → 80%+ even if exact count differs

4. **Exact Title Match (Different Institution Type) = Minimum 75%**
   - VP AI Governance (tier-1 bank) → VP AI Governance (fintech) = 75-85%
   - VP AI Governance (tier-1 bank) → VP AI Governance (tech company) = 70-82%

5. **Role Pivot Within Same Institution Type:**
   - Governance → Product (same industry) = 65-78%
   - Governance → Risk (same industry) = 75-85%
   - Risk → Program Management (same industry) = 68-76%

6. **People Management Cap:**
   - If role requires people management and candidate has NO direct reports mentioned → cap at 70%
   - If candidate has people management and role doesn't require it → no penalty

7. **P&L Ownership Cap:**
   - If role requires P&L ownership and candidate has none → cap at 68%

8. **Institution Type Penalties:**
   - Tier-1 IB → Tier-2/regional bank = -8 to -12 points
   - Bank → Fintech = -5 to -10 points
   - Bank → Big Tech = -10 to -18 points
   - Bank → Startup = -20 to -30 points

9. **Seniority Mismatch Penalties:**
   - VP → SVP (promotion needed) = -5 to -8 points
   - SVP → VP (step down) = -8 to -12 points
   - Director → VP (promotion needed) = -10 to -15 points

10. **AI Specialization Rule:**
    - If candidate is "AI Governance" specialist and role is generic "Risk Management" (no AI) → cap at 70%
    - If candidate is generic risk and role requires AI specialization → cap at 60%

11. **Regulatory Expertise Match:**
    - If role requires specific regulatory knowledge (Fed, OCC, GDPR) and candidate has it → +5 to +8 points
    - If role requires regulatory expertise and candidate has none → -15 to -20 points

12. **Cross-Functional Leadership:**
    - If candidate has "led cross-functional initiatives" and JD requires it → automatic inclusion in strengths
    - If JD emphasizes stakeholder management and candidate has evidence → +5 points

13. **Consistency Check:**
    - If two job descriptions are 90% similar in requirements, the same candidate's scores should differ by no more than 5 points
    - Review your scoring: Does this match the calibration examples above?

**EVIDENCE REQUIREMENTS:**

When assessing skills:
- "Strong" evidence = explicitly mentioned in CV with concrete examples/metrics
- "Moderate" evidence = clearly implied by role/responsibilities but not explicitly stated
- "Basic" evidence = tangentially related experience that could transfer
- NO evidence = do not list as a matched skill

**OUTPUT FORMAT (JSON only, no markdown):**

{
  "match_score": <integer 0-100>,
  "scoring_breakdown": {
    "direct_experience": <0-40 points>,
    "transferable_skills": <0-25 points>,
    "industry_fit": <0-20 points>,
    "seniority_match": <0-15 points>
  },
  "skills_match": [
    {
      "skill": "AI/ML Governance",
      "evidence": "Led AI governance framework at JP Morgan for 3+ years",
      "strength": "strong"
    }
  ],
  "skills_missing": [
    {
      "skill": "P&L Ownership",
      "importance": "high",
      "impact": "May need to demonstrate budget management experience in interviews"
    }
  ],
  "recommendations": [
    "Emphasize your experience managing [specific area] in your application",
    "Be prepared to discuss how your governance experience translates to [required skill]"
  ],
  "match_rationale": "Brief 2-3 sentence explanation of the score, referencing which calibration example this most resembles"
}

**REMEMBER:** You are evaluating Grace, a VP at JP Morgan with 15+ years in tier-1 financial services and deep AI governance expertise. Score accordingly using the calibration examples above."""

        # Insert Template #1
        conn.execute(text("""
            INSERT INTO prompt_templates (name, description, prompt_text, target_profile, is_default, version)
            VALUES (:name, :description, :prompt_text, :target_profile, :is_default, :version)
        """), {
            "name": "Default - General Professional",
            "description": "A balanced, professional evaluation suitable for most career backgrounds and industries.",
            "prompt_text": default_prompt,
            "target_profile": "General",
            "is_default": True,
            "version": "v1"
        })

        # Insert Template #2
        conn.execute(text("""
            INSERT INTO prompt_templates (name, description, prompt_text, target_profile, is_default, version)
            VALUES (:name, :description, :prompt_text, :target_profile, :is_default, :version)
        """), {
            "name": "Finance VP (Tier-1 IB)",
            "description": "Specialized evaluation for VP/SVP-level financial services professionals with AI/ML governance, risk, or product experience at tier-1 investment banks. Includes detailed calibration for banking roles.",
            "prompt_text": finance_vp_prompt,
            "target_profile": "VP Banking - AI/Governance",
            "is_default": False,
            "version": "v1"
        })

        conn.commit()
        print("✓ Seeded 2 initial prompt templates")

seed_initial_prompt_templates()


def assign_user_to_finance_vp_template():
    """Assign user_id=1 (Grace) to the Finance VP template if user exists"""
    engine = get_db_connection()
    if not engine:
        return

    try:
        with engine.connect() as conn:
            # First, check if user_id=1 exists in users table
            user_check = conn.execute(text("""
                SELECT COUNT(*) FROM users WHERE id = 1
            """))
            user_exists = user_check.fetchone()[0] > 0

            if not user_exists:
                # User doesn't exist yet, skip gracefully (don't crash)
                return

            # Check if user_id=1 already has a preference assigned
            result = conn.execute(text("""
                SELECT COUNT(*) FROM user_prompt_preferences WHERE user_id = 1
            """))
            count = result.fetchone()[0]

            if count > 0:
                return  # User already has a template assigned

            # Get the Finance VP template ID
            template = conn.execute(text("""
                SELECT id FROM prompt_templates WHERE name = 'Finance VP (Tier-1 IB)'
            """))
            template_row = template.fetchone()

            if not template_row:
                print("⚠ Finance VP template not found, skipping user assignment")
                return

            template_id = template_row[0]

            # Assign user_id=1 to this template
            conn.execute(text("""
                INSERT INTO user_prompt_preferences (user_id, template_id, is_active)
                VALUES (:user_id, :template_id, TRUE)
            """), {"user_id": 1, "template_id": template_id})

            conn.commit()
            print(f"✓ Assigned user_id=1 to Finance VP (Tier-1 IB) template (template_id={template_id})")
    except Exception as e:
        # Gracefully handle any errors during assignment (don't crash the app)
        print(f"⚠ Could not auto-assign template to user_id=1: {e}")
        return

assign_user_to_finance_vp_template()


def update_finance_vp_template_remove_persona():
    """Update the Finance VP template to remove persona name"""
    engine = get_db_connection()
    if not engine:
        return

    with engine.connect() as conn:
        # Updated prompt without persona name
        finance_vp_prompt = """You are an experienced recruiter at a top-tier executive search firm with 20+ years of experience placing VP, SVP, and ED-level candidates at tier-1 investment banks (Goldman Sachs, JP Morgan, Morgan Stanley, Citi, etc.). You specialize in evaluating senior financial services talent for AI/ML governance, risk management, and product leadership roles.

**YOUR EXPERTISE:**
You have personally placed 200+ VPs and 50+ SVPs/EDs in financial services. You understand what separates a 92% match from an 85% match. You provide calibrated, consistent evaluations that help candidates understand exactly where they stand.

**CANDIDATE CONTEXT (Grace):**
- 15+ years in financial services at tier-1 institutions (JP Morgan, Citi, Morgan Stanley)
- Currently VP, AI Governance at JP Morgan
- Deep expertise in AI/ML risk management, model validation, regulatory compliance
- Led enterprise-wide AI governance frameworks
- Looking for VP/SVP/ED roles in AI governance, risk management, or product management at financial institutions

**YOUR EVALUATION FRAMEWORK (100-Point System):**

**1. DIRECT EXPERIENCE MATCH (40 points)**
- Exact role type match (AI Governance, Risk, Product)
- Institution type match (tier-1 IB vs fintech vs tech)
- Relevant domain expertise (AI/ML, regulatory, model risk)
- Years of experience at appropriate level

**2. TRANSFERABLE SKILLS (25 points)**
- Technical skills (AI/ML, data governance, model validation)
- Leadership skills (stakeholder management, cross-functional leadership)
- Regulatory expertise (Fed, OCC, GDPR, AI regulations)
- Strategic capabilities (framework design, policy development)

**3. INDUSTRY FIT (20 points)**
- Financial services background
- Regulatory environment familiarity
- Enterprise scale experience
- Cultural fit (bank vs tech vs consulting)

**4. SENIORITY MATCH (15 points)**
- Title level alignment (VP to VP, VP to SVP, etc.)
- Scope of responsibility match
- People management experience
- P&L or budget ownership

**CALIBRATION EXAMPLES (Learn these scoring patterns):**

**90-100% Matches:**
1. VP AI Governance at Goldman Sachs → VP AI Governance at Morgan Stanley (95%)
2. VP Model Risk at JP Morgan → VP AI Risk at Citi (92%)
3. SVP AI Strategy at Bank of America → SVP AI Governance at Wells Fargo (94%)

**75-89% Matches:**
4. VP AI Governance at JP Morgan → VP AI Governance at Stripe (fintech) (82%)
   - Reason: Strong role match but institution type shift (bank → fintech)
5. VP AI Governance at JP Morgan → VP Product Management (AI) at Capital One (78%)
   - Reason: Pivot from governance to product, different institution tier
6. VP Model Risk at JP Morgan → Director AI Strategy at McKinsey (consulting) (76%)
   - Reason: Title step down, industry shift to consulting

**60-74% Matches:**
7. VP AI Governance at JP Morgan → VP Product Management at tech company (68%)
   - Reason: Industry shift (finance → tech), role pivot (governance → product)
8. VP Model Risk at JP Morgan → VP Program Management (AI initiatives) at Citi (72%)
   - Reason: Shift from risk to program management, less specialized role
9. Senior Manager AI Governance at JP Morgan → VP AI Governance at regional bank (65%)
   - Reason: Title promotion needed, institution tier drop

**55-70% Matches:**
10. VP AI Governance at JP Morgan → VP Risk Management (no AI focus) at Morgan Stanley (65%)
    - Reason: Loses AI specialization, becomes generic risk role
11. VP Model Risk at JP Morgan → Senior Director Analytics at tech company (62%)
    - Reason: Title ambiguity, industry shift, role is more analytics than governance

**Below 55% Matches:**
12. VP AI Governance at JP Morgan → Head of Data Science at startup (48%)
    - Reason: Major industry shift, completely different role (governance → IC technical)
13. VP AI Governance at JP Morgan → Chief Risk Officer at small fintech (52%)
    - Reason: Title seems senior but scope much smaller, institution scale mismatch
14. VP Model Risk at JP Morgan → Product Manager (AI tools) at Series A startup (45%)
    - Reason: Title demotion, industry shift, completely different scope/scale

**CONSISTENCY RULES (Apply these strictly):**

1. **Exact Title + Institution Type Match = Minimum 85%**
   - VP AI Governance (bank) → VP AI Governance (bank) = 85-95%
   - Only deduct for scope, team size, or institution prestige differences

2. **Competitor Bank Match = Minimum 88%**
   - JP Morgan ↔ Goldman/Morgan Stanley/Citi = very high match
   - These are peer institutions with comparable complexity

3. **Years of Experience Scope Match:**
   - If candidate has "X years managing Y use cases" and JD requires "managing Z use cases," score should not drop below 80%
   - Example: Candidate managed 50 AI models, JD requires managing AI models → 80%+ even if exact count differs

4. **Exact Title Match (Different Institution Type) = Minimum 75%**
   - VP AI Governance (tier-1 bank) → VP AI Governance (fintech) = 75-85%
   - VP AI Governance (tier-1 bank) → VP AI Governance (tech company) = 70-82%

5. **Role Pivot Within Same Institution Type:**
   - Governance → Product (same industry) = 65-78%
   - Governance → Risk (same industry) = 75-85%
   - Risk → Program Management (same industry) = 68-76%

6. **People Management Cap:**
   - If role requires people management and candidate has NO direct reports mentioned → cap at 70%
   - If candidate has people management and role doesn't require it → no penalty

7. **P&L Ownership Cap:**
   - If role requires P&L ownership and candidate has none → cap at 68%

8. **Institution Type Penalties:**
   - Tier-1 IB → Tier-2/regional bank = -8 to -12 points
   - Bank → Fintech = -5 to -10 points
   - Bank → Big Tech = -10 to -18 points
   - Bank → Startup = -20 to -30 points

9. **Seniority Mismatch Penalties:**
   - VP → SVP (promotion needed) = -5 to -8 points
   - SVP → VP (step down) = -8 to -12 points
   - Director → VP (promotion needed) = -10 to -15 points

10. **AI Specialization Rule:**
    - If candidate is "AI Governance" specialist and role is generic "Risk Management" (no AI) → cap at 70%
    - If candidate is generic risk and role requires AI specialization → cap at 60%

11. **Regulatory Expertise Match:**
    - If role requires specific regulatory knowledge (Fed, OCC, GDPR) and candidate has it → +5 to +8 points
    - If role requires regulatory expertise and candidate has none → -15 to -20 points

12. **Cross-Functional Leadership:**
    - If candidate has "led cross-functional initiatives" and JD requires it → automatic inclusion in strengths
    - If JD emphasizes stakeholder management and candidate has evidence → +5 points

13. **Consistency Check:**
    - If two job descriptions are 90% similar in requirements, the same candidate's scores should differ by no more than 5 points
    - Review your scoring: Does this match the calibration examples above?

**EVIDENCE REQUIREMENTS:**

When assessing skills:
- "Strong" evidence = explicitly mentioned in CV with concrete examples/metrics
- "Moderate" evidence = clearly implied by role/responsibilities but not explicitly stated
- "Basic" evidence = tangentially related experience that could transfer
- NO evidence = do not list as a matched skill

**OUTPUT FORMAT (JSON only, no markdown):**

{
  "match_score": <integer 0-100>,
  "scoring_breakdown": {
    "direct_experience": <0-40 points>,
    "transferable_skills": <0-25 points>,
    "industry_fit": <0-20 points>,
    "seniority_match": <0-15 points>
  },
  "skills_match": [
    {
      "skill": "AI/ML Governance",
      "evidence": "Led AI governance framework at JP Morgan for 3+ years",
      "strength": "strong"
    }
  ],
  "skills_missing": [
    {
      "skill": "P&L Ownership",
      "importance": "high",
      "impact": "May need to demonstrate budget management experience in interviews"
    }
  ],
  "recommendations": [
    "Emphasize your experience managing [specific area] in your application",
    "Be prepared to discuss how your governance experience translates to [required skill]"
  ],
  "match_rationale": "Brief 2-3 sentence explanation of the score, referencing which calibration example this most resembles"
}

**REMEMBER:** You are evaluating Grace, a VP at JP Morgan with 15+ years in tier-1 financial services and deep AI governance expertise. Score accordingly using the calibration examples above."""

        # Update the template
        result = conn.execute(text("""
            UPDATE prompt_templates
            SET prompt_text = :prompt_text,
                updated_at = CURRENT_TIMESTAMP
            WHERE name = 'Finance VP (Tier-1 IB)'
        """), {"prompt_text": finance_vp_prompt})

        conn.commit()

        if result.rowcount > 0:
            print(f"✓ Updated Finance VP template to remove persona name (affected {result.rowcount} row)")
        else:
            print("⚠ Finance VP template not found or not updated")

update_finance_vp_template_remove_persona()


def generate_reset_token():
    import secrets
    return secrets.token_urlsafe(32)

def create_password_reset_token(user_id):
    engine = get_db_connection()
    if not engine:
        return None
    
    token = generate_reset_token()
    # Token expires in 1 hour
    from datetime import datetime, timedelta
    expires_at = datetime.now() + timedelta(hours=1)
    
    try:
        with engine.connect() as conn:
            # First, invalidate any existing tokens for this user
            conn.execute(text("""
                UPDATE password_reset_tokens 
                SET used = TRUE 
                WHERE user_id = :user_id AND used = FALSE
            """), {"user_id": user_id})
            
            # Create new token
            conn.execute(text("""
                INSERT INTO password_reset_tokens (user_id, token, expires_at)
                VALUES (:user_id, :token, :expires_at)
            """), {
                "user_id": user_id,
                "token": token,
                "expires_at": expires_at
            })
            conn.commit()
        return token
    except Exception as e:
        print(f"Error creating reset token: {e}")
        return None


# ==================== AI MATCH TOOL: CV PARSING FUNCTIONS ====================

def extract_text_from_pdf(file_data):
    """Extract text from PDF file"""
    try:
        from PyPDF2 import PdfReader
        import io

        pdf_file = io.BytesIO(file_data)
        reader = PdfReader(pdf_file)

        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"

        return text.strip()
    except ImportError as e:
        print(f"PyPDF2 not installed: {e}")
        raise Exception("PyPDF2 library not installed. Please install it.")
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_data):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        import io

        docx_file = io.BytesIO(file_data)
        doc = Document(docx_file)

        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text.strip()
    except ImportError as e:
        print(f"python-docx not installed: {e}")
        raise Exception("python-docx library not installed. Please install it.")
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Failed to parse DOCX: {str(e)}")


def parse_cv(file_data, file_type):
    """Parse CV and extract text based on file type"""
    if file_type.lower() == 'pdf':
        return extract_text_from_pdf(file_data)
    elif file_type.lower() in ['docx', 'doc']:
        return extract_text_from_docx(file_data)
    else:
        raise Exception(f"Unsupported file type: {file_type}")


def save_cv_to_db(user_id, cv_name, file_data, file_type, extracted_text):
    """Save CV to database"""
    engine = get_db_connection()
    if not engine:
        return None

    try:
        file_size = len(file_data)

        with engine.connect() as conn:
            # Check if CV with same name exists, update if yes
            result = conn.execute(text("""
                SELECT id FROM user_cvs
                WHERE user_id = :user_id AND cv_name = :cv_name
            """), {"user_id": user_id, "cv_name": cv_name})

            existing_cv = result.fetchone()

            if existing_cv:
                # Update existing CV
                conn.execute(text("""
                    UPDATE user_cvs
                    SET file_data = :file_data, file_type = :file_type,
                        file_size = :file_size, extracted_text = :extracted_text,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE id = :cv_id
                """), {
                    "file_data": file_data,
                    "file_type": file_type,
                    "file_size": file_size,
                    "extracted_text": extracted_text,
                    "cv_id": existing_cv[0]
                })
                conn.commit()
                return existing_cv[0]
            else:
                # Insert new CV
                result = conn.execute(text("""
                    INSERT INTO user_cvs (user_id, cv_name, file_data, file_type, file_size, extracted_text)
                    VALUES (:user_id, :cv_name, :file_data, :file_type, :file_size, :extracted_text)
                    RETURNING id
                """), {
                    "user_id": user_id,
                    "cv_name": cv_name,
                    "file_data": file_data,
                    "file_type": file_type,
                    "file_size": file_size,
                    "extracted_text": extracted_text
                })
                conn.commit()
                cv_id = result.fetchone()[0]
                return cv_id
    except Exception as e:
        print(f"Error saving CV to database: {e}")
        return None


def get_user_cvs(user_id):
    """Get all CVs for a user"""
    engine = get_db_connection()
    if not engine:
        return []

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT id, cv_name, file_type, file_size, created_at, updated_at
                FROM user_cvs
                WHERE user_id = :user_id
                ORDER BY updated_at DESC
            """), {"user_id": user_id})

            cvs = []
            for row in result:
                cvs.append({
                    'id': row[0],
                    'cv_name': row[1],
                    'file_type': row[2],
                    'file_size': row[3],
                    'created_at': row[4],
                    'updated_at': row[5]
                })
            return cvs
    except Exception as e:
        print(f"Error getting user CVs: {e}")
        return []


def get_cv_by_id(cv_id, user_id):
    """Get specific CV by ID (with user verification)"""
    engine = get_db_connection()
    if not engine:
        return None

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT id, cv_name, file_data, file_type, extracted_text
                FROM user_cvs
                WHERE id = :cv_id AND user_id = :user_id
            """), {"cv_id": cv_id, "user_id": user_id})

            row = result.fetchone()
            if row:
                return {
                    'id': row[0],
                    'cv_name': row[1],
                    'file_data': row[2],
                    'file_type': row[3],
                    'extracted_text': row[4]
                }
            return None
    except Exception as e:
        print(f"Error getting CV: {e}")
        return None


# ==================== AI MATCH TOOL: MASTER TEMPLATE FUNCTIONS ====================

def save_master_template(user_id, template_text, filename):
    """Save or update user's master resume template"""
    engine = get_db_connection()
    if not engine:
        return None

    try:
        with engine.connect() as conn:
            # Deactivate any existing active templates for this user
            conn.execute(text("""
                UPDATE user_master_templates
                SET is_active = FALSE
                WHERE user_id = :user_id AND is_active = TRUE
            """), {"user_id": user_id})

            # Get the next version number
            version_result = conn.execute(text("""
                SELECT COALESCE(MAX(version), 0) + 1
                FROM user_master_templates
                WHERE user_id = :user_id
            """), {"user_id": user_id})
            version = version_result.fetchone()[0]

            # Insert new template
            result = conn.execute(text("""
                INSERT INTO user_master_templates
                (user_id, template_text, original_filename, version, is_active)
                VALUES (:user_id, :template_text, :filename, :version, TRUE)
                RETURNING id
            """), {
                "user_id": user_id,
                "template_text": template_text,
                "filename": filename,
                "version": version
            })

            template_id = result.fetchone()[0]
            conn.commit()

            return template_id

    except Exception as e:
        print(f"Error saving master template: {e}")
        import traceback
        traceback.print_exc()
        return None


def get_user_master_template(user_id):
    """Get user's active master resume template"""
    engine = get_db_connection()
    if not engine:
        return None

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT id, template_text, original_filename, version, created_at, updated_at
                FROM user_master_templates
                WHERE user_id = :user_id AND is_active = TRUE
                ORDER BY created_at DESC
                LIMIT 1
            """), {"user_id": user_id})

            row = result.fetchone()
            if row:
                return {
                    'id': row[0],
                    'template_text': row[1],
                    'original_filename': row[2],
                    'version': row[3],
                    'created_at': row[4],
                    'updated_at': row[5]
                }
            return None

    except Exception as e:
        print(f"Error getting master template: {e}")
        return None


def extract_text_from_docx(file_data):
    """Extract text from .docx file while preserving structure"""
    try:
        from docx import Document
        from io import BytesIO

        # Load document from binary data
        doc = Document(BytesIO(file_data))

        text_parts = []

        for para in doc.paragraphs:
            # Preserve heading levels
            if para.style.name.startswith('Heading'):
                text_parts.append(f"\n## {para.text}\n")
            elif para.text.strip():
                # Check if it's a list item (starts with bullet or number)
                text = para.text.strip()
                if text.startswith('•') or text.startswith('-') or (len(text) > 2 and text[0].isdigit() and text[1] in '.):'):
                    text_parts.append(f"  {text}")
                else:
                    text_parts.append(text)

        # Join with newlines to preserve structure
        extracted_text = '\n'.join(text_parts)

        return extracted_text

    except Exception as e:
        print(f"Error extracting text from docx: {e}")
        import traceback
        traceback.print_exc()
        return None


# ==================== CV CUSTOMIZATION: HELPER FUNCTIONS ====================

def parse_headlines_from_template(template_text):
    """Extract all headline variations from master template

    Handles both numbered (1. Headline) and unnumbered (plain text) formats.
    """
    import re

    headlines = []
    lines = template_text.split('\n')

    in_headlines_section = False

    print(f"\n=== PARSING HEADLINES ===")
    print(f"Total lines in template: {len(lines)}")

    for i, line in enumerate(lines):
        line_stripped = line.strip()

        # Detect start of HEADLINES section
        if 'HEADLINE' in line_stripped.upper() and ('VARIATION' in line_stripped.upper() or '(' in line_stripped):
            in_headlines_section = True
            print(f"Found HEADLINES section at line {i}: '{line_stripped[:50]}...'")
            continue

        # Detect end of HEADLINES section
        if in_headlines_section and line_stripped:
            # Check if this is a separator line
            if line_stripped.startswith('___'):
                print(f"End of HEADLINES section at line {i}: separator line")
                break

            # Check if this is an ALL CAPS section header (e.g., "CORE SKILLS & COMPETENCIES")
            # Must be all uppercase AND contain multiple words OR common section keywords
            upper_line = line_stripped.upper()
            if line_stripped.isupper() and len(line_stripped) > 10:
                print(f"End of HEADLINES section at line {i}: section header '{line_stripped[:50]}...'")
                break

            # Check if starts with known section keywords
            if upper_line.startswith('CORE SKILLS') or \
               upper_line.startswith('EXPERIENCE:') or \
               upper_line.startswith('EDUCATION:') or \
               upper_line.startswith('TECHNICAL ACUMEN:') or \
               upper_line.startswith('WORK HISTORY') or \
               upper_line.startswith('JP MORGAN'):
                print(f"End of HEADLINES section at line {i}: section keyword '{line_stripped[:50]}...'")
                break

        # Extract headlines (both numbered and unnumbered formats)
        if in_headlines_section and line_stripped:
            # Try numbered format first (e.g., "1. Headline text")
            match = re.match(r'^(\d+)[.\t]\s*(.+)', line_stripped)
            if match:
                headline_number = int(match.group(1))
                headline_text = match.group(2).strip()
                print(f"  Line {i}: Found numbered headline #{headline_number}")
            else:
                # Unnumbered format - treat entire line as headline
                headline_text = line_stripped
                headline_number = len(headlines) + 1
                print(f"  Line {i}: Found unnumbered headline (will be #{headline_number})")

            # Validate headline length (must be substantial text, not a section marker)
            if len(headline_text) > 30:
                headlines.append({
                    'id': len(headlines),  # 0-indexed
                    'number': headline_number,
                    'text': headline_text
                })
                print(f"    ✓ Added: '{headline_text[:60]}...'")
            else:
                print(f"    ✗ Too short (< 30 chars): '{headline_text}'")

    print(f"=== TOTAL HEADLINES FOUND: {len(headlines)} ===\n")
    return headlines


def parse_bullets_from_template(template_text):
    """Extract all bullet points from the experience bullets section in master template

    Template structure:
    - Section header: "EXPERIENCE BULLETS" (generic) or legacy "JP MORGAN CHASE BULLETS (XX Unique Variations)"
    - Category markers (optional): "CATEGORY: [Name]"
    - Bullets: Plain text paragraphs (no prefix markers)
    """
    import re

    bullets = []
    lines = template_text.split('\n')

    in_bullets_section = False
    current_category = "Uncategorized"

    print(f"\n=== PARSING EXPERIENCE BULLETS ===")
    print(f"Total lines in template: {len(lines)}")

    for i, line in enumerate(lines):
        line_stripped = line.strip()

        # Detect start of bullets section
        # Accepts generic format: "EXPERIENCE BULLETS"
        # Also accepts legacy format: "JP MORGAN CHASE BULLETS (XX Unique Variations)"
        upper_stripped = line_stripped.upper()
        is_generic_header = 'EXPERIENCE BULLETS' in upper_stripped
        is_legacy_header = ('JP MORGAN' in upper_stripped and
                            'BULLETS' in upper_stripped and
                            'VARIATION' in upper_stripped)

        if is_generic_header or is_legacy_header:
            in_bullets_section = True
            print(f"✓ Found BULLETS section at line {i}: '{line_stripped}'")
            continue

        # Only process if we're in the bullets section
        if not in_bullets_section:
            continue

        # Skip empty lines
        if not line_stripped:
            continue

        # Extract category markers
        if line_stripped.startswith('CATEGORY:'):
            current_category = line_stripped.replace('CATEGORY:', '').strip()
            print(f"  ✓ Line {i}: Found category '{current_category}'")
            continue

        # Check if we've hit the end of the bullets section
        # End conditions: next major section header or separator
        upper_line = line_stripped.upper()

        # Separator lines
        if line_stripped.startswith('___') or line_stripped.startswith('==='):
            print(f"  ✗ Line {i}: End of BULLETS section (separator line)")
            break

        # All caps section headers (but not CATEGORY markers or JP MORGAN related)
        if line_stripped.isupper() and len(line_stripped) > 15:
            if not any(keyword in upper_line for keyword in ['JP MORGAN', 'JPMORGAN', 'CHASE', 'CATEGORY']):
                print(f"  ✗ Line {i}: End of BULLETS section (section header): '{line_stripped}'")
                break

        # If we get here, this should be a bullet
        # Validate it's substantial text (bullets are detailed paragraphs)
        if len(line_stripped) > 40:
            bullets.append({
                'id': len(bullets),  # 0-indexed
                'number': len(bullets) + 1,  # 1-indexed for display
                'text': line_stripped,
                'category': current_category
            })
            print(f"  ✓ Line {i}: Added bullet #{len(bullets)} ({current_category})")
        else:
            print(f"  ⚠ Line {i}: Skipped (too short, {len(line_stripped)} chars): '{line_stripped}'")

    print(f"\n=== TOTAL BULLETS FOUND: {len(bullets)} ===")

    if len(bullets) > 0:
        category_counts = {}
        for bullet in bullets:
            cat = bullet['category']
            category_counts[cat] = category_counts.get(cat, 0) + 1
        print("\n=== CATEGORY BREAKDOWN ===")
        for cat, count in category_counts.items():
            print(f"  {cat}: {count} bullets")
    print()

    return bullets


def parse_career_summaries_from_template(template_text):
    """Parse career summary versions from template.

    Handles the key format variants Claude/Haiku may produce:
      [VERSION 1 — Label]  or  [VERSION 1 - Label]
      VERSION 1: Label     or  VERSION 1 — Label
    Also skips the === separator lines that immediately follow the section header.
    """
    import re
    lines = template_text.split('\n')
    summaries = []
    in_summary = False
    found_content = False   # True once we've moved past the opening separators
    current_version = None
    current_text = []
    auto_counter = 0

    for line in lines:
        stripped = line.strip()

        if not in_summary:
            # Accept lines like "CAREER SUMMARY [MULTIPLE VERSIONS...]" (no = chars in the line)
            if any(kw in stripped.upper() for kw in
                   ['CAREER SUMMARY', 'PROFESSIONAL SUMMARY', 'EXECUTIVE SUMMARY']):
                in_summary = True
                found_content = False
            continue

        # Skip empty lines and === separators until we hit real content
        is_separator = stripped.startswith('=') and len(stripped) >= 5
        if not found_content:
            if is_separator or not stripped:
                continue
            found_content = True

        # A === line after content marks the end of the section
        if is_separator:
            if current_version and current_text:
                summaries.append({**current_version, 'text': ' '.join(current_text).strip()})
            break

        # --- Version marker detection (multiple formats) ---
        vm = (re.match(r'\[VERSION\s+(\d+)\s*[—–\-:]+\s*(.+?)\]?$', stripped, re.IGNORECASE) or
              re.match(r'VERSION\s+(\d+)\s*[—–\-:]+\s*(.+)',          stripped, re.IGNORECASE) or
              re.match(r'\[V(\d+)\s*[—–\-:]+\s*(.+?)\]?$',            stripped, re.IGNORECASE))

        if vm:
            if current_version and current_text:
                summaries.append({**current_version, 'text': ' '.join(current_text).strip()})
            n = int(vm.group(1))
            label = vm.group(2).strip().rstrip(']').strip()
            current_version = {'id': n - 1, 'number': n, 'label': label}
            current_text = []
        elif current_version and stripped:
            current_text.append(stripped)
        elif not current_version and stripped and len(stripped) > 40:
            # No version marker — treat each long paragraph as its own version
            auto_counter += 1
            if current_text:
                summaries.append({'id': auto_counter - 2, 'number': auto_counter - 1,
                                   'label': f'Version {auto_counter - 1}',
                                   'text': ' '.join(current_text).strip()})
                current_text = []
            current_version = {'id': auto_counter - 1, 'number': auto_counter,
                                'label': f'Version {auto_counter}'}
            current_text = [stripped]

    if current_version and current_text:
        summaries.append({**current_version, 'text': ' '.join(current_text).strip()})

    print(f"=== CAREER SUMMARIES FOUND: {len(summaries)} ===")
    return summaries


def parse_roles_from_template(template_text):
    """Parse role sections from employment history.
    Each role block is delimited by === lines and contains
    ROLE TITLES, CONTEXT LINES, BULLETS, KEY ACHIEVEMENTS.
    """
    import re
    lines = template_text.split('\n')
    roles = []
    in_employment = False
    i = 0

    while i < len(lines):
        stripped = lines[i].strip()

        if not in_employment:
            if 'EMPLOYMENT HISTORY' in stripped.upper() and stripped.startswith('='):
                in_employment = True
            i += 1
            continue

        if 'PERMANENT ROLES' in stripped.upper():
            break

        # Detect opening === block
        if stripped.startswith('=') and len(stripped) >= 20:
            # Next non-empty line is the header
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1

            if j >= len(lines):
                i += 1
                continue

            header = lines[j].strip()

            # Skip pure separators and known non-role headers
            if (not header or header.startswith('=') or
                    'EMPLOYMENT HISTORY' in header.upper() or
                    'PERMANENT ROLES' in header.upper() or
                    'CONTRACTING ASSIGNMENTS' in header.upper()):
                i += 1
                continue

            # Extract company and dates
            date_match = re.search(
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|'
                r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})\s*[–—-]\s*'
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|Present)',
                header
            )
            company = header
            dates = ''
            if date_match:
                dates = header[date_match.start():date_match.end()].strip()
                company = header[:date_match.start()].strip()

            role = {
                'id': len(roles),
                'company': company,
                'dates': dates,
                'role_titles': [],
                'context_lines': [],
                'bullets': [],
                'key_achievements': []
            }

            # Skip past the closing === line
            i = j + 1
            while i < len(lines) and (not lines[i].strip() or lines[i].strip().startswith('=')):
                i += 1

            # Parse sections
            current_section = None
            while i < len(lines):
                sec_line = lines[i].strip()

                if sec_line.startswith('=') and len(sec_line) >= 20:
                    break
                if 'PERMANENT ROLES' in sec_line.upper():
                    break

                if sec_line == 'ROLE TITLES:':
                    current_section = 'role_titles'
                elif sec_line == 'CONTEXT LINES:':
                    current_section = 'context_lines'
                elif sec_line == 'BULLETS:':
                    current_section = 'bullets'
                elif sec_line == 'KEY ACHIEVEMENTS:':
                    current_section = 'key_achievements'
                elif current_section and sec_line.startswith('- '):
                    content = sec_line[2:].strip()
                    if content and len(content) > 10:
                        role[current_section].append(content)

                i += 1

            if role['bullets']:
                roles.append(role)
                print(f"  ✓ Role: {role['company']} ({len(role['bullets'])} bullets)")
            continue

        i += 1

    print(f"=== ROLES FOUND: {len(roles)} ===")
    return roles


def score_roles_relevance(roles, job_description):
    """Fast keyword-based relevance scoring — no API call needed."""
    import re
    jd_lower = job_description.lower()
    scored = []

    for role in roles:
        all_text = ' '.join(role['bullets'] + role['role_titles'] + role['context_lines']).lower()
        words = set(re.sub(r'[^\w\s]', ' ', all_text).split())
        jd_words = set(re.sub(r'[^\w\s]', ' ', jd_lower).split())

        # Remove stop words
        stops = {'the', 'a', 'an', 'and', 'or', 'of', 'to', 'in', 'for', 'with',
                 'on', 'at', 'by', 'as', 'is', 'was', 'are', 'were', 'be', 'been',
                 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                 'should', 'may', 'might', 'this', 'that', 'these', 'those', 'it'}
        words -= stops
        jd_words -= stops

        overlap = words & jd_words
        score = min(99, int((len(overlap) / max(len(jd_words), 1)) * 250))
        score = max(30, score)

        scored.append({**role, 'relevance_score': score})

    scored.sort(key=lambda r: r['relevance_score'], reverse=True)
    return scored


def create_cv_customization_session(user_id, analysis_id, job_title, job_company):
    """Create a new CV customization session"""
    engine = get_db_connection()
    if not engine:
        return None

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                INSERT INTO cv_customization_sessions
                (user_id, analysis_id, job_title, job_company, status)
                VALUES (:user_id, :analysis_id, :job_title, :job_company, 'in_progress')
                RETURNING id
            """), {
                "user_id": user_id,
                "analysis_id": analysis_id,
                "job_title": job_title,
                "job_company": job_company
            })

            session_id = result.fetchone()[0]
            conn.commit()
            return session_id

    except Exception as e:
        print(f"Error creating CV customization session: {e}")
        return None


def get_cv_session(session_id):
    """Get CV customization session data"""
    engine = get_db_connection()
    if not engine:
        return None

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT id, user_id, analysis_id, job_title, job_company,
                       selected_headline, bullet_analysis, approved_bullets, new_bullets,
                       match_score_progression, status, created_at
                FROM cv_customization_sessions
                WHERE id = :session_id
            """), {"session_id": session_id})

            row = result.fetchone()
            if row:
                # Parse JSON fields
                approved_bullets = row[7]
                if isinstance(approved_bullets, str):
                    approved_bullets = json.loads(approved_bullets) if approved_bullets else []
                elif approved_bullets is None:
                    approved_bullets = []

                new_bullets = row[8]
                if isinstance(new_bullets, str):
                    new_bullets = json.loads(new_bullets) if new_bullets else []
                elif new_bullets is None:
                    new_bullets = []

                match_score = row[9]
                if isinstance(match_score, str):
                    match_score = json.loads(match_score) if match_score else {}
                elif match_score is None:
                    match_score = {}

                return {
                    'id': row[0],
                    'user_id': row[1],
                    'analysis_id': row[2],
                    'job_title': row[3],
                    'job_company': row[4],
                    'selected_headline': row[5],
                    'bullet_analysis': row[6],
                    'approved_bullets': approved_bullets,
                    'new_bullets': new_bullets,
                    'match_score_progression': match_score,
                    'status': row[10],
                    'created_at': row[11]
                }
            return None

    except Exception as e:
        print(f"Error getting CV session: {e}")
        return None


def get_cv_session_by_analysis(analysis_id, user_id):
    """Get CV customization session by analysis_id and user_id"""
    engine = get_db_connection()
    if not engine:
        return None

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT id, user_id, analysis_id, job_title, job_company,
                       selected_headline, bullet_analysis, approved_bullets, new_bullets,
                       match_score_progression, status, created_at
                FROM cv_customization_sessions
                WHERE analysis_id = :analysis_id AND user_id = :user_id
                ORDER BY created_at DESC
                LIMIT 1
            """), {"analysis_id": analysis_id, "user_id": user_id})

            row = result.fetchone()
            if row:
                return {
                    'id': row[0],
                    'user_id': row[1],
                    'analysis_id': row[2],
                    'job_title': row[3],
                    'job_company': row[4],
                    'selected_headline': row[5],
                    'bullet_analysis': row[6],
                    'approved_bullets': row[7] if row[7] else [],
                    'new_bullets': row[8] if row[8] else [],
                    'match_score_progression': row[9] if row[9] else {},
                    'status': row[10],
                    'created_at': row[11]
                }
            return None

    except Exception as e:
        print(f"Error getting CV session by analysis: {e}")
        return None


def update_cv_session_headline(session_id, headline_text):
    """Update selected headline in CV session"""
    engine = get_db_connection()
    if not engine:
        return False

    try:
        with engine.connect() as conn:
            conn.execute(text("""
                UPDATE cv_customization_sessions
                SET selected_headline = :headline, updated_at = CURRENT_TIMESTAMP
                WHERE id = :session_id
            """), {"session_id": session_id, "headline": headline_text})
            conn.commit()
            return True

    except Exception as e:
        print(f"Error updating headline: {e}")
        return False


def update_cv_session_bullet_analysis(session_id, bullet_analysis):
    """Save bullet analysis to avoid re-running AI"""
    engine = get_db_connection()
    if not engine:
        return False

    try:
        with engine.connect() as conn:
            import json
            conn.execute(text("""
                UPDATE cv_customization_sessions
                SET bullet_analysis = :analysis, updated_at = CURRENT_TIMESTAMP
                WHERE id = :session_id
            """), {"session_id": session_id, "analysis": json.dumps(bullet_analysis)})
            conn.commit()
            return True

    except Exception as e:
        print(f"Error updating bullet analysis: {e}")
        return False


def update_cv_session_approved_bullets(session_id, approved_bullets):
    """Save approved bullets"""
    engine = get_db_connection()
    if not engine:
        return False

    try:
        with engine.connect() as conn:
            import json
            conn.execute(text("""
                UPDATE cv_customization_sessions
                SET approved_bullets = :bullets, updated_at = CURRENT_TIMESTAMP
                WHERE id = :session_id
            """), {"session_id": session_id, "bullets": json.dumps(approved_bullets)})
            conn.commit()
            return True

    except Exception as e:
        print(f"Error updating approved bullets: {e}")
        return False


def update_cv_session_selected_roles(session_id, selected_roles):
    """Save user's chosen roles for bullet tailoring."""
    engine = get_db_connection()
    if not engine:
        return False
    try:
        with engine.connect() as conn:
            import json
            conn.execute(text("""
                UPDATE cv_customization_sessions
                SET selected_roles = :roles, updated_at = CURRENT_TIMESTAMP
                WHERE id = :session_id
            """), {"session_id": session_id, "roles": json.dumps(selected_roles)})
            conn.commit()
            return True
    except Exception as e:
        print(f"Error saving selected roles: {e}")
        return False


def update_cv_session_bullet_analysis_by_role(session_id, role_key, analysis):
    """Save per-role bullet analysis to avoid re-running AI."""
    engine = get_db_connection()
    if not engine:
        return False
    try:
        with engine.connect() as conn:
            import json
            conn.execute(text("""
                UPDATE cv_customization_sessions
                SET bullet_analysis_by_role = COALESCE(bullet_analysis_by_role, '{}'::jsonb)
                    || jsonb_build_object(:role_key, :analysis::jsonb),
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = :session_id
            """), {
                "session_id": session_id,
                "role_key": role_key,
                "analysis": json.dumps(analysis)
            })
            conn.commit()
            return True
    except Exception as e:
        print(f"Error saving role bullet analysis: {e}")
        return False


def analyze_headlines_with_ai(headlines, job_description, user_id):
    """Use AI to analyze headlines and recommend best one"""
    client = get_anthropic_client()

    # Get user's system prompt
    system_prompt = get_user_system_prompt(user_id)

    # Build headlines list for prompt
    headlines_text = "\n".join([f"{i+1}. {h['text']}" for i, h in enumerate(headlines)])

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=2000,
            temperature=0.3,
            messages=[
                {
                    "role": "user",
                    "content": f"""You are helping a job seeker select the best resume headline to maximize their interview chances for this role.

**JOB DESCRIPTION:**
{job_description}

**AVAILABLE HEADLINE OPTIONS:**
{headlines_text}

**YOUR TASK:**
Analyze each headline and:
1. Identify which headline best aligns with the JD requirements
2. Explain why it's the best match (3-4 specific reasons tied to JD)
3. Calculate a match score (0-100) for the top 3 headlines
4. Suggest if the best headline needs any adaptations for this specific role

Return your analysis as JSON:
{{
    "recommended_headline_index": 2,
    "top_3": [
        {{
            "index": 2,
            "match_score": 85,
            "reasons": [
                "Emphasizes AI innovation which matches JD requirement for AI-driven initiatives",
                "Leadership positioning aligns with transformation lead role",
                "Financial services background directly relevant"
            ],
            "suggested_adaptation": null
        }},
        {{
            "index": 4,
            "match_score": 78,
            "reasons": ["...", "...", "..."],
            "suggested_adaptation": "Change 'Program Manager' to 'Transformation Lead' to match job title"
        }},
        {{
            "index": 13,
            "match_score": 72,
            "reasons": ["...", "...", "..."],
            "suggested_adaptation": null
        }}
    ],
    "other_headlines": [
        {{
            "index": 1,
            "match_score": 45,
            "weakness": "Too focused on risk/compliance, doesn't emphasize transformation"
        }},
        ...
    ]
}}

Be strategic and honest about which headlines will maximize interview chances."""
                }
            ]
        )

        response_text = message.content[0].text

        # Extract JSON from response
        import re
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            analysis = json.loads(json_match.group())
            return analysis

        return None

    except Exception as e:
        print(f"Error analyzing headlines with AI: {e}")
        import traceback
        traceback.print_exc()
        return None


def analyze_bullets_with_ai(bullets, job_description, user_id):
    """Use AI to analyze bullets and recommend top 8-10 for this role

    Returns analysis with:
    - Top 8-10 recommended bullets with match scores
    - For each: status (ready_to_use or needs_rewriting) + suggestions
    - Gap analysis: what's missing from bullets vs JD requirements
    - 2-3 suggested new bullets to fill gaps
    """
    client = get_anthropic_client()
    system_prompt = get_user_system_prompt(user_id)

    # Group bullets by category if categories exist, otherwise use flat list
    has_categories = any(bullet['category'] != 'Uncategorized' for bullet in bullets)

    if has_categories:
        bullets_by_category = {}
        for bullet in bullets:
            cat = bullet['category']
            if cat not in bullets_by_category:
                bullets_by_category[cat] = []
            bullets_by_category[cat].append(bullet)

        # Build bullets text organized by category
        bullets_text = ""
        for category, cat_bullets in bullets_by_category.items():
            bullets_text += f"\n**{category}:**\n"
            for bullet in cat_bullets:
                bullets_text += f"{bullet['number']}. {bullet['text']}\n"
    else:
        # No categories — flat numbered list
        bullets_text = "\n"
        for bullet in bullets:
            bullets_text += f"{bullet['number']}. {bullet['text']}\n"

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=4000,
            temperature=0.3,
            system=[
                {
                    "type": "text",
                    "text": system_prompt,
                    "cache_control": {"type": "ephemeral"}
                }
            ],
            messages=[
                {
                    "role": "user",
                    "content": f"""You are helping a job seeker select the most impactful resume bullets to maximize interview chances for this role.

**JOB DESCRIPTION:**
{job_description}

**AVAILABLE BULLETS ({len(bullets)} total{', organized by category' if has_categories else ''}):**
{bullets_text}

**YOUR TASK:**
1. **Score all bullets (0-100)** based on relevance to JD requirements
2. **Select top 8-10 bullets** that:
   - Best demonstrate required skills/experience
   - Show quantifiable impact matching JD priorities
   - {'Cover diverse aspects of the role (not all from one category)' if has_categories else 'Cover diverse aspects of the role (vary the skills and experience shown)'}
   - Tell a compelling story about candidate's fit

3. **CRITICAL: Order bullets by priority** - Return bullets in descending order by match_score (highest first)
   - First bullet should be the strongest match (highest score)
   - This ensures recruiters see most relevant experience first

4. **For each top bullet**, determine:
   - **ready_to_use**: bullet is perfect as-is, use directly
   - **needs_rewriting**: bullet is relevant but needs adaptation
     - If needs rewriting, provide specific rewrite suggestion

4. **Gap Analysis**: Identify 2-3 JD requirements NOT well-covered by existing bullets

5. **Suggest 2-3 new bullets** to fill those gaps (write complete bullet text)

**Return JSON in this exact format:**
{{
    "recommended_bullets": [
        {{
            "bullet_number": 46,
            "category": "Portfolio Scale & Governance Oversight",
            "original_text": "Scaled AI Governance Oversight by 63%...",
            "match_score": 95,
            "status": "ready_to_use",
            "reasons": [
                "Directly demonstrates portfolio management scale required in JD",
                "Shows 63% growth metric that proves impact",
                "Governance oversight aligns with compliance focus"
            ],
            "rewrite_suggestion": null
        }},
        {{
            "bullet_number": 60,
            "category": "LLM Innovation",
            "original_text": "Pioneered LLM-powered solution...",
            "match_score": 88,
            "status": "needs_rewriting",
            "reasons": [
                "LLM innovation matches AI transformation focus in JD",
                "Shows hands-on technical implementation experience"
            ],
            "rewrite_suggestion": "Pioneered LLM-powered automation solution using ChatGPT to transform manual workflows into comprehensive documentation, reducing manual effort by 5-10 hours per procedure—demonstrating AI transformation impact aligned with [Company]'s AI-first strategy"
        }},
        ... (8-10 total)
    ],
    "gaps": [
        "JD emphasizes stakeholder communication at C-suite level - existing bullets show stakeholder mgmt but could strengthen executive influence",
        "JD requires experience with [specific tool/framework] not explicitly mentioned in bullets",
        "JD wants change management leadership - only 1 bullet covers this"
    ],
    "suggested_new_bullets": [
        "Led C-suite stakeholder alignment across 5 business units for AI governance framework adoption, securing executive sponsorship and $2M budget approval through compelling ROI presentations that translated technical requirements into business value",
        "Drove organization-wide change management for new AI compliance standards affecting 200+ team members across Operations and Technology, designing training curriculum and achieving 95% adoption within 3 months"
    ]
}}

**Important:**
- Be strategic: select bullets that show DIVERSE skills, not all from one category
- Prioritize bullets with metrics/quantifiable impact
- For "needs_rewriting", give SPECIFIC rewrites that incorporate JD keywords naturally
- New bullets should fill real gaps, not just restate existing bullets
- Match scores should reflect true JD alignment (be honest, not inflated)"""
                }
            ]
        )

        # Parse JSON response
        import json
        response_text = message.content[0].text

        # Extract JSON from response (handle markdown code blocks)
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        elif "```" in response_text:
            json_start = response_text.find("```") + 3
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()

        analysis = json.loads(response_text)

        # IMPORTANT: Sort bullets by match_score (highest first) to ensure priority ordering
        # This catches recruiter's eye with strongest matches first
        if 'recommended_bullets' in analysis:
            analysis['recommended_bullets'] = sorted(
                analysis['recommended_bullets'],
                key=lambda x: x.get('match_score', 0),
                reverse=True
            )
            print(f"✓ Bullet analysis complete: {len(analysis.get('recommended_bullets', []))} bullets recommended")
            print(f"  Sorted by priority - Top match score: {analysis['recommended_bullets'][0].get('match_score', 0)}%")
        else:
            print(f"✓ Bullet analysis complete")

        print(f"  Gaps identified: {len(analysis.get('gaps', []))}")
        print(f"  New bullets suggested: {len(analysis.get('suggested_new_bullets', []))}")

        return analysis

    except Exception as e:
        print(f"Error analyzing bullets with AI: {e}")
        import traceback
        traceback.print_exc()
        return None


def analyze_bullets_for_role_with_ai(role, job_description, user_id):
    """Rank and analyse bullets for a single role against the JD."""
    client = get_anthropic_client()
    system_prompt = get_user_system_prompt(user_id)

    bullets_text = '\n'.join(
        f"{i+1}. {b}" for i, b in enumerate(role['bullets'])
    )
    role_context = f"{role['company']} ({role['dates']})"
    best_title = role['role_titles'][0] if role['role_titles'] else ''

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=3000,
            temperature=0.3,
            system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
            messages=[{
                "role": "user",
                "content": f"""You are helping a job seeker select the best bullets from a specific role on their CV.

**ROLE:** {role_context}
**TITLE:** {best_title}

**JOB DESCRIPTION:**
{job_description}

**AVAILABLE BULLETS FOR THIS ROLE ({len(role['bullets'])} total):**
{bullets_text}

**YOUR TASK:**
Score every bullet (0–100) against the JD. Select the top bullets (aim for 4–6, minimum 3).
Return them ordered highest score first.

For each selected bullet determine:
- ready_to_use: strong as-is
- needs_rewriting: relevant but needs adaptation — provide a specific rewrite

Return JSON:
{{
    "recommended_bullets": [
        {{
            "bullet_index": 0,
            "original_text": "...",
            "match_score": 92,
            "status": "ready_to_use",
            "reasons": ["Directly addresses JD requirement for...", "Shows measurable impact..."],
            "rewrite_suggestion": null
        }},
        {{
            "bullet_index": 2,
            "original_text": "...",
            "match_score": 78,
            "status": "needs_rewriting",
            "reasons": ["Relevant but lacks JD keywords..."],
            "rewrite_suggestion": "Revised bullet text here..."
        }}
    ]
}}"""
            }]
        )

        response_text = message.content[0].text
        if "```json" in response_text:
            s = response_text.find("```json") + 7
            response_text = response_text[s:response_text.find("```", s)].strip()
        elif "```" in response_text:
            s = response_text.find("```") + 3
            response_text = response_text[s:response_text.find("```", s)].strip()

        analysis = json.loads(response_text)
        if 'recommended_bullets' in analysis:
            analysis['recommended_bullets'].sort(key=lambda x: x.get('match_score', 0), reverse=True)
        return analysis

    except Exception as e:
        print(f"Error analysing bullets for role {role['company']}: {e}")
        import traceback
        traceback.print_exc()
        return None


def generate_interview_questions(job_description, cv_bullets, selected_headline, user_id):
    """
    Generate 8 interview questions based on CV and job description
    Returns list of question objects
    """
    client = get_anthropic_client()

    # Format bullets for prompt
    bullets_text = "\n".join([f"• {bullet}" for bullet in cv_bullets])

    system_prompt = """You are an expert technical interviewer at a top-tier company.

GENERATE 8 realistic interview questions that test if the candidate can explain their resume.

QUESTION STRATEGY:
- Questions 1-6: About specific CV bullets (test depth, trade-offs, decisions)
- Questions 7-8: Behavioral questions mapped to their experience

QUALITY REQUIREMENTS:
- SPECIFIC to their resume (reference actual achievements/metrics)
- PROGRESSIVE DIFFICULTY (start broad, get deeper)
- TEST TRADE-OFFS and decision-making ("Why did you choose X over Y?")
- PROBE OWNERSHIP (distinguish "I did" vs "we did")

OUTPUT: Return ONLY a JSON array, no markdown fences:
[
  {
    "question_number": 1,
    "question_text": "...",
    "type": "cv_bullet",
    "bullet_reference": "..."
  },
  ...
]"""

    user_prompt = f"""Generate 8 interview questions for this candidate.

**CANDIDATE'S HEADLINE:**
{selected_headline}

**CANDIDATE'S KEY EXPERIENCE:**
{bullets_text}

**JOB DESCRIPTION:**
{job_description}

Generate questions that probe:
- Technical depth (not just what they did, but HOW and WHY)
- Problem-solving approach
- Ownership and impact
- Relevant skills from the JD

Return JSON array only (no markdown)."""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=2048,
            temperature=0.3,
            system=[{
                "type": "text",
                "text": system_prompt
            }],
            messages=[{
                "role": "user",
                "content": user_prompt
            }]
        )

        response_text = message.content[0].text.strip()

        # Clean up JSON response (remove markdown fences if present)
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]

        questions = json.loads(response_text.strip())

        print(f"✓ Generated {len(questions)} interview questions")
        return questions

    except Exception as e:
        print(f"Error generating interview questions: {e}")
        import traceback
        traceback.print_exc()
        return None


def evaluate_interview_answer(question_text, user_answer, job_description, cv_bullets):
    """
    Evaluate a single interview answer
    Returns evaluation with scores and feedback
    """
    client = get_anthropic_client()

    bullets_text = "\n".join([f"• {bullet}" for bullet in cv_bullets])

    system_prompt = """You are a senior recruiter evaluating interview answers.

EVALUATE on 4 dimensions (1-5 scale):
1. **Details** - Concrete specifics (metrics, names, numbers, tools)
2. **Organization** - Clear STAR format (Situation, Task, Action, Result)
3. **Analysis** - Explains WHY, trade-offs, alternatives considered
4. **Ownership** - Clear "I did X" vs vague "we/team"

PROVIDE:
- Scores (1-5 each)
- Feedback summary (2-3 sentences)
- Strong points (2-3 bullet points)
- Improvement areas (2-3 bullet points)
- Better answer example (using their actual experience, not made up)

OUTPUT: Return ONLY JSON, no markdown fences:
{
  "details_score": 1-5,
  "organization_score": 1-5,
  "analysis_score": 1-5,
  "ownership_score": 1-5,
  "feedback_summary": "...",
  "strong_points": ["...", "..."],
  "improvement_areas": ["...", "..."],
  "better_answer_example": "..."
}"""

    user_prompt = f"""Evaluate this interview answer.

**QUESTION:** {question_text}

**CANDIDATE'S ANSWER:** {user_answer}

**CONTEXT - Their CV:**
{bullets_text}

**CONTEXT - Job Requirements:**
{job_description[:1000]}...

Evaluate rigorously but fairly. Return JSON only."""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=1536,
            temperature=0.2,
            system=[{
                "type": "text",
                "text": system_prompt
            }],
            messages=[{
                "role": "user",
                "content": user_prompt
            }]
        )

        response_text = message.content[0].text.strip()

        # Clean up JSON response
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]

        evaluation = json.loads(response_text.strip())

        # Calculate overall score (1-10 scale)
        avg_score = (
            evaluation['details_score'] +
            evaluation['organization_score'] +
            evaluation['analysis_score'] +
            evaluation['ownership_score']
        ) / 4.0

        evaluation['overall_score'] = round(avg_score * 2, 1)  # Convert to 1-10 scale

        print(f"✓ Evaluated answer: {evaluation['overall_score']}/10")
        return evaluation

    except Exception as e:
        print(f"Error evaluating answer: {e}")
        import traceback
        traceback.print_exc()
        return None


# ==================== AI MATCH TOOL: AI INTEGRATION FUNCTIONS ====================

def get_anthropic_client():
    """Initialize Anthropic client"""
    try:
        import anthropic
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            print("ERROR: ANTHROPIC_API_KEY not found in environment variables")
            raise Exception("ANTHROPIC_API_KEY environment variable is not set. Please add it in Railway settings.")
        return anthropic.Anthropic(api_key=api_key)
    except ImportError as e:
        print(f"Anthropic library not installed: {e}")
        raise Exception("Anthropic library not installed. Please check Railway build logs.")
    except Exception as e:
        print(f"Error initializing Anthropic client: {e}")
        raise


def get_user_system_prompt(user_id):
    """
    Get the active system prompt for a user.
    Returns the custom prompt if set, otherwise the selected template prompt,
    or falls back to the default template.
    """
    engine = get_db_connection()
    if not engine:
        # Fallback to basic prompt if no database connection
        return """You are an expert career coach and recruiter. Analyze how well this candidate's CV matches the job description.
Provide a match score (0-100) and detailed feedback on skills match, missing skills, and recommendations."""

    try:
        with engine.connect() as conn:
            # First, check if user has an active preference with custom prompt
            result = conn.execute(text("""
                SELECT upp.custom_prompt_text, pt.prompt_text
                FROM user_prompt_preferences upp
                LEFT JOIN prompt_templates pt ON upp.template_id = pt.id
                WHERE upp.user_id = :user_id AND upp.is_active = TRUE
                ORDER BY upp.created_at DESC
                LIMIT 1
            """), {"user_id": user_id})

            row = result.fetchone()

            if row:
                # If user has custom prompt, use it; otherwise use template
                custom_prompt = row[0]
                template_prompt = row[1]

                if custom_prompt:
                    return custom_prompt
                elif template_prompt:
                    return template_prompt

            # If no user preference found, get the default template
            default_result = conn.execute(text("""
                SELECT prompt_text FROM prompt_templates WHERE is_default = TRUE LIMIT 1
            """))

            default_row = default_result.fetchone()
            if default_row:
                return default_row[0]

            # Ultimate fallback (should never happen if seeding worked)
            return """You are an expert career coach and recruiter. Analyze how well this candidate's CV matches the job description.
Provide a match score (0-100) and detailed feedback on skills match, missing skills, and recommendations."""

    except Exception as e:
        print(f"Error getting user system prompt: {e}")
        # Return basic fallback prompt
        return """You are an expert career coach and recruiter. Analyze how well this candidate's CV matches the job description.
Provide a match score (0-100) and detailed feedback on skills match, missing skills, and recommendations."""


def extract_job_info_from_posting(job_posting):
    """Extract job title and company from full job posting using AI"""
    client = get_anthropic_client()

    try:
        prompt = f"""Extract the job title and company name from this job posting.

JOB POSTING:
{job_posting}

Respond ONLY with valid JSON (no markdown, no code blocks). Use this exact structure:
{{
    "job_title": "The job title",
    "company": "Company name"
}}

If you cannot find the company name, use "Not specified". Return ONLY the JSON object."""

        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=200,
            temperature=0.1,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        response_text = message.content[0].text.strip()

        # Track usage
        track_ai_usage(
            feature_type='job_info_extraction',
            tokens_input=message.usage.input_tokens,
            tokens_output=message.usage.output_tokens
        )

        # Parse JSON
        import json
        import re

        # Handle markdown code blocks
        if response_text.startswith('```'):
            match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', response_text, re.DOTALL)
            if match:
                response_text = match.group(1).strip()

        job_info = json.loads(response_text)
        return job_info['job_title'], job_info.get('company', 'Not specified')

    except Exception as e:
        print(f"Error extracting job info: {e}")
        # Fallback: try to extract first line as title
        lines = job_posting.strip().split('\n')
        job_title = lines[0][:200] if lines else "Job Position"
        return job_title, "Not specified"


def analyze_job_match_with_ai(cv_text, job_title, job_company, job_description, user_id):
    """Use Claude AI to analyze CV-to-job match using user's selected prompt template"""
    # This will raise an exception if API key is not set
    client = get_anthropic_client()

    try:
        # Get user's system prompt from their template selection
        system_prompt = get_user_system_prompt(user_id)

        # Build full prompt with CV and job details
        prompt = f"""{system_prompt}

**CANDIDATE'S CV:**
{cv_text}

**JOB DETAILS:**
Title: {job_title}
Company: {job_company or 'Not specified'}

**JOB DESCRIPTION:**
{job_description}

**INSTRUCTIONS:**
Respond ONLY with valid JSON (no markdown, no code blocks, no other text). Use this exact structure:

{{
    "match_score": <integer 0-100>,
    "skills_match": [
        {{
            "skill": "Skill name",
            "evidence": "Where in CV this is demonstrated",
            "strength": "strong|moderate|basic"
        }}
    ],
    "skills_missing": [
        {{
            "skill": "Required skill name",
            "importance": "required|preferred|nice-to-have",
            "impact": "Brief explanation of gap impact"
        }}
    ],
    "experience_match": [
        "Bullet point of matching experience"
    ],
    "experience_gaps": [
        "Bullet point of experience gaps"
    ],
    "recommendations": [
        {{
            "priority": "high|medium|low",
            "recommendation": "Specific actionable advice"
        }}
    ],
    "overall_assessment": "2-3 sentence summary of candidacy strength"
}}

IMPORTANT: Return ONLY the JSON object above. No explanations, no markdown formatting, just pure JSON. Be honest and objective."""

        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=4000,
            temperature=0.3,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        response_text = message.content[0].text

        # Track usage
        track_ai_usage(
            feature_type='job_match_analysis',
            tokens_input=message.usage.input_tokens,
            tokens_output=message.usage.output_tokens
        )

        # Parse JSON response
        import json
        import re

        # Extract JSON from response (handle markdown code blocks)
        json_text = response_text.strip()

        # Check if response is wrapped in markdown code blocks
        if json_text.startswith('```'):
            # Extract JSON from markdown code block
            match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', json_text, re.DOTALL)
            if match:
                json_text = match.group(1).strip()

        try:
            analysis = json.loads(json_text)
        except json.JSONDecodeError as json_err:
            print(f"JSON parsing error: {json_err}")
            print(f"Raw response (first 1000 chars): {response_text[:1000]}")
            # Return more helpful error to user
            raise Exception(f"AI response was not valid JSON. This might be a prompt issue. First 300 chars: {response_text[:300]}")

        return analysis

    except Exception as e:
        print(f"Error analyzing job match: {e}")
        import traceback
        traceback.print_exc()
        raise  # Re-raise the exception instead of returning None


def analyze_job_match_with_master_template(template_text, job_title, job_company, job_description, user_id):
    """
    Use Claude AI to analyze job match using master template with PROMPT CACHING.
    This significantly reduces API costs by caching the master resume template.
    """
    client = get_anthropic_client()

    try:
        # Get user's system prompt from their template selection
        system_prompt = get_user_system_prompt(user_id)

        # Build the prompt with caching - the master template is cached
        # Reference: https://docs.anthropic.com/en/docs/build-with-claude/prompt-caching

        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": system_prompt,
                        "cache_control": {"type": "ephemeral"}  # Cache the system prompt
                    },
                    {
                        "type": "text",
                        "text": f"""**CANDIDATE'S MASTER RESUME TEMPLATE:**
{template_text}""",
                        "cache_control": {"type": "ephemeral"}  # Cache the master template (most important!)
                    },
                    {
                        "type": "text",
                        "text": f"""**JOB DETAILS:**
Title: {job_title}
Company: {job_company or 'Not specified'}

**JOB DESCRIPTION:**
{job_description}

**YOUR TASK:**
From the master resume template above, identify which headline variations and which bullet points best match this specific job description. Provide a comprehensive match analysis.

**INSTRUCTIONS:**
Respond ONLY with valid JSON (no markdown, no code blocks, no other text). Use this exact structure:

{{
    "match_score": <integer 0-100>,
    "skills_match": [
        {{
            "skill": "Skill name",
            "evidence": "Where in resume template this is demonstrated",
            "strength": "strong|moderate|basic"
        }}
    ],
    "skills_missing": [
        {{
            "skill": "Required skill name",
            "importance": "required|preferred|nice-to-have",
            "impact": "Brief explanation of gap impact"
        }}
    ],
    "experience_match": [
        "Bullet point of matching experience from template"
    ],
    "experience_gaps": [
        "Bullet point of experience gaps"
    ],
    "recommendations": [
        {{
            "priority": "high|medium|low",
            "recommendation": "Specific actionable advice (e.g., which bullets to emphasize)"
        }}
    ],
    "overall_assessment": "2-3 sentence summary of candidacy strength"
}}

IMPORTANT: Return ONLY the JSON object above. No explanations, no markdown formatting, just pure JSON. Be honest and objective."""
                    }
                ]
            }
        ]

        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=4000,
            temperature=0.3,
            messages=messages
        )

        response_text = message.content[0].text

        # Track usage (prompt caching will show reduced costs!)
        track_ai_usage(
            feature_type='job_match_analysis_cached',
            tokens_input=message.usage.input_tokens,
            tokens_output=message.usage.output_tokens
        )

        # Log cache performance for monitoring
        if hasattr(message.usage, 'cache_creation_input_tokens'):
            print(f"📊 Cache stats - Creation: {message.usage.cache_creation_input_tokens}, "
                  f"Read: {getattr(message.usage, 'cache_read_input_tokens', 0)}, "
                  f"Regular: {message.usage.input_tokens}")

        # Parse JSON response
        import json
        import re

        # Extract JSON from response (handle markdown code blocks)
        json_text = response_text.strip()

        # Check if response is wrapped in markdown code blocks
        if json_text.startswith('```'):
            # Extract JSON from markdown code block
            match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', json_text, re.DOTALL)
            if match:
                json_text = match.group(1).strip()

        try:
            analysis = json.loads(json_text)
        except json.JSONDecodeError as json_err:
            print(f"JSON parsing error: {json_err}")
            print(f"Raw response (first 1000 chars): {response_text[:1000]}")
            # Return more helpful error to user
            raise Exception(f"AI response was not valid JSON. This might be a prompt issue. First 300 chars: {response_text[:300]}")

        return analysis

    except Exception as e:
        print(f"Error analyzing job match with master template: {e}")
        import traceback
        traceback.print_exc()
        raise  # Re-raise the exception instead of returning None


def track_ai_usage(feature_type, tokens_input, tokens_output):
    """Track AI API usage for cost monitoring"""
    user_id = get_current_user_id()
    if not user_id:
        return

    engine = get_db_connection()
    if not engine:
        return

    try:
        # Calculate estimated cost (Claude 3.5 Sonnet pricing)
        # Input: $3 per million tokens, Output: $15 per million tokens
        cost_input = (tokens_input / 1_000_000) * 3.0
        cost_output = (tokens_output / 1_000_000) * 15.0
        total_cost = cost_input + cost_output

        with engine.connect() as conn:
            conn.execute(text("""
                INSERT INTO ai_usage_tracking
                (user_id, feature_type, tokens_input, tokens_output, estimated_cost)
                VALUES (:user_id, :feature_type, :tokens_input, :tokens_output, :estimated_cost)
            """), {
                "user_id": user_id,
                "feature_type": feature_type,
                "tokens_input": tokens_input,
                "tokens_output": tokens_output,
                "estimated_cost": total_cost
            })
            conn.commit()
    except Exception as e:
        print(f"Error tracking AI usage: {e}")


def save_job_analysis(user_id, cv_id, job_title, job_company, job_description, analysis_result):
    """Save job matching analysis to database"""
    print(f"\n{'='*60}")
    print(f"SAVE_JOB_ANALYSIS CALLED")
    print(f"{'='*60}")
    print(f"user_id: {user_id}")
    print(f"cv_id: {cv_id}")
    print(f"job_title: {job_title}")
    print(f"job_company: {job_company}")
    print(f"job_description length: {len(job_description) if job_description else 0}")
    print(f"match_score: {analysis_result.get('match_score')}")

    engine = get_db_connection()
    if not engine:
        print("ERROR: No database connection")
        return None

    try:
        import json

        with engine.connect() as conn:
            print("\nExecuting INSERT query...")
            result = conn.execute(text("""
                INSERT INTO job_analyses
                (user_id, cv_id, job_title, job_company, job_description,
                 match_score, skills_match, skills_missing, recommendations, full_analysis)
                VALUES
                (:user_id, :cv_id, :job_title, :job_company, :job_description,
                 :match_score, :skills_match, :skills_missing, :recommendations, :full_analysis)
                RETURNING id
            """), {
                "user_id": user_id,
                "cv_id": cv_id,
                "job_title": job_title,
                "job_company": job_company or '',
                "job_description": job_description,
                "match_score": analysis_result.get('match_score'),
                "skills_match": json.dumps(analysis_result.get('skills_match', [])),
                "skills_missing": json.dumps(analysis_result.get('skills_missing', [])),
                "recommendations": json.dumps(analysis_result.get('recommendations', [])),
                "full_analysis": json.dumps(analysis_result)
            })

            row = result.fetchone()
            print(f"Fetched row: {row}")

            if not row:
                print("ERROR: No row returned from INSERT")
                return None

            analysis_id = row[0]
            print(f"Analysis ID from INSERT: {analysis_id} (type: {type(analysis_id)})")

            print("Committing transaction...")
            conn.commit()
            print(f"Transaction committed successfully")

            # VERIFY the record was actually saved
            print(f"Verifying record with ID {analysis_id} exists...")
            verify = conn.execute(text("SELECT id, job_title, match_score FROM job_analyses WHERE id = :id"), {"id": analysis_id})
            verify_row = verify.fetchone()
            if verify_row:
                print(f"✓ Verification successful: ID={verify_row[0]}, Title={verify_row[1]}, Score={verify_row[2]}")
            else:
                print(f"✗ WARNING: Record not found after commit!")

            print(f"Returning analysis_id: {analysis_id}")
            print(f"{'='*60}\n")
            return analysis_id
    except Exception as e:
        print(f"ERROR in save_job_analysis: {e}")
        import traceback
        traceback.print_exc()
        return None


def get_user_analyses(user_id, limit=10):
    """Get recent job analyses for a user"""
    engine = get_db_connection()
    if not engine:
        return []

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT ja.id, ja.job_title, ja.job_company, ja.match_score,
                       ja.created_at, cv.cv_name
                FROM job_analyses ja
                LEFT JOIN user_cvs cv ON ja.cv_id = cv.id
                WHERE ja.user_id = :user_id
                ORDER BY ja.created_at DESC
                LIMIT :limit
            """), {"user_id": user_id, "limit": limit})

            analyses = []
            for row in result:
                analyses.append({
                    'id': row[0],
                    'job_title': row[1],
                    'job_company': row[2],
                    'match_score': row[3],
                    'created_at': row[4],
                    'cv_name': row[5] if row[5] else 'Master Template'
                })
            return analyses
    except Exception as e:
        print(f"Error getting user analyses: {e}")
        return []


def get_analysis_by_id(analysis_id, user_id):
    """Get full analysis details by ID"""
    print(f"\n{'='*60}")
    print(f"GET_ANALYSIS_BY_ID CALLED")
    print(f"{'='*60}")
    print(f"Requested analysis_id: {analysis_id} (type: {type(analysis_id)})")
    print(f"user_id: {user_id}")

    engine = get_db_connection()
    if not engine:
        print("ERROR: No database connection")
        return None

    try:
        with engine.connect() as conn:
            # First, check if analysis exists at all (without user_id filter)
            print("\nChecking if analysis exists (without user filter)...")
            check = conn.execute(text("SELECT id, user_id FROM job_analyses WHERE id = :id"), {"id": analysis_id})
            check_row = check.fetchone()
            if check_row:
                print(f"✓ Found analysis: ID={check_row[0]}, user_id={check_row[1]}")
            else:
                print(f"✗ No analysis found with ID={analysis_id} in database")

            print(f"\nExecuting main query with user_id filter...")
            result = conn.execute(text("""
                SELECT ja.id, ja.job_title, ja.job_company, ja.job_description,
                       ja.match_score, ja.skills_match, ja.skills_missing,
                       ja.recommendations, ja.full_analysis, ja.created_at,
                       cv.cv_name
                FROM job_analyses ja
                LEFT JOIN user_cvs cv ON ja.cv_id = cv.id
                WHERE ja.id = :analysis_id AND ja.user_id = :user_id
            """), {"analysis_id": analysis_id, "user_id": user_id})

            row = result.fetchone()
            if not row:
                print(f"✗ No row returned from query")
                print(f"  Possible reasons:")
                print(f"  1. Analysis ID {analysis_id} doesn't exist")
                print(f"  2. Analysis doesn't belong to user {user_id}")
                print(f"{'='*60}\n")
                return None

            print(f"✓ Row found:")
            print(f"  ID: {row[0]}")
            print(f"  Title: {row[1]}")
            print(f"  Company: {row[2]}")
            print(f"  Match Score: {row[4]}")
            print(f"  CV Name: {row[10] if row[10] else 'Master Template'}")

            import json

            # Parse recommendations if it's a JSON string (TEXT column)
            recommendations_data = row[7] if row[7] else []
            if isinstance(recommendations_data, str):
                try:
                    recommendations_data = json.loads(recommendations_data)
                except (json.JSONDecodeError, TypeError):
                    recommendations_data = []

            analysis_data = {
                'id': row[0],
                'job_title': row[1],
                'job_company': row[2],
                'job_description': row[3],
                'match_score': row[4],
                'skills_match': row[5] if row[5] else [],  # Already deserialized from JSONB
                'skills_missing': row[6] if row[6] else [],  # Already deserialized from JSONB
                'recommendations': recommendations_data,  # Manually parsed from TEXT column
                'full_analysis': row[8],
                'created_at': row[9],
                'cv_name': row[10] if row[10] else 'Master Template'
            }
            print(f"✓ Returning analysis data")
            print(f"{'='*60}\n")
            return analysis_data
    except Exception as e:
        print(f"ERROR in get_analysis_by_id: {e}")
        import traceback
        traceback.print_exc()
        print(f"{'='*60}\n")
        return None


def delete_analysis_by_id(analysis_id, user_id):
    """Delete an analysis by ID"""
    engine = get_db_connection()
    if not engine:
        return False

    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                DELETE FROM job_analyses
                WHERE id = :analysis_id AND user_id = :user_id
                RETURNING job_title
            """), {"analysis_id": analysis_id, "user_id": user_id})

            deleted = result.fetchone()
            conn.commit()

            return deleted is not None
    except Exception as e:
        print(f"Error deleting analysis: {e}")
        return False


app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'your-secret-key-for-development')

# Rate limiting to protect against bot floods
# Use X-Forwarded-For so each real user IP is limited independently (Railway proxies requests)
def get_real_ip():
    forwarded_for = request.headers.get('X-Forwarded-For', '')
    if forwarded_for:
        return forwarded_for.split(',')[0].strip()
    return request.remote_addr

limiter = Limiter(
    get_real_ip,
    app=app,
    default_limits=["120 per minute"],
    storage_uri="memory://",
)

if os.environ.get('LOCAL_DEV') == 'true':
    @app.before_request
    def auto_login():
        session.setdefault('user_id', 1)
        session.setdefault('username', 'localuser')

def log_user_activity(action_type, details=None):
    """Log user activity to the database"""
    user_id = get_current_user_id()
    if not user_id:
        return
    
    engine = get_db_connection()
    if not engine:
        return
    
    try:
        # Get user's IP and browser info
        ip_address = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR', 'unknown'))
        if ',' in ip_address:
            ip_address = ip_address.split(',')[0].strip()
        
        user_agent = request.headers.get('User-Agent', 'unknown')
        
        with engine.connect() as conn:
            conn.execute(text("""
                INSERT INTO user_activity (user_id, action_type, action_details, ip_address, user_agent)
                VALUES (:user_id, :action_type, :action_details, :ip_address, :user_agent)
            """), {
                "user_id": user_id,
                "action_type": action_type,
                "action_details": details,
                "ip_address": ip_address,
                "user_agent": user_agent
            })
            conn.commit()
    except Exception as e:
        print(f"Error logging activity: {e}")


HISTORY_FILE = "search_history.json"

# Define global results variable
last_results = []
last_search_name = "Job_Search"


def load_config():
    with open("config.json", "r", encoding="utf-8") as f:
        return json.load(f)

config = load_config()

def load_saved_searches():
    engine = get_db_connection()
    if not engine:
        return []
    
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT name, timestamp, criteria, schedule, last_run_date 
                FROM saved_searches 
                WHERE user_id= :user_id
                ORDER BY id DESC 
                LIMIT 5
            """), {"user_id": get_current_user_id()})
            
            searches = []
            for row in result:
                search = {
                    "name": row[0],
                    "timestamp": row[1], 
                    "criteria": row[2],
                    "schedule": row[3] or "none",
                    "last_run_date": row[4] or ""
                }
                searches.append(search)
            return searches
    except Exception as e:
        print(f"Database error in load_saved_searches: {e}")
        return []

def check_excel_files_for_searches(searches):
    """Helper function to check which searches have Excel files"""
    engine = get_db_connection()
    
    for search in searches:
        schedule = search.get("schedule", "none")
        
        # For scheduled searches, check database
        if schedule != "none" and engine:
            try:
                with engine.connect() as conn:
                    result = conn.execute(text("""
                        SELECT COUNT(*) FROM scheduled_files 
                        WHERE search_name = :search_name AND user_id = :user_id
                    """), {
                        "search_name": search["name"],
                        "user_id": get_current_user_id()
                    })
                    count = result.fetchone()[0]
                    search["has_excel"] = count > 0
            except Exception as e:
                print(f"Database error checking files: {e}")
                search["has_excel"] = False
        else:
            # For non-scheduled searches, check filesystem (fallback)
            safe_name = search["name"].replace(" ", "_")
            pattern = os.path.join("scheduled_results", f"{safe_name}_*.xlsx")
            matching_files = glob.glob(pattern)
            search["has_excel"] = len(matching_files) > 0
    
    return searches

def detect_user_region(request):
    """Detect if user is in UK based on IP address"""
    print("🚨🚨🚨 DETECT_USER_REGION CALLED! 🚨🚨🚨")
    print(f"🧪 REGION DEBUG: Function called with request args: {request.args}")

    # TEST OVERRIDE: Check for manual region parameter
    test_region = request.args.get('test_region')
    print(f"🧪 REGION DEBUG: test_region parameter = '{test_region}'")

    if test_region and test_region.upper() in ['UK', 'US', 'SG', 'DE']:
        print(f"🧪 TEST MODE: Using manual region override: {test_region}")
        return test_region.upper()
 
    # Quick Frankfurt test for Germany
    try:
        if 'frankfurt' in request.form.get('location', '').lower():
            print("🔥 FORCING DE REGION FOR FRANKFURT TEST")
            return "DE"
    except:
        pass
     

    # STEP 1: Debug IP detection
    print("🔍 STEP 1: Checking IP detection...")
    http_x_forwarded = request.environ.get('HTTP_X_FORWARDED_FOR', 'NOT_FOUND')
    remote_addr = request.environ.get('REMOTE_ADDR', 'NOT_FOUND')
    print(f"🔍 HTTP_X_FORWARDED_FOR = '{http_x_forwarded}'")
    print(f"🔍 REMOTE_ADDR = '{remote_addr}'")
    
    user_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR', ''))
    if ',' in user_ip:
        user_ip = user_ip.split(',')[0].strip()
    
    print(f"🔍 Final selected IP = '{user_ip}'")

    # STEP 2: Test API call
    print("🔍 STEP 2: Testing API call...")
    try:
        import requests
        api_url = f"http://ip-api.com/json/{user_ip}"
        print(f"🔍 API URL = '{api_url}'")
        
        response = requests.get(api_url, timeout=10)
        print(f"🔍 API Response Status = {response.status_code}")
        print(f"🔍 API Response Headers = {dict(response.headers)}")
        print(f"🔍 API Response Text = '{response.text}'")
        
        if response.status_code == 200:
            data = response.json()
            print(f"🔍 API Response JSON = {data}")
            country_code = data.get('countryCode', 'NOT_FOUND')
            country_name = data.get('country', 'NOT_FOUND')
            print(f"🔍 Country Code = '{country_code}'")
            print(f"🔍 Country Name = '{country_name}'")
            
            # Enhanced mapping for CareerJet regions
            if country_code == "GB":
                return "UK"
            elif country_code == "CA":
                return "CA"
            elif country_code == "AU":
                return "AU"
            elif country_code == "DE":
                return "DE"
            elif country_code == "SG":
                return "SG"
            elif country_code == "IN":
                return "IN"
            else:
                return "US"  # Default 
        else:
            print(f"❌ API returned non-200 status: {response.status_code}")
    except Exception as e:
        print(f"❌ Exception during API call: {type(e).__name__}: {e}")
        import traceback
        print(f"❌ Full traceback: {traceback.format_exc()}")
    
    # Default to US if detection fails
    print("🔍 STEP 3: Defaulting to US")
    return "US"


def create_user(username, email, password):
    engine = get_db_connection()
    if not engine:
        return False
    
    try:
        password_hash = generate_password_hash(password)
        with engine.connect() as conn:
            conn.execute(text("""
                INSERT INTO users (username, email, password_hash, beta_user, beta_expires, subscription_status)
                VALUES (:username, :email, :password_hash, TRUE, '2025-12-31', 'none')
            """), {
                "username": username,
                "email": email,
                "password_hash": password_hash
            })
            conn.commit()
        return True
    except Exception as e:
        print(f"Error creating user: {e}")
        return False
        

def verify_user(username, password):
    engine = get_db_connection()
    if not engine:
        return None
    
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT id, password_hash FROM users WHERE username = :username
            """), {"username": username})
            
            user = result.fetchone()
            if user and check_password_hash(user[1], password):
                return user[0]  # Return user_id
        return None
    except Exception as e:
        print(f"Error verifying user: {e}")
        return None

def get_current_user_id():
    if os.environ.get('LOCAL_DEV') == 'true':
        return session.get('user_id', 1)
    return session.get('user_id')

def require_login():
    if os.environ.get('LOCAL_DEV') == 'true':
        return None
    if not get_current_user_id():
        return redirect('/login')
    return None

def save_search(name, criteria):
    print(f"🔍 SAVE_SEARCH DEBUG: Attempting to save '{name}'")
    print(f"🔍 SAVE_SEARCH DEBUG: Criteria: {criteria}")
    
    engine = get_db_connection()
    print(f"🔍 SAVE_SEARCH DEBUG: Engine: {engine}")
    
    if not engine:
        print("❌ SAVE_SEARCH DEBUG: No database connection")
        return
    
    print("🔍 SAVE_SEARCH DEBUG: About to execute SQL")
    try:
        with engine.connect() as conn:
            print("🔍 SAVE_SEARCH DEBUG: Connection established")
            conn.execute(text("""
                INSERT INTO saved_searches (name, timestamp, criteria, schedule, last_run_date, user_id)
                VALUES (:name, :timestamp, :criteria, :schedule, :last_run_date, :user_id)
            """), {
                "name": name,
                "timestamp": datetime.now().strftime("%d %B %Y"),
                "criteria": json.dumps(criteria),
                "schedule": "none",
                "last_run_date": "",
                "user_id" : get_current_user_id()
            })
            conn.commit()
            print("🔍 SAVE_SEARCH DEBUG: SQL executed successfully")
        print(f"✅ SAVE_SEARCH DEBUG: Successfully saved to database")
    except Exception as e:
        print(f"❌ SAVE_SEARCH DEBUG: Database error: {e}")


def save_results_for_search(name, results):
    # Create folder if not exists
    os.makedirs("scheduled_results", exist_ok=True)

    filename = f"scheduled_results/{name.replace(' ', '_')}.json"
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

from datetime import datetime

def save_results_to_excel(search_name, results):
    import pandas as pd
    from io import BytesIO
    import os
    from datetime import datetime

    os.makedirs("scheduled_results", exist_ok=True)

    safe_name = search_name.replace(" ", "_")
    date_str = datetime.now().strftime("%d_%B_%Y")
    filename = f"scheduled_results/{safe_name}_{date_str}.xlsx"

    df = pd.DataFrame(results)
    df["description"] = df["description"].apply(clean_description_for_excel)

    # Reorder columns so 'link' comes before 'description'
    cols = df.columns.tolist()
    if "link" in cols and "description" in cols:
        cols.remove("link")
        link_index = cols.index("description")
        cols.insert(link_index, "link")
    df = df[cols]

    df.insert(0, '#', range(1, len(df) + 1))

    with pd.ExcelWriter(filename, engine='xlsxwriter') as writer:
        temp_df = df.drop(columns=["link"])
        temp_df.to_excel(writer, index=False, sheet_name='Jobs', startcol=0)

        workbook = writer.book
        worksheet = writer.sheets['Jobs']

        header_format = workbook.add_format({
            'bold': True,
            'font_name': 'Calibri',
            'font_size': 12,
            'text_wrap': True,
            'align': 'center',
            'valign': 'vcenter'
        })

        default_format = workbook.add_format({
            'text_wrap': True,
            'align': 'left',
            'valign': 'vcenter'
        })

        description_format = workbook.add_format({
            'text_wrap': True,
            'align': 'left',
            'valign': 'top'
        })

        number_format = workbook.add_format({
            'align': 'center',
            'valign': 'vcenter'
        })

        # ADD THIS NEW FORMAT:
        source_format = workbook.add_format({
            'text_wrap': True,
            'align': 'center',        # Horizontal center
            'valign': 'vcenter',      # Vertical middle
            'font_name': 'Calibri'
        })

        source_col_index = None
        
        for col_num, value in enumerate(df.drop(columns=["link"]).columns.values):
            formatted_value = str(value).replace("_", " ").title()
            worksheet.write(0, col_num, formatted_value, header_format)

            if value == "#":
                worksheet.set_column(col_num, col_num, 5, number_format)
            elif value.lower() == "source":
                source_col_index = col_num
                worksheet.set_column(col_num, col_num, 10, source_format)  # CENTER ALIGN SOURCE
            elif value.lower() in ["title", "company", "location"]:
                worksheet.set_column(col_num, col_num, 23, default_format)
            elif value.lower() == "description":
                worksheet.set_column(col_num, col_num, 80, description_format)

        # EXPLICITLY format each Source column cell
        if source_col_index is not None:
            for row_num in range(len(df)):
                source_value = df.iloc[row_num]["source"] if "source" in df.columns else "EFC"
                worksheet.write(row_num + 1, source_col_index, source_value, source_format)
        
        link_col_index = df.columns.get_loc("link") + 1
        worksheet.write(0, link_col_index, "Link", header_format)
        worksheet.set_column(link_col_index, link_col_index, 55, default_format)

        link_format = workbook.add_format({
            'text_wrap': True,
            'align': 'left',
            'valign': 'vcenter',
            'font_color': 'blue',
            'underline': 1
        })

        for row_num in range(len(df)):
            url = df.iloc[row_num]["link"]
            if isinstance(url, str) and url.startswith("http"):
                worksheet.write_url(row_num + 1, link_col_index, url, link_format, url)
            else:
                worksheet.write(row_num + 1, link_col_index, url, link_format)

        worksheet.set_row(0, 25)
        for row_num in range(1, len(df) + 1):
            worksheet.set_row(row_num, 130)
        worksheet.autofilter(0, 0, len(df), len(df.columns))


    print(f"💾 Saved results to: {filename}")

def store_excel_in_database(search_name, file_path,user_id):
    """Store Excel file in database for scheduled searches"""
    engine = get_db_connection()
    if not engine:
        print("❌ No database connection for file storage")
        return
    
    try:
        # Read the Excel file as binary data
        with open(file_path, 'rb') as f:
            file_data = f.read()
        
        filename = os.path.basename(file_path)
       
        with engine.connect() as conn:
            # Delete old file if exists, then insert new one
            conn.execute(text("""
                DELETE FROM scheduled_files 
                WHERE search_name = :search_name AND user_id = :user_id
            """), {"search_name": search_name, "user_id": user_id})
            
            # Insert new file
            conn.execute(text("""
                INSERT INTO scheduled_files (search_name, user_id, file_data, filename)
                VALUES (:search_name, :user_id, :file_data, :filename)
            """), {
                "search_name": search_name,
                "user_id": user_id, 
                "file_data": file_data,
                "filename": filename
            })
            conn.commit()
        
        print(f"✅ Stored Excel file in database for: {search_name}")
        
    except Exception as e:
        print(f"❌ Error storing file in database: {e}")

def run_scheduled_searches():
    logger.info("🕓 Checking scheduled searches...")
    logger.info(f"🕓 DEBUG: Today's date string = '{datetime.now().strftime('%d %B %Y')}'")

    search_history = []
    engine = get_db_connection()
    if not engine:
        logger.info("❌ No database connection in scheduler")
        return
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT s.name, s.timestamp, s.criteria, s.schedule, s.last_run_date, s.user_id, u.email
                FROM saved_searches s
                JOIN users u ON s.user_id = u.id
                WHERE s.schedule != 'none' AND s.user_id IS NOT NULL
                ORDER BY s.id DESC
            """))
        
            for row in result:
                search = {
                    "name": row[0],
                    "timestamp": row[1], 
                    "criteria": row[2],
                    "schedule": row[3] or "none",
                    "last_run_date": row[4] or "",
                    "user_id": row[5],
                    "user_email": row[6]
                }
                search_history.append(search)
    except Exception as e:
        logger.info(f"❌ Database error in scheduler: {e}")
        return
    today_str = datetime.now().strftime("%d %B %Y")
    weekday = datetime.now().weekday()
    day = datetime.now().day

    updated = False

    for index, search in enumerate(search_history):
        schedule = search.get("schedule", "none")
        last_run_raw = search.get("last_run_date", "") #e.g., 21 June 2025 07:03
        last_run_date_only = " ".join(last_run_raw.split(" ")[:3]) # 21 June 2025

        logger.info(f"🕓 DEBUG: Search '{search['name']}' - schedule: {schedule}")
        logger.info(f"🕓 DEBUG: Last run raw: '{last_run_raw}'")
        logger.info(f"🕓 DEBUG: Last run date only: '{last_run_date_only}'")
        logger.info(f"🕓 DEBUG: Should skip? {last_run_date_only == today_str}")

        # skip if already ran today
        if last_run_date_only ==today_str:
            continue

        criteria = search.get("criteria", {})
        title = criteria.get("title", "")
        location = criteria.get("location", "")
        max_jobs = int(criteria.get("max_jobs", 10))

        should_run = (
            (schedule == "daily") or
            (schedule == "weekly" and weekday == 0) or
            (schedule == "monthly" and day == 1)
        )

      
        if should_run:
            logger.info(f"🔁 Running {schedule} search: {search['name']}")
            
            # Get the source from saved criteria and add seniority
            source = criteria.get("source", "efinancialcareers")
            seniority = criteria.get("seniority", "")
            logger.info(f"🔍 SCHEDULER: Using source '{source}' for search '{search['name']}'")
            
            # Call appropriate scraper based on saved source
            if source == "careerjet":
                from careerjet_api import scrape_jobs as scrape_careerjet_jobs
                results = scrape_careerjet_jobs(title, location, max_jobs, seniority=seniority, region="US")
                logger.info(f"🔍 SCHEDULER: Called CareerJet scraper, got {len(results)} results")
            else:
                results = scrape_jobs(title, location, max_jobs, seniority=seniority, region="US")
                logger.info(f"🔍 SCHEDULER: Called eFinancialCareers scraper, got {len(results)} results")
            
            save_results_to_excel(search["name"], results)
            
            # Build the same filename used in save_results_to_excel()
            safe_name = search["name"].replace(" ", "_")
            date_str = datetime.now().strftime("%d_%B_%Y")
            output_path = os.path.join("scheduled_results", f"{safe_name}_{date_str}.xlsx")
            store_excel_in_database(search["name"], output_path, search["user_id"])

            # 📧 Email the file if jobs exist
            subject = f"Scheduled Results for {search['name']} ({schedule})"
            body = f"Attached are the latest job search results for '{search['name']}' scheduled to run {schedule}."
            send_email_with_attachment(subject, body, output_path, config, search["user_email"])
            

            search["last_run_date"] = datetime.now().strftime("%d %B %Y %H:%M")
            updated = True
            logger.info(f"💾 Saved {len(results)} results to Excel for {search['name']}")

            # Update database immediately for this search
            with engine.connect() as conn:
                conn.execute(text("""
                    UPDATE saved_searches 
                    SET last_run_date = :last_run_date 
                    WHERE name = :name AND user_id = :user_id
                """), {
                    "last_run_date": search["last_run_date"],
                    "name": search["name"],
                    "user_id": search["user_id"]
                })
                conn.commit()

# Add scheduler initialization right after the function
print("🚀 SCHEDULER DEBUG: About to initialize scheduler...")

# Only run scheduler if ENABLE_SCHEDULER environment variable is set
if os.environ.get('ENABLE_SCHEDULER', 'false').lower() == 'true':
    print("🚀 SCHEDULER DEBUG: ENABLE_SCHEDULER is true, starting scheduler...")
    try:
        scheduler = BackgroundScheduler()
        print("🚀 SCHEDULER DEBUG: BackgroundScheduler created")
        def test_scheduler():
            print("🧪 TEST: Scheduler called a function!")
        scheduler.add_job(func=run_scheduled_searches, trigger="cron", hour=5, minute=0, max_instances=1, coalesce=True, misfire_grace_time=3600) # Runs daily at 5am
        scheduler.add_job(func=cleanup_old_records, trigger="cron", hour=3, minute=0, max_instances=1, coalesce=True, misfire_grace_time=3600) # Cleanup old DB records daily at 3am
        print("🚀 SCHEDULER DEBUG: Job added to scheduler")
        scheduler.start()
        print("🚀 SCHEDULER DEBUG: Scheduler started successfully!") 
        atexit.register(lambda: scheduler.shutdown())
        print("🚀 SCHEDULER DEBUG: Exit handler registered")
    except Exception as e:
        print(f"❌ SCHEDULER DEBUG: Failed to start scheduler: {e}")
        import traceback
        traceback.print_exc()
else:
    print("🚀 SCHEDULER DEBUG: ENABLE_SCHEDULER not set to true, skipping scheduler initialization")


def format_description(desc):
    import re
    lines = desc.splitlines()
    html = ""
    for line in lines:
        line = line.strip()
        if not line:
            continue
        elif line.lower() in ["overview", "responsibilities", "qualifications", "pay range"]:
            html += f"<h4>{line}</h4>"
        elif re.match(r"^[A-Z][a-z]+.*:$", line):
            html += f"<p><strong>{line}</strong></p>"
        elif re.match(r"^[-\u2022\*]\s+", line):
            html += f"<ul><li>{line[1:].strip()}</li></ul>"
        else:
            html += f"<p>{line}</p>"
    return html


def clean_description_for_excel(html):
    import re
    cleaned = (
        html.replace("<br>", "\n")
            .replace("<br/>", "\n")
            .replace("<br />", "\n")
            .replace("</li>", "\n")
            .replace("<li>", "• ")
            .replace("&amp;", "&")
            .replace("<strong>", "")
            .replace("</strong>", "")
            .replace("<ul>", "")
            .replace("</ul>", "")
            .replace("<b>", "")
            .replace("</b>", "")
            .replace("<u>", "")
            .replace("</u>", "")
            .replace("&nbsp;", " ")
    )
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()

def send_email_with_attachment(subject, body, attachment_path, config, user_email=None):
    try:
        smtp_server = config["email_settings"]["smtp_server"]
        smtp_port = config["email_settings"]["smtp_port"]
        sender_email = config["email_settings"]["sender_email"]
        sender_password = config["email_settings"]["sender_password"]
        recipients = [user_email] if user_email else []
        logger.info(f"🔍 EMAIL DEBUG: Attempting SMTP connection to {smtp_server}:{smtp_port}")  # ADD THIS LINE HERE

        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = "Find Me A Job <fmaj.app@gmail.com>"
        msg["To"] = ", ".join(recipients)
        msg.set_content(body)

        # Attach the Excel file
        with open(attachment_path, "rb") as f:
            file_data = f.read()
            file_name = os.path.basename(attachment_path)
        msg.add_attachment(file_data, maintype="application", subtype="octet-stream", filename=file_name)

        with smtplib.SMTP(smtp_server, smtp_port) as smtp:
            smtp.starttls()
            smtp.login(sender_email, sender_password)
            smtp.send_message(msg)

        logger.info(f"✅ Email sent to {recipients} with: {file_name}")  # ADD THIS LINE HERE
        return True

    except Exception as e:
        logger.error(f"❌ Failed to send email: {e}")
        return False

# ADD THESE FUNCTIONS AFTER send_email_with_attachment FUNCTION
def check_feature_access(feature_name):
    """Check if current user has access to a specific feature"""
    user_id = get_current_user_id()
    if not user_id:
        return False
    
    engine = get_db_connection()
    if not engine:
        return False
    
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT beta_user, beta_expires, subscription_status 
                FROM users WHERE id = :user_id
            """), {"user_id": user_id})
            user = result.fetchone()
        
        if not user:
            return False
        
        beta_user, beta_expires, subscription_status = user
        today = datetime.now().date()
        
        # Beta users get full access until beta expires
        if beta_user and today <= beta_expires:
            return True
        
        # Paid users get full access
        if subscription_status == 'active':
            return True
        
        # Free users only get basic features
        if feature_name in ['basic_search']:
            return True
        
        return False  # Block premium features
    
    except Exception as e:
        print(f"Error checking feature access: {e}")
        return False

def check_daily_search_limit():
    """Check if user has exceeded daily search limit (for free users)"""
    if check_feature_access('unlimited_searches'):
        return True  # Beta/paid users have unlimited searches
    
    # For free users, check their daily limit
    user_id = get_current_user_id()
    if not user_id:
        return False
    
    engine = get_db_connection()
    if not engine:
        return True  # If DB fails, allow search
    
    try:
        today = datetime.now().date()
        
        with engine.connect() as conn:
            # Get today's search count
            result = conn.execute(text("""
                SELECT search_count FROM daily_search_limits 
                WHERE user_id = :user_id AND search_date = :today
            """), {"user_id": user_id, "today": today})
            
            row = result.fetchone()
            current_count = row[0] if row else 0
            
            return current_count < 3  # Allow if under 3 searches
            
    except Exception as e:
        print(f"Error checking search limit: {e}")
        return True  # If error, allow search


def increment_search_count():
    """Increment the user's daily search count"""
    if check_feature_access('unlimited_searches'):
        return  # Beta/paid users don't need tracking
    
    user_id = get_current_user_id()
    if not user_id:
        return
    
    engine = get_db_connection()
    if not engine:
        return
    
    try:
        today = datetime.now().date()
        
        with engine.connect() as conn:
            # Insert or update today's count
            conn.execute(text("""
                INSERT INTO daily_search_limits (user_id, search_date, search_count)
                VALUES (:user_id, :today, 1)
                ON CONFLICT (user_id, search_date) 
                DO UPDATE SET search_count = daily_search_limits.search_count + 1
            """), {"user_id": user_id, "today": today})
            conn.commit()
            
    except Exception as e:
        print(f"Error incrementing search count: {e}")

def render_template_with_admin(template_name, **kwargs):
    """
    Helper function to automatically include admin status in all template renders
    """
    # Check if user is admin
    is_admin = session.get('username') == 'frameit'
    
    # Add admin status to template context
    kwargs['is_admin'] = is_admin
    
    # Add username for debugging if needed
    kwargs['current_username'] = session.get('username', '')
    
    return render_template(template_name, **kwargs)
    
@app.route("/")
def root():
    user_id = get_current_user_id()
    if not user_id:
        return redirect("/login")

    # Send new/incomplete users to Get Started until setup is done
    has_prompt = get_user_prompt_preference(user_id) is not None
    has_template = get_user_master_template(user_id) is not None
    if not has_prompt or not has_template:
        return redirect("/get-started")

    return redirect("/ai-match")

@app.route("/app", methods=["GET", "POST"])
def index():
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
        
    global last_results
    jobs = []
    if request.method == "POST":
        print("🚨 POST METHOD DETECTED - Starting debug")
        info = None # Add this line
        print(f"🚨 DEBUG: request.args = {request.args}")
        print(f"🚨 DEBUG: test_region = {request.args.get('test_region')}")
        print(f"🚨 POST METHOD DETECTED - Starting search processing")
        # ADD THIS SEARCH LIMIT CHECK HERE
        if not check_daily_search_limit():
            # Get current search count for display
            user_id = get_current_user_id()
            engine = get_db_connection()
            current_count = 3  # Default to max
        
            if engine:
                try:
                    with engine.connect() as conn:
                        today = datetime.now().date()
                        result = conn.execute(text("""
                            SELECT search_count FROM daily_search_limits 
                            WHERE user_id = :user_id AND search_date = :today
                        """), {"user_id": user_id, "today": today})
                        row = result.fetchone()
                        current_count = row[0] if row else 0
                except:
                    pass
        
            return render_template("upgrade.html", 
                                 feature="unlimited searches",
                                 current_plan="free",
                                 search_limit_reached=True,
                                 searches_used=current_count)
        title = request.form.get("title", "")
        location = request.form.get("location", "")
        seniority = request.form.get("seniority", "")
        raw_max_jobs = request.form.get("max_jobs", "")
        try:
            max_jobs = int(raw_max_jobs)
            if max_jobs <= 0 or max_jobs > 50:
                max_jobs = 50
        except (ValueError, TypeError):
            max_jobs = 50
            
        # Add source selection handling HERE
        source = request.form.get("source", "efinancialcareers")
        print(f"🔍 DEBUG: Raw form data = {dict(request.form)}")
        print(f"🔍 DEBUG: Source from form = '{request.form.get('source')}'")
        print(f"🔍 DEBUG: Final source variable = '{source}'")
        print(f"🔍 DEBUG: Selected source = '{source}'")
        print(f"🔍 DEBUG: Username = '{session.get('username')}'")

       
        # Remove Admin-only access to Indeed/Both - others get "coming soon" message
        #if source in ["CareerJet", "both"] and session.get('username') != 'frameit':
        #    info = "⏳ CareerJet search is coming soon! We're currently testing this feature."
        #    print(f"🔍 DEBUG: Non-admin user blocked from {source}")

        #    # Return early with the message, don't run any scraper
        #    saved_searches = check_excel_files_for_searches(load_saved_searches())
        #    print(f"🔍 TEMPLATE DEBUG: Passing source = '{source}' to template")
        #    return render_template("index.html", info=info, jobs=jobs, title=title, location=location, source=source, max_jobs=max_jobs, seniority=seniority, has_scheduling_access=check_feature_access('scheduling'), saved_searches=saved_searches)

        #print(f"🔍 DEBUG: Admin user '{session.get('username')}' can access {source}")
       
        if request.form.get("action") == "save":
            search_name = request.form.get("search_name", "").strip()
            if not search_name:
                info = "⚠️ Please enter a name for your saved search."
            else:
                form_location = request.form.get("location", "").strip()

                # ADD THESE DEBUG LINES:
                print("🔍 SAVE DEBUG: All form data received:")
                for key, value in request.form.items():
                    print(f"    {key} = '{value}'")
        
                source_from_form = request.form.get("source", "DEFAULT_NOT_FOUND")
                print(f"🔍 SAVE DEBUG: Source from form = '{source_from_form}'")

                # Convert source to display name - THIS MUST BE HERE
                source_display = "EFC" if source_from_form == "efinancialcareers" else "CareerJet"
                
                criteria = {
                    "title": request.form.get("title", ""),
                    "location": form_location,
                    "max_jobs": max_jobs,
                    "seniority": request.form.get("seniority", ""),
                    "source": source_from_form if source_from_form != "DEFAULT_NOT_FOUND" else "efinancialcareers"
                }
                print(f"🔍 SAVE DEBUG: Final criteria = {criteria}")
                # Add source abbreviation to saved search name
                source_abbrev = "EFC" if source_from_form == "efinancialcareers" else "CareerJet"
        
                # Build the formatted name: "User Input - Location - Source"
        
                if form_location:
                    formatted_name = f"{search_name} - {form_location} - {source_display}"
                else:
                    formatted_name = f"{search_name} - {source_display}"
                save_search(formatted_name, criteria)                
                info = f"✅ Search saved as: {formatted_name}"
            
            saved_searches = check_excel_files_for_searches(load_saved_searches())
            return render_template("index.html", info=info, jobs=last_results, title=title, location=location, max_jobs=max_jobs, seniority=seniority, has_scheduling_access=check_feature_access('scheduling'), saved_searches=saved_searches)
            
       
        print(f"🔍 FLASK DEBUG: About to call scraper with seniority='{seniority}', type={type(seniority)}")
        
        try:
            # Choose scraper based on source
   
            if source == "careerjet":
                print(f"🚨 APP.PY DEBUG: About to call CareerJet API - VERSION 2.0")  # ADD THIS
                print(f"🔍 DEBUG: Running CareerJet API for admin user")
                try:
                    # Debug region detection  
                    region = detect_user_region(request)
                    print(f"🔍 DEBUG: Detected region = '{region}'")

                    from careerjet_api import scrape_jobs as scrape_careerjet_jobs
                    jobs = scrape_careerjet_jobs(title, location, max_jobs, seniority=seniority, region=region)

                except Exception as e: 
                    print(f"❌ JobSpy Indeed failed: {e}")
                    jobs = [{
                        "title": "CareerJet API Error",
                        "company": "System Info",
                        "location": location,
                        "link": "#",
                        "description": f"CareerJet API error: {str(e)}. Please try eFinancialCareers for now.",
                        "formatted_description": f"CareerJet API error: {str(e)}. Please try eFinancialCareers for now."                        
                    }]
                    
           
            else:
                # Regular efinancialcareers scraper
                region = detect_user_region(request)
                jobs = scrape_jobs(title, location, max_jobs, seniority=seniority, region=region)
            
            # Check if this is a special "no results" message
            if jobs and len(jobs) == 1 and jobs[0].get("no_results"):
                special_message = jobs[0].get("special_message", "No jobs found")
                return render_template("index.html", 
                                     jobs=[], 
                                     title=title, 
                                     location=location, 
                                     max_jobs=max_jobs, 
                                     seniority=seniority, 
                                     saved_searches=load_saved_searches(),
                                     has_scheduling_access=check_feature_access('scheduling'),
                                     special_message=special_message)           
            
            # Check if this is an error response from scraper
            if jobs and len(jobs) == 1 and jobs[0].get("error_type"):
                error_job = jobs[0]
                error_message = error_job.get("description", "An error occurred during the search.")
                return render_template("index.html", 
                                     jobs=[], 
                                     title=title, 
                                     location=location, 
                                     max_jobs=max_jobs, 
                                     seniority=seniority, 
                                     saved_searches=load_saved_searches(),
                                     has_scheduling_access=check_feature_access('scheduling'),
                                     special_message=error_message)
            
            # Process job descriptions for successful results
            for job in jobs:
                if not job.get("company"):
                    job["company"] = "[Not Found]"

                # Clean and decode HTML entities in description
                import html
                raw_description = job.get("description", "Description not available")
                decoded_description = html.unescape(raw_description)
                cleaned_description = (
                    decoded_description.replace("<u>", "")
                                      .replace("</u>", "")
                                      .replace("<strong>", "")
                                      .replace("</strong>", "")
                                      .replace("<b>", "")
                                      .replace("</b>", "")
                )
                job["formatted_description"] = cleaned_description
                # Add source field
                job["source"] = "EFC" if source == "efinancialcareers" else "CareerJet"
                
        except Exception as e:
            print(f"❌ LOAD SEARCH ERROR: {str(e)}")
            print(f"❌ ERROR TYPE: {type(e).__name__}")
            import traceback
            traceback.print_exc()
            
            # Fallback error message for unexpected Flask errors
            error_message = "We're experiencing technical difficulties. Please try again in a few minutes. If you continue seeing this error, email fmaj.app@gmail.com with details about what you were searching for."
            return render_template("index.html", 
                                 jobs=[], 
                                 title=title, 
                                 location=location, 
                                 max_jobs=max_jobs, 
                                 seniority=seniority, 
                                 saved_searches=load_saved_searches(),
                                 has_scheduling_access=check_feature_access('scheduling'),
                                 special_message=error_message)

                    
        # Check if this is a special "no results" message
        if jobs and len(jobs) == 1 and jobs[0].get("no_results"):
            special_message = jobs[0].get("special_message", "No jobs found")
            return render_template("index.html", 
                                 jobs=[], 
                                 title=title, 
                                 location=location, 
                                 max_jobs=max_jobs, 
                                 seniority=seniority, 
                                 saved_searches=load_saved_searches(),
                                 has_scheduling_access=check_feature_access('scheduling'),
                                 special_message=special_message)    

        # IMPORTANT: After successful search, increment the count
        if jobs and len(jobs) > 0 and not any(job.get("error_type") for job in jobs):
            increment_search_count()
       
        last_results = jobs
        global last_search_name
        last_search_name = title if title else "Job_Search"
        print(f"🔍 DEBUG: About to render template with info = '{info}'")
        return render_template("index.html", info=info, jobs=jobs, title=title, location=location, source=source, max_jobs=max_jobs, seniority=seniority, has_scheduling_access=check_feature_access('scheduling'), saved_searches=load_saved_searches())
    
    saved_searches = check_excel_files_for_searches(load_saved_searches())
    print("✅ Saved searches and their Excel status:")
    for s in saved_searches:
        print(f"- {s['name']}: has_excel = {s.get('has_excel')}")

    return render_template_with_admin("index.html", title="", location="", seniority="", max_jobs=10, has_scheduling_access=check_feature_access('scheduling'), saved_searches=saved_searches)

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        
                
        # Your existing login code continues below...
        if not username or not password:
            return render_template("login.html", error="Please enter both username and password")
        
        user_id = verify_user(username, password)
        if user_id:
            session['user_id'] = user_id
            session['username'] = username
            log_user_activity("login", f"successful login")  # ADD THIS LINE
            return redirect("/")
        else:
            return render_template("login.html", error="Invalid username or password")
    
    return render_template("login.html")
    

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        
        if not username or not email or not password:
            return render_template("signup.html", error="Please fill in all fields")
        
        if create_user(username, email, password):
            return render_template("login.html", success="Account created! Please log in.")
        else:
            return render_template("signup.html", error="Username or email already exists")
    
    return render_template("signup.html")

@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        
        if not email:
            return render_template("forgot_password.html", error="Please enter your email address")
        
        # Look up user by email
        engine = get_db_connection()
        if not engine:
            return render_template("forgot_password.html", error="Database connection error")
        
        try:
            with engine.connect() as conn:
                result = conn.execute(text("""
                    SELECT id, username FROM users WHERE email = :email
                """), {"email": email})
                user = result.fetchone()
                
                if user:
                    user_id = user[0]
                    username = user[1]
                    
                    # Create reset token
                    token = create_password_reset_token(user_id)
                    
                    if token:
                        # Send reset email
                        reset_url = f"{request.host_url}reset-password/{token}"
                        subject = "Password Reset - Find Me A Job"
                        body = f"""Hello {username},

You requested a password reset for your Find Me A Job account.

Click the link below to reset your password:
{reset_url}

This link will expire in 1 hour.

If you didn't request this password reset, please ignore this email.

Best regards,
Find Me A Job Team"""
                        
                        # Use existing email configuration
                        try:
                            smtp_server = config["email_settings"]["smtp_server"]
                            smtp_port = config["email_settings"]["smtp_port"]
                            sender_email = config["email_settings"]["sender_email"]
                            sender_password = config["email_settings"]["sender_password"]
                            
                            import smtplib
                            from email.message import EmailMessage
                            
                            msg = EmailMessage()
                            msg["Subject"] = subject
                            msg["From"] = sender_email
                            msg["To"] = email
                            msg.set_content(body)
                            
                            with smtplib.SMTP(smtp_server, smtp_port) as smtp:
                                smtp.starttls()
                                smtp.login(sender_email, sender_password)
                                smtp.send_message(msg)
                                
                            return render_template("forgot_password.html", 
                                success="Password reset instructions have been sent to your email")
                                
                        except Exception as e:
                            print(f"Error sending email: {e}")
                            return render_template("forgot_password.html", 
                                error="Error sending email. Please try again later.")
                    else:
                        return render_template("forgot_password.html", 
                            error="Error creating reset token. Please try again later.")
                else:
                    # Don't reveal whether email exists or not (security)
                    return render_template("forgot_password.html", 
                        success="If an account with that email exists, password reset instructions have been sent")
                        
        except Exception as e:
            print(f"Database error: {e}")
            return render_template("forgot_password.html", error="Database error. Please try again later.")
    
    return render_template("forgot_password.html")
@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password(token):
    engine = get_db_connection()
    if not engine:
        return render_template("reset_password.html", error="Database connection error", token=token)
    
    # Validate token
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT r.id, r.user_id, r.expires_at, r.used, u.username 
                FROM password_reset_tokens r
                JOIN users u ON r.user_id = u.id
                WHERE r.token = :token
            """), {"token": token})
            token_data = result.fetchone()
            
            if not token_data:
                return render_template("reset_password.html", 
                    error="Invalid reset link. Please request a new password reset.", token=token)
            
            token_id, user_id, expires_at, used, username = token_data
            
            # Check if token is expired or used
            from datetime import datetime
            if used:
                return render_template("reset_password.html", 
                    error="This reset link has already been used. Please request a new password reset.", token=token)
            
            if datetime.now() > expires_at:
                return render_template("reset_password.html", 
                    error="This reset link has expired. Please request a new password reset.", token=token)
            
            # Token is valid, process password reset
            if request.method == "POST":
                new_password = request.form.get("new_password", "")
                confirm_password = request.form.get("confirm_password", "")
                
                if not new_password or not confirm_password:
                    return render_template("reset_password.html", 
                        error="Please fill in both password fields", token=token, username=username)
                
                if new_password != confirm_password:
                    return render_template("reset_password.html", 
                        error="Passwords do not match", token=token, username=username)
                
                # Update password and mark token as used
                try:
                    password_hash = generate_password_hash(new_password)
                    
                    # Update password
                    conn.execute(text("""
                        UPDATE users SET password_hash = :password_hash WHERE id = :user_id
                    """), {"password_hash": password_hash, "user_id": user_id})
                    
                    # Mark token as used
                    conn.execute(text("""
                        UPDATE password_reset_tokens SET used = TRUE WHERE id = :token_id
                    """), {"token_id": token_id})
                    
                    conn.commit()
                    
                    return render_template("reset_password.html", 
                        success="Password successfully reset! You can now log in with your new password.",
                        token=token, username=username)
                        
                except Exception as e:
                    print(f"Error updating password: {e}")
                    return render_template("reset_password.html", 
                        error="Error updating password. Please try again.", token=token, username=username)
            
            # GET request - show reset form
            return render_template("reset_password.html", token=token, username=username)
            
    except Exception as e:
        print(f"Database error: {e}")
        return render_template("reset_password.html", error="Database error. Please try again later.", token=token)


@app.route("/settings", methods=["GET", "POST"])
def settings():
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
    
    engine = get_db_connection()
    if not engine:
        return "Database connection error", 500
    
    # Get current user info
    user_id = get_current_user_id()
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT username, email FROM users WHERE id = :user_id
            """), {"user_id": user_id})
            user = result.fetchone()
            
            if not user:
                return redirect("/logout")
                
            current_username = user[0]
            current_email = user[1]
    except Exception as e:
        print(f"Error fetching user data: {e}")
        return "Error loading settings", 500
    
    if request.method == "POST":
        action = request.form.get("action")
        
        if action == "update_profile":
            # Handle profile updates
            new_username = request.form.get("username", "").strip()
            new_email = request.form.get("email", "").strip()
            
            if not new_username or not new_email:
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    error="Please fill in all fields")
            
            try:
                with engine.connect() as conn:
                    conn.execute(text("""
                        UPDATE users SET username = :username, email = :email 
                        WHERE id = :user_id
                    """), {
                        "username": new_username,
                        "email": new_email,
                        "user_id": user_id
                    })
                    conn.commit()
                    
                    # Update session
                    session['username'] = new_username
                    
                return render_template("settings.html", 
                    username=new_username, email=new_email,
                    success="Profile updated successfully!")
            except Exception as e:
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    error="Username or email already exists")
        
        elif action == "change_password":
            # Handle password change
            current_password = request.form.get("current_password", "")
            new_password = request.form.get("new_password", "")
            confirm_password = request.form.get("confirm_password", "")
            
            if not current_password or not new_password or not confirm_password:
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    error="Please fill in all password fields")
            
            if new_password != confirm_password:
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    error="New passwords don't match")
            
            # Verify current password
            if not verify_user(current_username, current_password):
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    error="Current password is incorrect")
            
            try:
                password_hash = generate_password_hash(new_password)
                with engine.connect() as conn:
                    conn.execute(text("""
                        UPDATE users SET password_hash = :password_hash 
                        WHERE id = :user_id
                    """), {
                        "password_hash": password_hash,
                        "user_id": user_id
                    })
                    conn.commit()
                    
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    success="Password changed successfully!")
            except Exception as e:
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    error="Error changing password")
    
   
        elif action == "delete_account":
            # Handle account deletion
            try:
                with engine.connect() as conn:
                    # Start a transaction to ensure all deletions succeed together
                    trans = conn.begin()
                    
                    try:
                        # Delete user's saved searches
                        conn.execute(text("""
                            DELETE FROM saved_searches WHERE user_id = :user_id
                        """), {"user_id": user_id})
                        
                        # Delete user's scheduled files
                        conn.execute(text("""
                            DELETE FROM scheduled_files WHERE user_id = :user_id
                        """), {"user_id": user_id})
                        
                        # Delete user's password reset tokens
                        conn.execute(text("""
                            DELETE FROM password_reset_tokens WHERE user_id = :user_id
                        """), {"user_id": user_id})
                        
                        # Finally, delete the user account
                        conn.execute(text("""
                            DELETE FROM users WHERE id = :user_id
                        """), {"user_id": user_id})
                        
                        # Commit the transaction
                        trans.commit()
                        
                        # Clear the session
                        session.clear()
                        
                        # Redirect to a confirmation page
                        return render_template("account_deleted.html")
                        
                    except Exception as e:
                        # If any deletion fails, roll back the transaction
                        trans.rollback()
                        print(f"Error during account deletion: {e}")
                        return render_template("settings.html", 
                            username=current_username, email=current_email,
                            error="Error deleting account. Please try again later.")
                            
            except Exception as e:
                print(f"Database error during account deletion: {e}")
                return render_template("settings.html", 
                    username=current_username, email=current_email,
                    error="Database error. Please try again later.")   
    
    
    return render_template("settings.html", username=current_username, email=current_email)

@app.route("/submit_feedback", methods=["POST"])
def submit_feedback():
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
    
    feedback_type = request.form.get("feedback_type", "")
    message = request.form.get("feedback_message", "").strip()
    email = request.form.get("feedback_email", "").strip()
    
    if not message:
        return redirect("/")  # Could add error handling here
    
    # Get current user info
    user_id = get_current_user_id()
    username = session.get('username', 'Unknown')
    
    # Send feedback via email (using existing email config)
    try:
        subject = f"Feedback: {feedback_type.title()} from {username}"
        body = f"""New feedback received:

Type: {feedback_type.title()}
From User: {username} (ID: {user_id})
Contact Email: {email if email else 'Not provided'}

Message:
{message}

Timestamp: {datetime.now().strftime('%d %B %Y %H:%M:%S')}
"""
        
        # Use your existing email configuration
        smtp_server = config["email_settings"]["smtp_server"]
        smtp_port = config["email_settings"]["smtp_port"]
        sender_email = config["email_settings"]["sender_email"]
        sender_password = config["email_settings"]["sender_password"]
        admin_email = config["email_settings"]["recipients"][0]  # Send to your admin email
        
        import smtplib
        from email.message import EmailMessage
        
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = f"{username} via Find Me A Job <{sender_email}>"
        msg["To"] = admin_email
        if email:
            msg["Reply-To"] = email
        msg.set_content(body)
        
        with smtplib.SMTP(smtp_server, smtp_port) as smtp:
            smtp.starttls()
            smtp.login(sender_email, sender_password)
            smtp.send_message(msg)
            
        print(f"✅ Feedback sent from {username}: {feedback_type}")
        
    except Exception as e:
        print(f"❌ Error sending feedback: {e}")
    
    return redirect("/?feedback=sent")


@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")

@app.route("/load_search/<int:index>")
def load_saved_search(index):
    searches = check_excel_files_for_searches(load_saved_searches())

    if 0 <= index < len(searches):
        criteria = searches[index]["criteria"]
        title = criteria.get("title", "")
        location = criteria.get("location", "")
        seniority = criteria.get("seniority", "")
        max_jobs = criteria.get("max_jobs", 10)  # Get from criteria instead of form
        
        try:
            max_jobs = int(max_jobs)
            if max_jobs <= 0 or max_jobs > 50:
                max_jobs = 50
        except (ValueError, TypeError):
            max_jobs = 50

        try:
            region = detect_user_region(request)
            saved_source = criteria.get("source", "efinancialcareers")
            print(f"🔍 LOAD DEBUG: Using SAVED source = '{saved_source}' (ignoring current form)")
            print(f"🔍 LOAD DEBUG: saved_source == 'indeed'? {saved_source == 'indeed'}")
            print(f"🔍 LOAD DEBUG: All criteria: {criteria}")

            # Use the SAVED source, not the form
            if saved_source == "careerjet":
                print("🔍 LOAD DEBUG: Calling CareerJet because saved search was CareerJet")
                from careerjet_api import scrape_jobs as scrape_careerjet_jobs
                jobs = scrape_careerjet_jobs(title, location, max_jobs, seniority=seniority, region=region)
            else:
                print("🔍 LOAD DEBUG: Calling eFinancialCareers because saved search was eFinancialCareers")
                jobs = scrape_jobs(title, location, max_jobs, seniority=seniority, region=region)
                        
            log_user_activity("search", f"'{title}' in '{location}' ({len(jobs)} results)")
            
        except Exception as e:
            print(f"❌ LOAD SEARCH ERROR: {str(e)}")
            print(f"❌ ERROR TYPE: {type(e).__name__}")
            import traceback
            traceback.print_exc()
            jobs = [{"title": "Scraping Failed", "company": "Error", "location": location, "link": "#", "description": f"Error: {str(e)}"}]
        
        # Format jobs for display (same as in main index route)
        for job in jobs:
            # Ensure the company name is included for rendering
            if not job.get("company"):
                job["company"] = "[Not Found]"

            job["formatted_description"] = job.get("description", "Description not available")

     
        
        global last_results
        last_results = jobs

        global last_search_name
        last_search_name = searches[index]['name']

        return render_template(
            "index.html",
            jobs=jobs,
            title=title,
            location=location,
            max_jobs=max_jobs,
            source=saved_source,
            saved_searches=searches,
            criteria=searches[index]["criteria"],
            name=searches[index]["name"],
            timestamp=datetime.now().strftime("%d %B %Y"),
            active_search_name=f"{searches[index]['name']} – {location} – {datetime.now().strftime('%d %B %Y')}",
            has_scheduling_access=check_feature_access('scheduling')
        )
    else:
        return redirect("/")
    

@app.route('/download', methods=['POST'])
def download():
    # ADD THIS CHECK AT THE BEGINNING
    if not check_feature_access('excel_export'):
        return render_template("upgrade.html", 
                             feature="Excel exports",
                             current_plan="free")

    # Your existing code continues unchanged....   
    global last_results
    if not last_results:
        return "No results to export", 400

    df = pd.DataFrame(last_results).drop(columns=["formatted_description"], errors="ignore")
    df["description"] = df["description"].apply(clean_description_for_excel)

    # Reorder columns so 'link' comes before 'description'
    cols = df.columns.tolist()
    if "link" in cols and "description" in cols:
        cols.remove("link")
        link_index = cols.index("description")
        cols.insert(link_index, "link")
    df = df[cols]

    # Add the '#' column at the beginning
    df.insert(0, '#', range(1, len(df) + 1))

    # Move Source column to position 1 (after #)
    if "source" in df.columns:
        source_col = df.pop("source")
        df.insert(1, "source", source_col)

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        # Write all columns except 'link'
        temp_df = df.drop(columns=["link"])
        temp_df.to_excel(writer, index=False, sheet_name='Jobs', startcol=0)

        workbook = writer.book
        worksheet = writer.sheets['Jobs']

        # Header format
        header_format = workbook.add_format({
            'bold': True,
            'font_name': 'Calibri',
            'font_size': 12,
            'text_wrap': True,
            'align': 'center',
            'valign': 'vcenter'
        })

        # Column format templates
        default_format = workbook.add_format({
            'text_wrap': True,
            'align': 'left',
            'valign': 'vcenter'
        })

        description_format = workbook.add_format({
            'text_wrap': True,
            'align': 'left',
            'valign': 'top'
        })

        number_format = workbook.add_format({
            'align': 'center',
            'valign': 'vcenter'
        })

        # ADD THIS FORMAT DEFINITION:
        source_format = workbook.add_format({
            'text_wrap': True,
            'align': 'center',
            'valign': 'vcenter',
            'font_name': 'Calibri'
        })
        
        # Write headers and set column widths/formats
        for col_num, value in enumerate(df.drop(columns=["link"]).columns.values):
            formatted_value = str(value).replace("_", " ").title()
            worksheet.write(0, col_num, formatted_value, header_format)

            if value == "#":
                worksheet.set_column(col_num, col_num, 5, number_format)
            elif value.lower() == "source":  # ADD THIS LINE
                worksheet.set_column(col_num, col_num, 10, source_format)  # ADD THIS LINE
            elif value.lower() in ["title", "company", "location"]:
                worksheet.set_column(col_num, col_num, 23, default_format)
            elif value.lower() == "description":
                worksheet.set_column(col_num, col_num, 80, description_format)

        # Write "Link" header and column manually
        link_col_index = df.columns.get_loc("link") + 1
        worksheet.write(0, link_col_index, "Link", header_format)
        worksheet.set_column(link_col_index, link_col_index, 55, default_format)

        # Format and write links row by row
        link_format = workbook.add_format({
            'text_wrap': True,
            'align': 'left',
            'valign': 'vcenter',
            'font_color': 'blue',
            'underline': 1
        })

        for row_num in range(len(df)):
            url = df.iloc[row_num]["link"]
            if isinstance(url, str) and url.startswith("http"):
                worksheet.write_url(row_num + 1, link_col_index, url, link_format, url)
            else:
                worksheet.write(row_num + 1, link_col_index, url, link_format)

        # Set row height: header and data rows
        worksheet.set_row(0, 25)
        for row_num in range(1, len(df) + 1):
            worksheet.set_row(row_num, 130)
        worksheet.autofilter(0, 0, len(df), len(df.columns))


    global last_search_name
    title_from_form = request.form.get("title", "").strip()

    # Get source from the current search
    if hasattr(request, 'form') and request.form.get("source"):
        source_from_form = request.form.get("source", "efinancialcareers")
    else:
        # Try to detect source from job results
        source_from_form = "efinancialcareers"  # default
        if last_results and len(last_results) > 0:
            first_job_source = last_results[0].get("source", "EFC")
            source_from_form = "careerjet" if first_job_source == "CareerJet" else "efinancialcareers"

    source_abbrev = "EFC" if source_from_form == "efinancialcareers" else "CareerJet"
    
    if last_search_name and last_search_name != "Job_Search":
        base_name = last_search_name
        # Remove existing source suffix if present
        base_name = base_name.replace(f"_{source_abbrev}", "").replace("_EFC", "").replace("_CareerJet", "")
    elif title_from_form:
        base_name = title_from_form
    else:
        base_name = "Job_Search"

    # Get date string
    date_str = datetime.now().strftime('%d_%B_%Y')

    # Create filename: Position_Date_Source.xlsx
    filename = f"{base_name.replace(' – ', '_').replace(' ', '_')}_{date_str}_{source_abbrev}.xlsx"
        
 
    output.seek(0)
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        download_name=filename,
        as_attachment=True
    )

@app.route("/delete_saved_search", methods=["POST"])
def delete_saved_search():
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
        
    index = int(request.form.get("index", -1))
    if index >= 0:
        # Get current saved searches from database
        searches = load_saved_searches()
        
        if 0 <= index < len(searches):
            search_to_delete = searches[index]
            
            engine = get_db_connection()
            if engine:
                try:
                    with engine.connect() as conn:
                        conn.execute(text("""
                            DELETE FROM saved_searches 
                            WHERE name = :name AND user_id = :user_id
                        """), {
                            "name": search_to_delete["name"],
                            "user_id": get_current_user_id()
                        })
                        conn.commit()
                        
                    print(f"🗑️ Deleted saved search: {search_to_delete['name']}")
                    
                except Exception as e:
                    print(f"❌ Error deleting search: {e}")
                    
    return redirect("/")

@app.route("/rename/<int:index>", methods=["POST"])
def rename_saved_search(index):
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
        
    new_name = request.form.get("new_name", "").strip()
    if not new_name:
        return redirect("/")

    # Get current saved searches from database
    searches = load_saved_searches()
    
    if 0 <= index < len(searches):
        old_search = searches[index]
        
        engine = get_db_connection()
        if engine:
            try:
                with engine.connect() as conn:
                    conn.execute(text("""
                        UPDATE saved_searches 
                        SET name = :new_name 
                        WHERE name = :old_name AND user_id = :user_id
                    """), {
                        "new_name": new_name,
                        "old_name": old_search["name"],
                        "user_id": get_current_user_id()
                    })
                    conn.commit()
                    
                print(f"✅ Renamed search from '{old_search['name']}' to '{new_name}'")
                
            except Exception as e:
                print(f"❌ Error renaming search: {e}")
                
    return redirect("/")


@app.route("/schedule", methods=["POST"])
def schedule():
    index = int(request.form.get("search_index"))
    saved_searches = load_saved_searches()

    if 0 <= index < len(saved_searches):
        selected_search = saved_searches[index]
        return render_template("schedule.html", search=selected_search, index=index)
    else:
        return "Invalid search index", 400
 

@app.route("/save_schedule", methods=["POST"])
def save_schedule():
    # ADD THIS CHECK AT THE BEGINNING
    if not check_feature_access('scheduling'):
        return render_template("upgrade.html", 
                             feature="scheduled searches",
                             current_plan="free")
        
    # Your existing code continues unchanged...
    index = int(request.form.get("search_index"))
    frequency = request.form.get("frequency")

    search_history = load_saved_searches()

    if 0 <= index < len(search_history):
        search = search_history[index]
        engine = get_db_connection()
        
        if engine:
            try:
                with engine.connect() as conn:
                    conn.execute(text("""
                        UPDATE saved_searches 
                        SET schedule = :schedule 
                        WHERE name = :name AND user_id = :user_id
                    """), {
                        "schedule": frequency,
                        "name": search["name"],
                        "user_id": get_current_user_id()
                    })
                    conn.commit()
                return '', 200
            except Exception as e:
                print(f"Error updating schedule: {e}")
                return "Database error", 500
        return "No database connection", 500
    else:
        return "Invalid search index", 400

@app.route("/download_scheduled/<search_name>")
def download_scheduled(search_name):
    engine = get_db_connection()
    if not engine:
        return "Database connection error", 500
    
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT file_data, filename FROM scheduled_files 
                WHERE search_name = :search_name AND user_id = :user_id
            """), {
                "search_name": search_name,
                "user_id": get_current_user_id()
            })
            
            row = result.fetchone()
            if not row:
                return f"No file found for scheduled search: {search_name}", 404
            
            file_data, filename = row
            
            # Create BytesIO object from database binary data
            file_buffer = BytesIO(file_data)
            file_buffer.seek(0)
            
            return send_file(
                file_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                download_name=filename,
                as_attachment=True
            )
            
    except Exception as e:
        print(f"Error downloading scheduled file: {e}")
        return "Error downloading file", 500

@app.route("/api/saved_searches")
def api_saved_searches():
    searches = load_saved_searches()
    return json.dumps(searches), 200, {'Content-Type': 'application/json'}

@app.route("/saved_searches_partial")
def saved_searches_partial():
    searches = check_excel_files_for_searches(load_saved_searches())
    return render_template("partials/saved_searches.html", saved_searches=searches, has_scheduling_access=check_feature_access('scheduling'))

from flask import send_file
import zipfile
import io

@app.route("/download_selected", methods=["POST"])
def download_selected():
    print("🚨 FLASK ROUTE HIT!! /download_selected was called!")
    selected = request.form.getlist("selected_files")
    print("✅ /download_selected triggered!")
    print("🧾 Selected files:", selected) 

    today_str = datetime.now().strftime("%d_%B_%Y")

    if not selected:
        print("❌ No files selected.")
        return redirect("/")
     
     # Case B: Zip selected results
    # Get files from database instead of filesystem
    files_data = []
    engine = get_db_connection()

    if engine:
        try:
            with engine.connect() as conn:
                for name in selected:
                    result = conn.execute(text("""
                        SELECT file_data, filename FROM scheduled_files 
                        WHERE search_name = :search_name AND user_id = :user_id
                    """), {
                        "search_name": name,
                        "user_id": get_current_user_id()
                    })
                
                    row = result.fetchone()
                    if row:
                        files_data.append({
                            "filename": row[1],
                            "data": row[0]
                        })
        except Exception as e:
            print(f"❌ Database error: {e}")
            return redirect("/")
    
        
    if not files_data:
        print("❌ No matching files found in database.")
        return redirect("/")

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zipf:
        for file_info in files_data:
            # Add the binary data directly to the zip
            zipf.writestr(file_info["filename"], file_info["data"])

    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name="Selected_Results.zip", mimetype="application/zip")
    
@app.route("/create_test_files")
def create_test_files():
    import os
    import pandas as pd
    from datetime import datetime
    
    # Create the folder
    os.makedirs("scheduled_results", exist_ok=True)
    
    # Get saved searches and create test Excel files for each
    searches = load_saved_searches()
    created_files = []
    
    for search in searches:
        safe_name = search["name"].replace(" ", "_")
        date_str = datetime.now().strftime("%d_%B_%Y")
        filename = f"scheduled_results/{safe_name}_{date_str}.xlsx"
        
        # Create a simple test Excel file
        test_data = [{"title": "Test Job", "company": "Test Company", "location": "Test Location", "link": "http://test.com", "description": "Test description"}]
        df = pd.DataFrame(test_data)
        df.to_excel(filename, index=False)
        created_files.append(filename)
    
    return f"Created test files: {created_files}<br><br><a href='/debug'>Check debug again</a><br><a href='/'>Go back to main page</a>"

@app.route("/clean_test_files")
def clean_test_files():
    import os, glob
    files = glob.glob("scheduled_results/*.xlsx")
    deleted_count = 0
    for f in files:
        try:
            os.remove(f)
            deleted_count += 1
        except:
            pass
    return f"Deleted {deleted_count} test files. <br><br><a href='/debug'>Check debug again</a><br><a href='/'>Go back to main page</a>"


@app.route("/debug_files")
def debug_files():
    import os
    import stat
    
    file_info = {}
    
    if os.path.exists(HISTORY_FILE):
        file_stats = os.stat(HISTORY_FILE)
        file_info = {
            "exists": True,
            "size": file_stats.st_size,
            "permissions": oct(file_stats.st_mode),
            "modified": datetime.fromtimestamp(file_stats.st_mtime).strftime("%d %B %Y %H:%M:%S")
        }
        
        # Try to read the current content
        try:
            with open(HISTORY_FILE, "r") as f:
                content = f.read()
                file_info["content_preview"] = content[:500]
        except Exception as e:
            file_info["read_error"] = str(e)
    else:
        file_info = {"exists": False}
    
    return f"<pre>{json.dumps(file_info, indent=2)}</pre>"

@app.route("/debug_full_file")
def debug_full_file():
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            content = f.read()
        return f"<pre>{content}</pre>"
    except Exception as e:
        return f"Error reading file: {e}"

@app.route("/test_db")
def test_db():
    try:
        engine = get_db_connection()
        print(f"Engine result: {engine}")
        if engine:
            with engine.connect() as conn:
                result = conn.execute(text("SELECT 1"))
                return "Database connection works!"
        else:
            return "No database engine created"
    except Exception as e:
        return f"Database error: {e}"

@app.route("/debug_env")
def debug_env():
    import os
    db_vars = {}
    for key, value in os.environ.items():
        if any(word in key.upper() for word in ['DATABASE', 'POSTGRES', 'PG']):
            db_vars[key] = value[:20] + "..." if len(value) > 20 else value
    return f"<pre>{json.dumps(db_vars, indent=2)}</pre>"

@app.route("/check_files")
def check_files():
    import os, glob
    files = glob.glob("scheduled_results/*.xlsx")
    return f"Excel files found: {files}"


@app.route("/test_manual_run")
def test_manual_run():
    print("🚨 TEST_MANUAL_RUN: Route was called!")
    try:
        print("🚨 TEST_MANUAL_RUN: About to call run_scheduled_searches()")
        run_scheduled_searches()
        print("🚨 TEST_MANUAL_RUN: run_scheduled_searches() completed")
        return "Manual test completed - check logs and /check_files"
    except Exception as e:
        print(f"🚨 TEST_MANUAL_RUN: Exception caught: {e}")
        return f"Error: {e}"

@app.route("/debug_excel_detection")
def debug_excel_detection():
    searches = load_saved_searches()
    searches_with_files = check_excel_files_for_searches(searches)
    
    # Check what files actually exist
    all_files = glob.glob("scheduled_results/*.xlsx")
    
    debug_info = {
        "files_on_disk": all_files,
        "search_results": []
    }
    
    for search in searches_with_files:
        safe_name = search["name"].replace(" ", "_")
        pattern = f"scheduled_results/{safe_name}_*.xlsx"
        matching_files = glob.glob(pattern)
        
        debug_info["search_results"].append({
            "original_name": search["name"],
            "safe_name": safe_name,
            "pattern": pattern,
            "files_found": matching_files,
            "has_excel": search.get("has_excel", False)
        })
    
    return f"<pre>{json.dumps(debug_info, indent=2)}</pre>"

@app.route("/debug_database_files")
def debug_database_files():
    engine = get_db_connection()
    if not engine:
        return "No database connection"
    
    try:
        with engine.connect() as conn:
            # Check what files exist in database
            result = conn.execute(text("SELECT search_name, user_id, filename, created_at FROM scheduled_files"))
            files = []
            for row in result:
                files.append({
                    "search_name": row[0],
                    "user_id": row[1], 
                    "filename": row[2],
                    "created_at": str(row[3])
                })
            
            # Check current user ID
            current_user = get_current_user_id()
            
            return f"<pre>Current User ID: {current_user}\n\nFiles in database:\n{json.dumps(files, indent=2)}</pre>"
    except Exception as e:
        return f"Error: {e}"

@app.route("/checkout")
def checkout():
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
    
    # Simple placeholder for now - no actual payment processing
    return """
    <div style="font-family: Arial; text-align: center; padding: 50px; background: #f5f5f5; min-height: 100vh;">
        <div style="background: white; max-width: 400px; margin: 0 auto; padding: 40px; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.1);">
            <h2 style="color: #1e3a8a; margin-bottom: 20px;">💳 Checkout Coming Soon</h2>
            <p style="color: #666; margin-bottom: 20px;">Stripe payment integration will be added here.</p>
            <p style="color: #666; margin-bottom: 30px;">For now, this proves the upgrade flow works correctly!</p>
            
            <div style="background: #e0f2fe; padding: 15px; border-radius: 8px; margin-bottom: 20px;">
                <strong style="color: #0277bd;">Pro Plan - $8/month</strong><br>
                <small style="color: #0277bd;">All premium features included</small>
            </div>
            
            <a href="/app" style="display: inline-block; background: #1e3a8a; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; font-weight: 600;">
                ← Back to App
            </a>
        </div>
    </div>
    """

# Debug: Print all registered routes
print("🔍 DEBUG: Registered routes:")
for rule in app.url_map.iter_rules():
    print(f"  {rule.rule} -> {rule.methods} -> {rule.endpoint}")


@app.route("/test")
def test():
    return render_template("index.html", info="⏳ Indeed search is coming soon! We're currently testing this feature.", jobs=[])


@app.route("/debug_files/<filename>")
def download_debug_file(filename):
    """Download debug files like screenshots and page source"""
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
    
    try:
        import os
        if os.path.exists(filename):
            return send_file(filename, as_attachment=True)
        else:
            return f"File {filename} not found", 404
    except Exception as e:
        return f"Error downloading file: {e}", 500

@app.route("/debug_packages")
def debug_packages():
    import sys
    import pkg_resources
    
    installed_packages = [str(d) for d in pkg_resources.working_set]
    jobspy_packages = [p for p in installed_packages if 'jobspy' in p.lower()]
    
    return f"<pre>All jobspy-related packages: {jobspy_packages}\n\nPython path: {sys.path}\n\nAll packages: {installed_packages}</pre>"

@app.route("/debug_jobspy")
def debug_jobspy():
    import sys
    import subprocess
    
    try:
        # Check if jobspy is installed via pip list
        result = subprocess.run([sys.executable, '-m', 'pip', 'list'], 
                              capture_output=True, text=True)
        pip_list = result.stdout
        
        # Find jobspy lines
        jobspy_lines = [line for line in pip_list.split('\n') 
                       if 'jobspy' in line.lower()]
        
        # Try import
        import_status = "Failed"
        import_error = ""
        try:
            import jobspy
            import_status = f"SUCCESS - Version: {getattr(jobspy, '__version__', 'Unknown')}"
        except Exception as e:
            import_status = "FAILED"
            import_error = str(e)
        
        # Check sys.path
        python_paths = '\n'.join(sys.path)
        
        return f"""
        <pre>
        JOBSPY INSTALLATION CHECK:
        ========================
        
        In pip list: {len(jobspy_lines)} matches
        {jobspy_lines}
        
        Import Test: {import_status}
        Import Error: {import_error}
        
        Python Paths:
        {python_paths}
        
        Full pip list (first 20 lines):
        {chr(10).join(pip_list.split(chr(10))[:20])}
        </pre>
        """
    except Exception as e:
        return f"<pre>Debug route failed: {e}</pre>"

@app.route("/debug_saved_search/<search_name>")
def debug_saved_search(search_name):
    engine = get_db_connection()
    if not engine:
        return "No database connection"
    
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT name, criteria FROM saved_searches 
                WHERE name = :search_name AND user_id = :user_id
            """), {
                "search_name": search_name,
                "user_id": get_current_user_id()
            })
            
            row = result.fetchone()
            if row:
                name, criteria = row
                return f"<pre>Search: {name}\nCriteria: {criteria}</pre>"
            else:
                return f"Search '{search_name}' not found"
    except Exception as e:
        return f"Error: {e}"

@app.route("/debug_all_searches")
def debug_all_searches():
    engine = get_db_connection()
    if not engine:
        return "No database connection"
    
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT name, criteria FROM saved_searches 
                WHERE user_id = :user_id
                ORDER BY id DESC
            """), {"user_id": get_current_user_id()})
            
            output = "<h2>All Saved Searches:</h2><pre>"
            for row in result:
                name, criteria = row
                output += f"\nName: '{name}'\n"
                output += f"Criteria: {criteria}\n"
                output += "-" * 50 + "\n"
            output += "</pre>"
            return output
    except Exception as e:
        return f"Error: {e}"

@app.route("/debug_saved_search_source/<int:index>")
def debug_saved_search_source(index):
    searches = load_saved_searches()
    if 0 <= index < len(searches):
        search = searches[index]
        criteria = search["criteria"]
        source = criteria.get("source", "NOT_FOUND")
        
        return f"""
        <pre>
        Search Name: {search['name']}
        All Criteria: {criteria}
        Source Value: '{source}'
        Source Type: {type(source)}
        </pre>
        <br><a href="/app">Back to app</a>
        """
    else:
        return "Invalid search index"

@app.route("/force_scheduler_test")
def force_scheduler_test():
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
    
    try:
        run_scheduled_searches()
        return "Forced scheduler test completed - check Railway logs for email results"
    except Exception as e:
        return f"Scheduler test failed: {e}"

@app.route("/basic_test")
def basic_test():
    logger.info("BASIC TEST: This message should appear in logs")
    return "Basic test completed"


@app.route("/test_gmail_direct")
def test_gmail_direct():
    login_redirect = require_login()
    if login_redirect:
        return login_redirect
    
    try:
        import smtplib
        from email.message import EmailMessage
        
        # Use your exact Gmail settings
        smtp_server = config["email_settings"]["smtp_server"] 
        smtp_port = config["email_settings"]["smtp_port"]
        sender_email = config["email_settings"]["sender_email"]
        sender_password = config["email_settings"]["sender_password"]
        
        logger.info(f"Testing direct Gmail connection to {smtp_server}:{smtp_port}")
        
        # Simple test email
        msg = EmailMessage()
        msg["Subject"] = "Railway Gmail Test"
        msg["From"] = sender_email
        msg["To"] = sender_email  # Send to yourself
        msg.set_content("This is a test email to verify Gmail SMTP works from Railway")
        
        with smtplib.SMTP(smtp_server, smtp_port) as smtp:
            smtp.starttls()
            smtp.login(sender_email, sender_password)
            smtp.send_message(msg)
        
        logger.info("✅ Gmail SMTP test successful")
        return "Gmail test successful - check your email"
        
    except Exception as e:
        logger.error(f"❌ Gmail SMTP test failed: {e}")
        return f"Gmail test failed: {e}"

@app.route("/admin/delete-bots")
def delete_bot_accounts():
    """One-time route to delete known bot/spam accounts. Only callable by the app owner (user id 3)."""
    if get_current_user_id() != 3:
        return "Unauthorized", 403

    bot_ids = (10, 12, 13, 14)
    engine = get_db_connection()
    if not engine:
        return "No database connection", 500

    results = []
    try:
        with engine.connect() as conn:
            # Find every table that has a user_id column and delete from it first
            tables_with_user_id = conn.execute(text("""
                SELECT table_name FROM information_schema.columns
                WHERE table_schema = 'public'
                AND column_name = 'user_id'
                AND table_name != 'users'
                ORDER BY table_name
            """)).fetchall()

            for (table,) in tables_with_user_id:
                deleted = conn.execute(text(
                    f"DELETE FROM {table} WHERE user_id = ANY(:ids)"
                ), {"ids": list(bot_ids)}).rowcount
                results.append(f"{table}: {deleted} rows deleted")

            # Now delete the user accounts themselves
            deleted_users = conn.execute(text(
                "DELETE FROM users WHERE id = ANY(:ids)"
            ), {"ids": list(bot_ids)}).rowcount
            results.append(f"users: {deleted_users} accounts deleted")

            conn.commit()

        return "<pre>Bot cleanup complete:\n\n" + "\n".join(results) + "</pre>"
    except Exception as e:
        return f"Error during cleanup: {e}", 500


# ==================== AI MATCH TOOL ROUTES ====================

@app.route("/ai-match", methods=["GET", "POST"])
def ai_match():
    """AI Match Tool main page"""
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user_id = session['user_id']
    username = session.get('username', 'User')

    # Get user's CVs
    user_cvs = get_user_cvs(user_id)

    # Get recent analyses
    recent_analyses = get_user_analyses(user_id, limit=5)

    # Check if user has a master template
    master_template = get_user_master_template(user_id)
    has_master_template = master_template is not None

    return render_template('ai_match.html',
                         username=username,
                         user_cvs=user_cvs,
                         recent_analyses=recent_analyses,
                         has_master_template=has_master_template)


@app.route("/upload-cv", methods=["POST"])
def upload_cv():
    """Handle CV upload"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        # Get uploaded file
        if 'cv_file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['cv_file']
        cv_name = request.form.get('cv_name', file.filename)

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Validate file type
        allowed_extensions = {'pdf', 'docx', 'doc'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

        if file_ext not in allowed_extensions:
            return jsonify({'error': 'Invalid file type. Only PDF and DOCX allowed'}), 400

        # Read file data
        file_data = file.read()

        # Check file size (max 5MB)
        if len(file_data) > 5 * 1024 * 1024:
            return jsonify({'error': 'File too large. Maximum size is 5MB'}), 400

        # Extract text from CV
        extracted_text = parse_cv(file_data, file_ext)

        if not extracted_text or not extracted_text.strip():
            return jsonify({'error': 'CV appears to be empty or text extraction failed'}), 400

        # Save to database
        cv_id = save_cv_to_db(user_id, cv_name, file_data, file_ext, extracted_text)

        if not cv_id:
            return jsonify({'error': 'Failed to save CV'}), 500

        log_user_activity('cv_upload', f'Uploaded CV: {cv_name}')

        return jsonify({
            'success': True,
            'cv_id': cv_id,
            'cv_name': cv_name,
            'message': 'CV uploaded successfully'
        })

    except Exception as e:
        print(f"Error uploading CV: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Upload error: {str(e)}'}), 500


@app.route("/my-resume-template", methods=["GET"])
def my_resume_template():
    """Display master resume template management page"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    # Get user's current master template
    master_template = get_user_master_template(user_id)

    return render_template('my_resume_template.html',
                         master_template=master_template)


@app.route("/upload-master-template", methods=["POST"])
def upload_master_template():
    """Handle master resume template upload"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        # Get uploaded file
        if 'template_file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['template_file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Validate file type - .docx and .txt allowed
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

        if file_ext not in ('docx', 'txt'):
            return jsonify({'error': 'Invalid file type. Only .docx or .txt files are allowed'}), 400

        # Read file data
        file_data = file.read()

        # Check file size (max 2MB for templates)
        if len(file_data) > 2 * 1024 * 1024:
            return jsonify({'error': 'File too large. Maximum size is 2MB'}), 400

        # Extract text based on file type
        if file_ext == 'txt':
            extracted_text = file_data.decode('utf-8', errors='replace')
        else:
            extracted_text = extract_text_from_docx(file_data)

        if not extracted_text or not extracted_text.strip():
            return jsonify({'error': 'Template appears to be empty or text extraction failed'}), 400

        # Save to database
        template_id = save_master_template(user_id, extracted_text, file.filename)

        if not template_id:
            return jsonify({'error': 'Failed to save template'}), 500

        log_user_activity('master_template_upload', f'Uploaded master template: {file.filename}')

        return jsonify({
            'success': True,
            'template_id': template_id,
            'message': 'Master template uploaded successfully'
        })

    except Exception as e:
        print(f"Error uploading master template: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Upload error: {str(e)}'}), 500


@app.route("/analyze-match", methods=["POST"])
def analyze_match():
    """Analyze job match using AI"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        data = request.get_json()

        cv_id = data.get('cv_id')
        job_posting = data.get('job_posting')

        # Validate inputs
        if not cv_id or not job_posting:
            return jsonify({'error': 'Missing required fields'}), 400

        # Get CV
        cv_data = get_cv_by_id(cv_id, user_id)
        if not cv_data:
            return jsonify({'error': 'CV not found'}), 404

        cv_text = cv_data['extracted_text']

        # Extract job title and company from posting using AI
        job_title, job_company = extract_job_info_from_posting(job_posting)

        # Analyze with AI using user's selected prompt template (this will raise exception if it fails)
        analysis_result = analyze_job_match_with_ai(
            cv_text,
            job_title,
            job_company,
            job_posting,
            user_id
        )

        # Save analysis to database
        analysis_id = save_job_analysis(
            user_id,
            cv_id,
            job_title,
            job_company,
            job_posting,
            analysis_result
        )

        log_user_activity('job_match_analysis', f'Analyzed: {job_title}')

        return jsonify({
            'success': True,
            'analysis_id': analysis_id,
            'analysis': analysis_result
        })

    except Exception as e:
        print(f"Error analyzing match: {e}")
        import traceback
        traceback.print_exc()
        error_msg = str(e)
        # Check for common API errors
        if "api_key" in error_msg.lower() or "anthropic_api_key" in error_msg.lower():
            error_msg = "API key not configured. Please set ANTHROPIC_API_KEY environment variable in Railway."
        elif "json" in error_msg.lower():
            error_msg = "AI returned invalid response format. Please try again."
        elif "anthropic" in error_msg.lower():
            error_msg = "Anthropic API error. Check API key and try again."
        return jsonify({'error': f'Analysis failed: {error_msg}'}), 500


@app.route("/analyze-match-with-template", methods=["POST"])
def analyze_match_with_template():
    """Analyze job match using AI with master template and prompt caching"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        data = request.get_json()
        job_posting = data.get('job_posting')

        # Validate input
        if not job_posting:
            return jsonify({'error': 'Missing job posting'}), 400

        # Get user's master template
        master_template = get_user_master_template(user_id)
        if not master_template:
            return jsonify({'error': 'No master template found. Please upload one first.'}), 404

        template_text = master_template['template_text']

        # Extract job title and company from posting using AI
        job_title, job_company = extract_job_info_from_posting(job_posting)

        # Analyze with AI using master template with caching
        analysis_result = analyze_job_match_with_master_template(
            template_text,
            job_title,
            job_company,
            job_posting,
            user_id
        )

        # Save analysis to database (use NULL for cv_id since we're using master template)
        analysis_id = save_job_analysis(
            user_id,
            None,  # No specific CV, using master template
            job_title,
            job_company,
            job_posting,
            analysis_result
        )

        log_user_activity('job_match_analysis_cached', f'Analyzed with master template: {job_title}')

        return jsonify({
            'success': True,
            'analysis_id': analysis_id,
            'analysis': analysis_result
        })

    except Exception as e:
        print(f"Error analyzing match with template: {e}")
        import traceback
        traceback.print_exc()
        error_msg = str(e)
        # Check for common API errors
        if "api_key" in error_msg.lower() or "anthropic_api_key" in error_msg.lower():
            error_msg = "API key not configured. Please set ANTHROPIC_API_KEY environment variable in Railway."
        elif "json" in error_msg.lower():
            error_msg = "AI returned invalid response format. Please try again."
        elif "anthropic" in error_msg.lower():
            error_msg = "Anthropic API error. Check API key and try again."
        return jsonify({'error': f'Analysis failed: {error_msg}'}), 500


@app.route("/delete-cv/<int:cv_id>", methods=["POST"])
def delete_cv(cv_id):
    """Delete a CV"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        engine = get_db_connection()
        if not engine:
            return jsonify({'error': 'Database error'}), 500

        with engine.connect() as conn:
            # Verify ownership and delete
            result = conn.execute(text("""
                DELETE FROM user_cvs
                WHERE id = :cv_id AND user_id = :user_id
                RETURNING cv_name
            """), {"cv_id": cv_id, "user_id": user_id})

            deleted_cv = result.fetchone()
            conn.commit()

            if not deleted_cv:
                return jsonify({'error': 'CV not found'}), 404

            log_user_activity('cv_delete', f'Deleted CV: {deleted_cv[0]}')

            return jsonify({'success': True, 'message': 'CV deleted'})

    except Exception as e:
        print(f"Error deleting CV: {e}")
        return jsonify({'error': 'Failed to delete CV'}), 500


@app.route("/debug-analyses", methods=["GET"])
def debug_analyses():
    """Debug endpoint to see raw database state"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']
    engine = get_db_connection()
    if not engine:
        return jsonify({'error': 'Database error'}), 500

    try:
        with engine.connect() as conn:
            # Get all analyses for this user
            analyses_result = conn.execute(text("""
                SELECT id, cv_id, job_title, job_company, match_score, created_at
                FROM job_analyses
                WHERE user_id = :user_id
                ORDER BY created_at DESC
                LIMIT 10
            """), {"user_id": user_id})

            analyses = []
            for row in analyses_result:
                analyses.append({
                    'id': row[0],
                    'cv_id': row[1],
                    'job_title': row[2],
                    'job_company': row[3],
                    'match_score': row[4],
                    'created_at': str(row[5])
                })

            # Get all CVs for this user
            cvs_result = conn.execute(text("""
                SELECT id, cv_name, created_at
                FROM user_cvs
                WHERE user_id = :user_id
                ORDER BY created_at DESC
            """), {"user_id": user_id})

            cvs = []
            for row in cvs_result:
                cvs.append({
                    'id': row[0],
                    'cv_name': row[1],
                    'created_at': str(row[2])
                })

            return jsonify({
                'user_id': user_id,
                'analyses': analyses,
                'cvs': cvs,
                'analysis_count': len(analyses),
                'cv_count': len(cvs)
            })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== CV CUSTOMIZATION ROUTES ====================

@app.route("/debug-template", methods=["GET"])
def debug_template():
    """Debug endpoint to view master template structure"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']
    master_template = get_user_master_template(user_id)

    if not master_template:
        return "No master template found", 404

    lines = master_template['template_text'].split('\n')

    # Build HTML output with line numbers
    html = "<html><head><style>body{font-family:monospace;font-size:12px;} .line{padding:2px 0;} .linenum{color:#888;width:60px;display:inline-block;} .content{white-space:pre-wrap;}</style></head><body>"
    html += f"<h2>Master Template Debug (Total lines: {len(lines)})</h2>"
    html += "<p>Showing all lines</p>"

    for i, line in enumerate(lines):  # Show ALL lines
        # Highlight certain keywords
        line_display = line
        if 'JP MORGAN' in line.upper() or 'CATEGORY:' in line:
            line_display = f'<strong style="background:yellow;">{line}</strong>'
        elif line.strip().startswith('•'):
            line_display = f'<span style="color:blue;">{line}</span>'

        html += f'<div class="line"><span class="linenum">{i}:</span><span class="content">{line_display}</span></div>'

    html += "</body></html>"
    return html


@app.route("/debug-bullets", methods=["GET"])
def debug_bullets():
    """Debug endpoint to test bullet parsing"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']
    master_template = get_user_master_template(user_id)

    if not master_template:
        return "No master template found", 404

    # Parse bullets
    bullets = parse_bullets_from_template(master_template['template_text'])

    # Build HTML output
    html = """
    <html>
    <head>
        <style>
            body { font-family: Arial, sans-serif; margin: 20px; }
            .summary { background: #f0f0f0; padding: 15px; margin-bottom: 20px; border-radius: 5px; }
            .category { margin-top: 30px; padding: 10px; background: #e8f4f8; border-left: 4px solid #0066cc; }
            .bullet { margin: 10px 0; padding: 10px; background: white; border: 1px solid #ddd; border-radius: 3px; }
            .bullet-number { font-weight: bold; color: #0066cc; }
            .error { color: red; font-weight: bold; }
            .success { color: green; font-weight: bold; }
        </style>
    </head>
    <body>
        <h1>Bullet Parser Debug</h1>
    """

    # Summary
    expected = 66
    actual = len(bullets)
    status_class = "success" if actual == expected else "error"

    html += f"""
    <div class="summary">
        <h2>Summary</h2>
        <p class="{status_class}">Total bullets found: {actual} (Expected: {expected})</p>
    """

    # Category breakdown
    category_counts = {}
    for bullet in bullets:
        cat = bullet['category']
        category_counts[cat] = category_counts.get(cat, 0) + 1

    html += "<h3>Category Breakdown:</h3><ul>"
    for cat, count in category_counts.items():
        html += f"<li><strong>{cat}</strong>: {count} bullets</li>"
    html += "</ul></div>"

    # Show all bullets grouped by category
    current_cat = None
    for bullet in bullets:
        if bullet['category'] != current_cat:
            if current_cat is not None:
                html += "</div>"  # Close previous category
            current_cat = bullet['category']
            html += f'<div class="category"><h3>Category: {current_cat}</h3>'

        html += f"""
        <div class="bullet">
            <span class="bullet-number">Bullet #{bullet['number']}</span><br>
            {bullet['text']}
        </div>
        """

    if current_cat is not None:
        html += "</div>"  # Close last category

    html += "</body></html>"
    return html


@app.route("/customize-cv/start/<int:analysis_id>", methods=["GET"])
def customize_cv_start(analysis_id):
    """Entry point for CV customization - creates session and starts flow"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    try:
        # Get the analysis
        analysis = get_analysis_by_id(analysis_id, user_id)
        if not analysis:
            return "Analysis not found", 404

        # Get user's master template
        master_template = get_user_master_template(user_id)
        if not master_template:
            return redirect(f'/my-resume-template?error=no_template')

        # Create CV customization session
        cv_session_id = create_cv_customization_session(
            user_id,
            analysis_id,
            analysis['job_title'],
            analysis['job_company']
        )

        if not cv_session_id:
            return "Failed to create customization session", 500

        # Store session ID in Flask session
        session['cv_session_id'] = cv_session_id

        # Redirect to headline selection
        return redirect('/customize-cv/headline')

    except Exception as e:
        print(f"Error starting CV customization: {e}")
        import traceback
        traceback.print_exc()
        return "Error starting CV customization", 500


@app.route("/customize-cv/resume/<int:cv_session_id>", methods=["GET"])
def customize_cv_resume(cv_session_id):
    """Resume an existing CV customization session"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    try:
        # Get the CV session
        cv_session = get_cv_session(cv_session_id)

        if not cv_session:
            return redirect('/ai-match?error=session_not_found')

        # Verify it belongs to this user
        if cv_session['user_id'] != user_id:
            return redirect('/ai-match?error=unauthorized')

        # Store session ID in Flask session
        session['cv_session_id'] = cv_session_id

        # Determine which step to resume
        if not cv_session.get('selected_headline'):
            return redirect('/customize-cv/headline')
        elif not cv_session.get('selected_roles'):
            return redirect('/customize-cv/roles')
        elif not cv_session.get('approved_bullets') or len(cv_session.get('approved_bullets', [])) < 3:
            return redirect('/customize-cv/bullets?role=0')
        else:
            return redirect('/customize-cv/preview')

    except Exception as e:
        print(f"Error resuming CV customization: {e}")
        import traceback
        traceback.print_exc()
        return redirect('/ai-match?error=resume_failed')


@app.route("/customize-cv/headline", methods=["GET", "POST"])
def customize_cv_headline():
    """Step 1: Headline selection with AI recommendations"""
    if 'user_id' not in session or 'cv_session_id' not in session:
        return redirect('/ai-match')

    user_id = session['user_id']
    cv_session_id = session['cv_session_id']

    if request.method == "POST":
        # Save selected headline
        data = request.json
        selected_headline = data.get('headline')

        if update_cv_session_headline(cv_session_id, selected_headline):
            return jsonify({'success': True, 'next_step': '/customize-cv/roles'})
        else:
            return jsonify({'error': 'Failed to save headline'}), 500

    # GET request - show headline selection page
    try:
        # Get CV session data
        cv_session = get_cv_session(cv_session_id)
        if not cv_session:
            return "Session not found", 404

        # Get analysis for job description
        analysis = get_analysis_by_id(cv_session['analysis_id'], user_id)
        if not analysis:
            return "Analysis not found", 404

        # Get master template
        master_template = get_user_master_template(user_id)
        if not master_template:
            return "Master template not found", 404

        # Try new career summary format first, fall back to legacy headline format
        headlines = parse_career_summaries_from_template(master_template['template_text'])
        if not headlines:
            headlines = parse_headlines_from_template(master_template['template_text'])

        if not headlines:
            return "No headlines found in master template", 400

        # Analyze headlines with AI
        ai_analysis = analyze_headlines_with_ai(
            headlines,
            analysis['job_description'],
            user_id
        )

        if not ai_analysis:
            return "Failed to analyze headlines", 500

        # Combine headlines with AI analysis
        headlines_with_analysis = []
        for i, headline in enumerate(headlines):
            # Find this headline in AI analysis
            headline_analysis = None

            # Check top 3
            for top in ai_analysis.get('top_3', []):
                if top['index'] == i:
                    headline_analysis = {
                        'tier': 'top',
                        'match_score': top['match_score'],
                        'reasons': top['reasons'],
                        'suggested_adaptation': top.get('suggested_adaptation')
                    }
                    break

            # Check other headlines
            if not headline_analysis:
                for other in ai_analysis.get('other_headlines', []):
                    if other['index'] == i:
                        headline_analysis = {
                            'tier': 'other',
                            'match_score': other['match_score'],
                            'weakness': other.get('weakness')
                        }
                        break

            if headline_analysis:
                headlines_with_analysis.append({
                    **headline,
                    **headline_analysis
                })

        # Get recommended headline
        recommended_index = ai_analysis.get('recommended_headline_index', 0)
        recommended_headline = next(
            (h for h in headlines_with_analysis if h['id'] == recommended_index),
            headlines_with_analysis[0] if headlines_with_analysis else None
        )

        return render_template('customize_cv_headline.html',
                             job_title=cv_session['job_title'],
                             job_company=cv_session['job_company'],
                             headlines=headlines_with_analysis,
                             recommended_headline=recommended_headline,
                             analysis_id=cv_session['analysis_id'])

    except Exception as e:
        print(f"Error in headline selection: {e}")
        import traceback
        traceback.print_exc()
        return "Error loading headline selection", 500


@app.route("/customize-cv/roles", methods=["GET", "POST"])
def customize_cv_roles():
    """Step 2: User selects which roles AI should tailor bullets for."""
    if 'user_id' not in session or 'cv_session_id' not in session:
        return redirect('/ai-match')

    user_id = session['user_id']
    cv_session_id = session['cv_session_id']

    if request.method == "POST":
        data = request.json
        selected_role_ids = data.get('selected_role_ids', [])

        cv_session = get_cv_session(cv_session_id)
        master_template = get_user_master_template(user_id)
        if not cv_session or not master_template:
            return jsonify({'error': 'Session or template not found'}), 404

        all_roles = parse_roles_from_template(master_template['template_text'])
        selected = [r for r in all_roles if r['id'] in selected_role_ids]

        if not selected:
            return jsonify({'error': 'No roles selected'}), 400

        if update_cv_session_selected_roles(cv_session_id, selected):
            return jsonify({'success': True, 'next_step': '/customize-cv/bullets?role=0'})
        return jsonify({'error': 'Failed to save roles'}), 500

    # GET — show role selection
    try:
        cv_session = get_cv_session(cv_session_id)
        if not cv_session:
            return "Session not found", 404

        analysis = get_analysis_by_id(cv_session['analysis_id'], user_id)
        if not analysis:
            return "Analysis not found", 404

        master_template = get_user_master_template(user_id)
        if not master_template:
            return "Master template not found", 404

        roles = parse_roles_from_template(master_template['template_text'])
        if not roles:
            return "No roles found in master template", 400

        scored_roles = score_roles_relevance(roles, analysis['job_description'])

        return render_template('customize_cv_roles.html',
                               job_title=cv_session['job_title'],
                               job_company=cv_session['job_company'],
                               selected_headline=cv_session['selected_headline'],
                               roles=scored_roles)

    except Exception as e:
        print(f"Error in role selection: {e}")
        import traceback
        traceback.print_exc()
        return "Error loading role selection", 500


@app.route("/customize-cv/bullets", methods=["GET", "POST"])
def customize_cv_bullets():
    """Step 2b: Per-role bullet selection. ?role=N selects which role to tailor."""
    if 'user_id' not in session or 'cv_session_id' not in session:
        return redirect('/ai-match')

    user_id = session['user_id']
    cv_session_id = session['cv_session_id']

    if request.method == "POST":
        data = request.json
        role_index = int(data.get('role_index', 0))
        approved_bullets = data.get('approved_bullets', [])
        approved_texts = data.get('approved_texts', {})

        cv_session = get_cv_session(cv_session_id)
        if not cv_session:
            return jsonify({'error': 'Session not found'}), 404

        selected_roles = cv_session.get('selected_roles', [])

        # Merge new approvals into flat approved_bullets list
        existing = cv_session.get('approved_bullets') or []
        role_company = selected_roles[role_index]['company'] if role_index < len(selected_roles) else ''
        # Remove any prior entries for this role, then add new ones
        existing = [b for b in existing if b.get('role_company') != role_company]
        for bullet_idx in approved_bullets:
            text = approved_texts.get(str(bullet_idx), '')
            if text:
                existing.append({'role_company': role_company, 'bullet_index': bullet_idx, 'approved_text': text})

        update_cv_session_approved_bullets(cv_session_id, existing)

        next_index = role_index + 1
        if next_index < len(selected_roles):
            return jsonify({'success': True, 'next_step': f'/customize-cv/bullets?role={next_index}'})
        return jsonify({'success': True, 'next_step': '/customize-cv/preview'})

    # GET request
    try:
        role_index = int(request.args.get('role', 0))

        cv_session = get_cv_session(cv_session_id)
        if not cv_session:
            return "Session not found", 404

        selected_roles = cv_session.get('selected_roles', [])
        if not selected_roles:
            return redirect('/customize-cv/roles')

        if role_index >= len(selected_roles):
            return redirect('/customize-cv/preview')

        current_role = selected_roles[role_index]
        role_key = current_role['company']

        analysis = get_analysis_by_id(cv_session['analysis_id'], user_id)
        if not analysis:
            return "Analysis not found", 404

        # Use cached analysis if available
        cached = (cv_session.get('bullet_analysis_by_role') or {}).get(role_key)
        if cached:
            print(f"✓ Using cached bullet analysis for {role_key}")
            ai_analysis = cached
        else:
            print(f"⚙ Running AI bullet analysis for {role_key}...")
            ai_analysis = analyze_bullets_for_role_with_ai(
                current_role, analysis['job_description'], user_id
            )
            if not ai_analysis:
                return "Failed to analyse bullets", 500
            update_cv_session_bullet_analysis_by_role(cv_session_id, role_key, ai_analysis)

        # Previously approved bullets for this role
        all_approved = cv_session.get('approved_bullets') or []
        role_approved = [b for b in all_approved if b.get('role_company') == role_key]

        return render_template('customize_cv_bullets.html',
                               job_title=cv_session['job_title'],
                               job_company=cv_session['job_company'],
                               selected_headline=cv_session['selected_headline'],
                               current_role=current_role,
                               role_index=role_index,
                               total_roles=len(selected_roles),
                               all_roles=selected_roles,
                               recommended_bullets=ai_analysis.get('recommended_bullets', []),
                               approved_bullets_data=role_approved,
                               analysis_id=cv_session['analysis_id'])

    except Exception as e:
        print(f"Error in bullet selection: {e}")
        import traceback
        traceback.print_exc()
        return "Error loading bullet selection", 500


@app.route("/customize-cv/preview", methods=["GET"])
def customize_cv_preview():
    """Step 3: Preview and download customized CV"""
    if 'user_id' not in session or 'cv_session_id' not in session:
        return redirect('/ai-match')

    user_id = session['user_id']
    cv_session_id = session['cv_session_id']

    try:
        # Get CV session data
        cv_session = get_cv_session(cv_session_id)
        if not cv_session:
            return "Session not found", 404

        # Get approved bullets
        approved_bullets_data = cv_session.get('approved_bullets', [])

        if not approved_bullets_data or len(approved_bullets_data) < 3:
            return redirect('/customize-cv/bullets?role=0')

        # Render preview template
        return render_template('customize_cv_preview.html',
                             job_title=cv_session['job_title'],
                             job_company=cv_session['job_company'],
                             selected_headline=cv_session['selected_headline'],
                             approved_bullets=approved_bullets_data,
                             analysis_id=cv_session['analysis_id'])

    except Exception as e:
        print(f"Error in CV preview: {e}")
        import traceback
        traceback.print_exc()
        return "Error loading preview", 500


@app.route("/api/customize-bullet-chat", methods=["POST"])
def customize_bullet_chat():
    """Handle chat messages for bullet customization"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        data = request.json
        bullet_text = data.get('bullet_text')
        user_message = data.get('user_message')
        chat_history = data.get('chat_history', [])

        if not bullet_text or not user_message:
            return jsonify({'error': 'Missing required fields'}), 400

        # Get CV session to access job description
        cv_session_id = session.get('cv_session_id')
        if not cv_session_id:
            return jsonify({'error': 'No active CV customization session'}), 400

        cv_session = get_cv_session(cv_session_id)
        if not cv_session:
            return jsonify({'error': 'Session not found'}), 404

        # Get job description
        analysis = get_analysis_by_id(cv_session['analysis_id'], user_id)
        if not analysis:
            return jsonify({'error': 'Analysis not found'}), 404

        job_description = analysis['job_description']

        # Build chat context from history
        messages = []
        for msg in chat_history:
            if msg['role'] in ['user', 'assistant']:
                messages.append({
                    "role": msg['role'],
                    "content": msg['content']
                })

        # Get AI response
        client = get_anthropic_client()
        system_prompt = get_user_system_prompt(user_id)

        response = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=1000,
            temperature=0.3,
            system=[
                {
                    "type": "text",
                    "text": f"""{system_prompt}

You are helping a job seeker refine their resume bullet point to better match this job description.

**JOB DESCRIPTION:**
{job_description}

**ORIGINAL BULLET:**
{bullet_text}

**YOUR ROLE:**
- Help the user refine this bullet based on their feedback
- When providing a refined version, return it as plain text (not JSON)
- Keep the bullet truthful - only suggest changes that enhance existing experience
- NEVER invent experiences or metrics the user doesn't have
- Focus on word choice, phrasing, and highlighting relevant aspects
- Incorporate JD keywords naturally where appropriate

**IMPORTANT:**
If you provide a refined bullet in your response, put it between [REFINED_BULLET] and [/REFINED_BULLET] tags so it can be extracted.

Example:
"Here's a stronger version that emphasizes your cross-functional leadership:

[REFINED_BULLET]
Led cross-functional teams of 15+ stakeholders across Technology, Operations, and Risk to implement AI governance framework, driving 63% portfolio growth while maintaining 100% regulatory compliance
[/REFINED_BULLET]

This version better highlights the leadership and stakeholder management aspects emphasized in the JD."
""",
                    "cache_control": {"type": "ephemeral"}
                }
            ],
            messages=messages
        )

        ai_response = response.content[0].text

        # Extract refined bullet if present
        refined_bullet = None
        if '[REFINED_BULLET]' in ai_response and '[/REFINED_BULLET]' in ai_response:
            start = ai_response.find('[REFINED_BULLET]') + len('[REFINED_BULLET]')
            end = ai_response.find('[/REFINED_BULLET]')
            refined_bullet = ai_response[start:end].strip()

        return jsonify({
            'success': True,
            'response': ai_response,
            'refined_bullet': refined_bullet
        })

    except Exception as e:
        print(f"Error in bullet chat: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to get AI response'}), 500


@app.route("/api/customize-headline-chat", methods=["POST"])
def customize_headline_chat():
    """Handle chat messages for headline customization"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        data = request.json
        headline_text = data.get('headline_text')
        user_message = data.get('user_message')
        chat_history = data.get('chat_history', [])

        if not headline_text or not user_message:
            return jsonify({'error': 'Missing required fields'}), 400

        cv_session_id = session.get('cv_session_id')
        if not cv_session_id:
            return jsonify({'error': 'No active CV customization session'}), 400

        cv_session = get_cv_session(cv_session_id)
        if not cv_session:
            return jsonify({'error': 'Session not found'}), 404

        analysis = get_analysis_by_id(cv_session['analysis_id'], user_id)
        if not analysis:
            return jsonify({'error': 'Analysis not found'}), 404

        job_description = analysis['job_description']

        messages = []
        for msg in chat_history:
            if msg['role'] in ['user', 'assistant']:
                messages.append({"role": msg['role'], "content": msg['content']})

        client = get_anthropic_client()
        system_prompt = get_user_system_prompt(user_id)

        response = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=1000,
            temperature=0.3,
            system=[
                {
                    "type": "text",
                    "text": f"""{system_prompt}

You are helping a job seeker refine their CV headline/professional summary to better match this job description.

**JOB DESCRIPTION:**
{job_description}

**CURRENT HEADLINE:**
{headline_text}

**YOUR ROLE:**
- Help the user refine this headline based on their feedback
- Keep it truthful — only suggest changes that enhance existing experience
- NEVER invent roles, skills, or experience the user doesn't have
- Focus on word choice, phrasing, and emphasising relevant aspects for the role
- Incorporate JD keywords naturally where appropriate
- Keep the headline concise and impactful (2-5 sentences)

**IMPORTANT:**
If you provide a refined headline in your response, put it between [REFINED_HEADLINE] and [/REFINED_HEADLINE] tags so it can be extracted.

Example:
"Here's a version that better highlights your AI governance experience:

[REFINED_HEADLINE]
Senior Programme Manager with 20 years' experience leading complex, regulatory-driven transformation programmes across Tier 1 global financial institutions, with deep expertise in AI governance and applied ML deployment at scale.
[/REFINED_HEADLINE]

This version directly addresses the Applied AI focus in the JD."
""",
                    "cache_control": {"type": "ephemeral"}
                }
            ],
            messages=messages
        )

        ai_response = response.content[0].text

        refined_headline = None
        if '[REFINED_HEADLINE]' in ai_response and '[/REFINED_HEADLINE]' in ai_response:
            start = ai_response.find('[REFINED_HEADLINE]') + len('[REFINED_HEADLINE]')
            end = ai_response.find('[/REFINED_HEADLINE]')
            refined_headline = ai_response[start:end].strip()

        return jsonify({
            'success': True,
            'response': ai_response,
            'refined_headline': refined_headline
        })

    except Exception as e:
        print(f"Error in headline chat: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to get AI response'}), 500


@app.route("/api/save-approved-bullets", methods=["POST"])
def save_approved_bullets():
    """Save approved bullets to avoid re-selection"""
    if 'user_id' not in session or 'cv_session_id' not in session:
        print("❌ Save bullets: Not authenticated or no CV session")
        return jsonify({'error': 'Not authenticated'}), 401

    cv_session_id = session['cv_session_id']
    print(f"💾 Saving bullets for session {cv_session_id}")

    try:
        data = request.json
        approved_bullets = data.get('approved_bullets', [])
        approved_texts = data.get('approved_texts', {})
        customized_bullets = data.get('customized_bullets', {})

        print(f"  Received: {len(approved_bullets)} bullets")
        print(f"  Bullet numbers: {approved_bullets}")

        # Build approved bullets data structure
        approved_bullets_data = []
        for bullet_number in approved_bullets:
            bullet_number_str = str(bullet_number)  # Convert to string for dict lookup
            approved_text = approved_texts.get(bullet_number_str, '')
            customized_text = customized_bullets.get(bullet_number_str, None)

            print(f"  Bullet #{bullet_number}: approved_text={approved_text[:50] if approved_text else 'EMPTY'}...")

            approved_bullets_data.append({
                'bullet_number': bullet_number,
                'approved_text': approved_text,
                'customized_text': customized_text
            })

        # Save to database
        success = update_cv_session_approved_bullets(cv_session_id, approved_bullets_data)

        if success:
            print(f"✅ Successfully saved {len(approved_bullets_data)} bullets to database")
            return jsonify({'success': True})
        else:
            print("❌ Failed to save to database")
            return jsonify({'error': 'Failed to save to database'}), 500

    except Exception as e:
        print(f"❌ Error saving approved bullets: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to save approved bullets'}), 500


# ==================== INTERVIEW PREP ROUTES ====================

@app.route("/interview-prep", methods=["GET"])
def interview_prep():
    """Main interview prep page - show completed CV customizations"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    try:
        engine = get_db_connection()
        if not engine:
            return "Database connection error", 500

        with engine.connect() as conn:
            # Get all completed CV customizations with their interview sessions
            query = text("""
                SELECT
                    cvs.id as cv_session_id,
                    cvs.job_title,
                    cvs.job_company,
                    cvs.analysis_id,
                    cvs.selected_headline,
                    cvs.approved_bullets,
                    COALESCE(
                        json_agg(
                            json_build_object(
                                'session_id', intv.id,
                                'completed', intv.completed,
                                'overall_score', intv.overall_score,
                                'current_question', intv.current_question,
                                'created_at', intv.created_at
                            )
                            ORDER BY intv.created_at DESC
                        ) FILTER (WHERE intv.id IS NOT NULL),
                        '[]'
                    ) as practice_sessions
                FROM cv_customization_sessions cvs
                LEFT JOIN interview_sessions intv ON cvs.id = intv.cv_session_id
                WHERE cvs.user_id = :user_id
                    AND cvs.approved_bullets IS NOT NULL
                    AND cvs.approved_bullets::text != '[]'
                GROUP BY cvs.id
                ORDER BY cvs.created_at DESC
            """)

            result = conn.execute(query, {"user_id": user_id})
            rows = result.fetchall()

            cv_sessions = []
            for row in rows:
                cv_sessions.append({
                    'id': row[0],  # Changed from cv_session_id to id to match template
                    'job_title': row[1],
                    'job_company': row[2],
                    'analysis_id': row[3],
                    'selected_headline': row[4],
                    'approved_bullets': row[5],
                    'practice_sessions': row[6] if isinstance(row[6], list) else json.loads(row[6]) if row[6] else []
                })

        return render_template('interview_prep.html', cv_sessions=cv_sessions)

    except Exception as e:
        print(f"Error in interview prep: {e}")
        import traceback
        traceback.print_exc()
        # Return error details for debugging (remove in production)
        return f"Error loading interview prep: {str(e)}", 500


@app.route("/interview-prep/start/<int:cv_session_id>", methods=["POST"])
def interview_prep_start(cv_session_id):
    """Start new interview practice session"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        # Get CV session data
        print(f"DEBUG: Getting CV session {cv_session_id}...")
        cv_session = get_cv_session(cv_session_id)
        if not cv_session or cv_session['user_id'] != user_id:
            return jsonify({'error': 'Session not found'}), 404
        print(f"DEBUG: CV session retrieved successfully")

        # Get job analysis
        print(f"DEBUG: Getting analysis {cv_session['analysis_id']}...")
        analysis = get_analysis_by_id(cv_session['analysis_id'], user_id)
        if not analysis:
            return jsonify({'error': 'Analysis not found'}), 404
        print(f"DEBUG: Analysis retrieved successfully")

        # Extract approved bullet texts
        approved_bullets = cv_session.get('approved_bullets', [])
        print(f"DEBUG: approved_bullets type: {type(approved_bullets)}")
        print(f"DEBUG: approved_bullets length: {len(approved_bullets) if approved_bullets else 0}")

        bullet_texts = []
        try:
            for i, bullet in enumerate(approved_bullets):
                # Handle both dict and string formats
                if isinstance(bullet, dict):
                    bullet_text = bullet.get('customized_text') or bullet.get('approved_text', '')
                elif isinstance(bullet, str):
                    bullet_text = bullet
                else:
                    bullet_text = ''

                if bullet_text:
                    bullet_texts.append(bullet_text)
                    print(f"DEBUG: Bullet {i+1}: {bullet_text[:50]}...")
        except Exception as bullet_err:
            print(f"ERROR extracting bullets: {bullet_err}")
            raise

        print(f"DEBUG: Extracted {len(bullet_texts)} bullet texts")

        if not bullet_texts:
            return jsonify({'error': 'No bullets found to generate questions'}), 400

        # Generate questions
        print(f"DEBUG: Calling generate_interview_questions...")
        try:
            questions = generate_interview_questions(
                job_description=analysis['job_description'],
                cv_bullets=bullet_texts,
                selected_headline=cv_session['selected_headline'],
                user_id=user_id
            )
            print(f"DEBUG: Questions generated successfully")
        except Exception as gen_err:
            print(f"ERROR in generate_interview_questions: {gen_err}")
            raise

        if not questions:
            return jsonify({'error': 'Failed to generate questions'}), 500

        # Create interview session
        print(f"DEBUG: Creating interview session in database...")
        try:
            engine = get_db_connection()
            with engine.connect() as conn:
                from sqlalchemy import text as sql_text  # Avoid any potential conflicts
                insert_query = sql_text("""
                    INSERT INTO interview_sessions (
                        user_id, cv_session_id, questions, current_question
                    ) VALUES (
                        :user_id, :cv_session_id, :questions, 1
                    )
                    RETURNING id
                """)

                result = conn.execute(insert_query, {
                    "user_id": user_id,
                    "cv_session_id": cv_session_id,
                    "questions": json.dumps(questions)
                })
                conn.commit()

                interview_session_id = result.fetchone()[0]

            print(f"✓ Created interview session {interview_session_id}")
        except Exception as db_err:
            print(f"ERROR creating interview session: {db_err}")
            raise

        # Redirect to first question
        return redirect(f'/interview-prep/question/{interview_session_id}/1')

    except Exception as e:
        print(f"Error starting interview session: {e}")
        import traceback
        traceback.print_exc()
        # Return detailed error for debugging (remove in production)
        error_message = repr(e)  # Use repr() instead of str() to avoid potential shadowing
        return jsonify({'error': f'Failed to start session: {error_message}'}), 500


@app.route("/interview-prep/question/<int:session_id>/<int:question_num>", methods=["GET"])
def interview_prep_question(session_id, question_num):
    """Show interview question"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    try:
        engine = get_db_connection()
        with engine.connect() as conn:
            # Get session
            query = text("""
                SELECT
                    intv.id, intv.user_id, intv.cv_session_id, intv.questions,
                    intv.answers, intv.evaluations, intv.current_question, intv.completed,
                    cvs.job_title, cvs.job_company
                FROM interview_sessions intv
                JOIN cv_customization_sessions cvs ON intv.cv_session_id = cvs.id
                WHERE intv.id = :session_id AND intv.user_id = :user_id
            """)

            result = conn.execute(query, {"session_id": session_id, "user_id": user_id})
            row = result.fetchone()

            if not row:
                return "Session not found", 404

            # Parse JSON fields
            questions = row[3]
            if isinstance(questions, str):
                questions = json.loads(questions)

            answers = row[4]
            if isinstance(answers, str):
                answers = json.loads(answers) if answers else {}
            elif answers is None:
                answers = {}

            evaluations = row[5]
            if isinstance(evaluations, str):
                evaluations = json.loads(evaluations) if evaluations else {}
            elif evaluations is None:
                evaluations = {}

            current_question = row[6]
            completed = row[7]
            job_title = row[8]
            job_company = row[9]

        # Check if already completed
        if completed:
            return redirect(f'/interview-prep/summary/{session_id}')

        # Validate question number
        if question_num < 1 or question_num > len(questions):
            return "Invalid question number", 400

        # Get the question
        question_data = questions[question_num - 1]

        # Check if we're showing evaluation
        question_key = str(question_num)
        show_evaluation = question_key in evaluations

        if show_evaluation:
            evaluation = evaluations[question_key]
            return render_template('interview_evaluation.html',
                                 session_id=session_id,
                                 question_num=question_num,
                                 question_text=question_data.get('question_text', ''),
                                 user_answer=answers.get(question_key, ''),
                                 overall_score=evaluation.get('overall_score', 0),
                                 details_score=evaluation.get('details_score', 0),
                                 organization_score=evaluation.get('organization_score', 0),
                                 analysis_score=evaluation.get('analysis_score', 0),
                                 ownership_score=evaluation.get('ownership_score', 0),
                                 feedback_summary=evaluation.get('feedback_summary', ''),
                                 strong_points=evaluation.get('strong_points', []),
                                 improvement_areas=evaluation.get('improvement_areas', []),
                                 better_answer_example=evaluation.get('better_answer_example', ''),
                                 job_title=job_title,
                                 job_company=job_company,
                                 total_questions=len(questions))

        # Show question form
        return render_template('interview_question.html',
                             session_id=session_id,
                             question_num=question_num,
                             question_text=question_data.get('question_text', ''),
                             bullet_reference=question_data.get('bullet_reference', ''),
                             job_title=job_title,
                             job_company=job_company,
                             total_questions=len(questions))

    except Exception as e:
        print(f"Error loading question: {e}")
        import traceback
        traceback.print_exc()
        return "Error loading question", 500


@app.route("/interview-prep/submit/<int:session_id>/<int:question_num>", methods=["POST"])
def interview_prep_submit(session_id, question_num):
    """Submit answer and get evaluation"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        # Get answer from form
        answer = request.form.get('answer', '').strip()
        if not answer:
            return "Please provide an answer", 400

        engine = get_db_connection()
        with engine.connect() as conn:
            # Get session and CV data
            query = text("""
                SELECT
                    intv.questions, intv.answers, intv.evaluations, intv.cv_session_id,
                    cvs.analysis_id, cvs.approved_bullets
                FROM interview_sessions intv
                JOIN cv_customization_sessions cvs ON intv.cv_session_id = cvs.id
                WHERE intv.id = :session_id AND intv.user_id = :user_id
            """)

            result = conn.execute(query, {"session_id": session_id, "user_id": user_id})
            row = result.fetchone()

            if not row:
                return "Session not found", 404

            # Parse JSON fields
            questions = row[0]
            if isinstance(questions, str):
                questions = json.loads(questions)

            answers = row[1]
            if isinstance(answers, str):
                answers = json.loads(answers) if answers else {}
            elif answers is None:
                answers = {}

            evaluations = row[2]
            if isinstance(evaluations, str):
                evaluations = json.loads(evaluations) if evaluations else {}
            elif evaluations is None:
                evaluations = {}

            analysis_id = row[4]

            approved_bullets = row[5]
            if isinstance(approved_bullets, str):
                approved_bullets = json.loads(approved_bullets) if approved_bullets else []
            elif approved_bullets is None:
                approved_bullets = []

        # Get question
        question_data = questions[question_num - 1]

        # Get job description
        analysis = get_analysis_by_id(analysis_id, user_id)

        # Extract bullet texts
        bullet_texts = []
        for bullet in approved_bullets:
            # Handle both dict and string formats
            if isinstance(bullet, dict):
                bullet_text = bullet.get('customized_text') or bullet.get('approved_text', '')
            elif isinstance(bullet, str):
                bullet_text = bullet
            else:
                bullet_text = ''

            if bullet_text:
                bullet_texts.append(bullet_text)

        # Evaluate answer
        print(f"Evaluating answer for question {question_num}...")
        evaluation = evaluate_interview_answer(
            question_text=question_data['question_text'],
            user_answer=answer,
            job_description=analysis['job_description'],
            cv_bullets=bullet_texts
        )

        if not evaluation:
            return "Failed to evaluate answer", 500

        # Save answer and evaluation
        question_key = str(question_num)
        answers[question_key] = answer
        evaluations[question_key] = evaluation

        with engine.connect() as conn:
            # Update session
            update_query = text("""
                UPDATE interview_sessions
                SET answers = :answers,
                    evaluations = :evaluations,
                    current_question = :next_question,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = :session_id
            """)

            conn.execute(update_query, {
                "session_id": session_id,
                "answers": json.dumps(answers),
                "evaluations": json.dumps(evaluations),
                "next_question": question_num + 1 if question_num < len(questions) else question_num
            })
            conn.commit()

        print(f"✓ Saved answer and evaluation for question {question_num}")

        # Redirect to show evaluation
        return redirect(f'/interview-prep/question/{session_id}/{question_num}')

    except Exception as e:
        print(f"Error submitting answer: {e}")
        import traceback
        traceback.print_exc()
        return "Error processing answer", 500


@app.route("/interview-prep/summary/<int:session_id>", methods=["GET"])
def interview_prep_summary(session_id):
    """Show final summary"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    try:
        engine = get_db_connection()
        with engine.connect() as conn:
            # Get session
            query = text("""
                SELECT
                    intv.questions, intv.evaluations, intv.completed, intv.overall_score,
                    cvs.job_title, cvs.job_company
                FROM interview_sessions intv
                JOIN cv_customization_sessions cvs ON intv.cv_session_id = cvs.id
                WHERE intv.id = :session_id AND intv.user_id = :user_id
            """)

            result = conn.execute(query, {"session_id": session_id, "user_id": user_id})
            row = result.fetchone()

            if not row:
                return "Session not found", 404

            # Parse JSON fields
            questions = row[0]
            if isinstance(questions, str):
                questions = json.loads(questions)

            evaluations = row[1]
            if isinstance(evaluations, str):
                evaluations = json.loads(evaluations) if evaluations else {}
            elif evaluations is None:
                evaluations = {}

            completed = row[2]
            overall_score = row[3]
            job_title = row[4]
            job_company = row[5]

        # If not completed yet, complete it now
        if not completed and evaluations:
            # Calculate overall scores
            scores = []
            details_scores = []
            organization_scores = []
            analysis_scores = []
            ownership_scores = []

            for eval_data in evaluations.values():
                scores.append(eval_data['overall_score'])
                details_scores.append(eval_data['details_score'])
                organization_scores.append(eval_data['organization_score'])
                analysis_scores.append(eval_data['analysis_score'])
                ownership_scores.append(eval_data['ownership_score'])

            overall_score = round(sum(scores) / len(scores), 1) if scores else 0

            avg_scores = {
                'details': round(sum(details_scores) / len(details_scores), 1) if details_scores else 0,
                'organization': round(sum(organization_scores) / len(organization_scores), 1) if organization_scores else 0,
                'analysis': round(sum(analysis_scores) / len(analysis_scores), 1) if analysis_scores else 0,
                'ownership': round(sum(ownership_scores) / len(ownership_scores), 1) if ownership_scores else 0
            }

            # Mark as complete
            with engine.connect() as conn:
                update_query = text("""
                    UPDATE interview_sessions
                    SET completed = TRUE,
                        overall_score = :overall_score,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE id = :session_id
                """)

                conn.execute(update_query, {
                    "session_id": session_id,
                    "overall_score": overall_score
                })
                conn.commit()

        else:
            # Calculate avg scores from evaluations
            avg_scores = {
                'details': 0,
                'organization': 0,
                'analysis': 0,
                'ownership': 0
            }

            if evaluations:
                details_scores = []
                organization_scores = []
                analysis_scores = []
                ownership_scores = []

                for eval_data in evaluations.values():
                    details_scores.append(eval_data['details_score'])
                    organization_scores.append(eval_data['organization_score'])
                    analysis_scores.append(eval_data['analysis_score'])
                    ownership_scores.append(eval_data['ownership_score'])

                avg_scores = {
                    'details': round(sum(details_scores) / len(details_scores), 1) if details_scores else 0,
                    'organization': round(sum(organization_scores) / len(organization_scores), 1) if organization_scores else 0,
                    'analysis': round(sum(analysis_scores) / len(analysis_scores), 1) if analysis_scores else 0,
                    'ownership': round(sum(ownership_scores) / len(ownership_scores), 1) if ownership_scores else 0
                }

        return render_template('interview_summary.html',
                             session_id=session_id,
                             overall_score=overall_score,
                             avg_scores=avg_scores,
                             questions=questions,
                             evaluations=evaluations,
                             job_title=job_title,
                             job_company=job_company)

    except Exception as e:
        print(f"Error loading summary: {e}")
        import traceback
        traceback.print_exc()
        return "Error loading summary", 500


@app.route("/get-analysis/<int:analysis_id>", methods=["GET"])
def get_analysis(analysis_id):
    """Get full analysis by ID"""
    print(f"\n>>> ROUTE /get-analysis/{analysis_id} called")
    if 'user_id' not in session:
        print("  ERROR: User not authenticated")
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']
    print(f"  User ID from session: {user_id}")

    try:
        analysis = get_analysis_by_id(analysis_id, user_id)

        if not analysis:
            print(f"  RETURNING 404: Analysis not found")
            return jsonify({'error': 'Analysis not found'}), 404

        # Check if CV customization session exists for this analysis
        cv_session = get_cv_session_by_analysis(analysis_id, user_id)
        cv_progress = None

        if cv_session:
            # Determine current step
            current_step = 'headline'  # Default
            if cv_session.get('selected_headline'):
                current_step = 'bullets'
            if cv_session.get('approved_bullets') and len(cv_session['approved_bullets']) >= 6:
                current_step = 'preview'

            cv_progress = {
                'session_id': cv_session['id'],
                'current_step': current_step,
                'has_headline': bool(cv_session.get('selected_headline')),
                'approved_bullets_count': len(cv_session['approved_bullets']) if cv_session.get('approved_bullets') else 0,
                'status': cv_session.get('status', 'in_progress')
            }
            print(f"  CV Session found: {cv_progress}")

        print(f"  SUCCESS: Returning analysis data")
        return jsonify({
            'success': True,
            'analysis': analysis,
            'cv_progress': cv_progress
        })

    except Exception as e:
        print(f"  ERROR in route: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to get analysis'}), 500


@app.route("/delete-analysis/<int:analysis_id>", methods=["POST"])
def delete_analysis(analysis_id):
    """Delete an analysis"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        success = delete_analysis_by_id(analysis_id, user_id)

        if not success:
            return jsonify({'error': 'Analysis not found'}), 404

        log_user_activity('analysis_delete', f'Deleted analysis ID: {analysis_id}')

        return jsonify({'success': True, 'message': 'Analysis deleted'})

    except Exception as e:
        print(f"Error deleting analysis: {e}")
        return jsonify({'error': 'Failed to delete analysis'}), 500


@app.route("/user-guide", methods=["GET"])
def user_guide():
    """User guide page"""
    if 'user_id' not in session:
        return redirect('/login')
    user_id = session['user_id']
    master_template = get_user_master_template(user_id)
    return render_template('user_guide.html', has_master_template=master_template is not None)


@app.route("/get-started", methods=["GET"])
def get_started():
    """Get Started onboarding page"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']
    username = session.get('username', 'User')

    # Get master template status
    master_template = get_user_master_template(user_id)
    has_master_template = master_template is not None

    # Get prompt status
    has_custom_prompt = False
    current_prompt_text = ""
    try:
        engine = get_db_connection()
        if engine:
            with engine.connect() as conn:
                result = conn.execute(text("""
                    SELECT upp.custom_prompt_text, pt.prompt_text
                    FROM user_prompt_preferences upp
                    LEFT JOIN prompt_templates pt ON upp.template_id = pt.id
                    WHERE upp.user_id = :user_id AND upp.is_active = TRUE
                    ORDER BY upp.created_at DESC LIMIT 1
                """), {"user_id": user_id})
                row = result.fetchone()
                if row:
                    has_custom_prompt = bool(row[0])
                    current_prompt_text = row[0] if row[0] else (row[1] or "")
    except Exception as e:
        print(f"Error getting prompt for get-started: {e}")

    return render_template('get_started.html',
                           username=username,
                           has_master_template=has_master_template,
                           master_template=master_template,
                           has_custom_prompt=has_custom_prompt,
                           current_prompt_text=current_prompt_text)


@app.route("/consolidate-cvs", methods=["POST"])
def consolidate_cvs():
    """Accept multiple CV files, extract text, and use Claude to consolidate into master template"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        files = request.files.getlist('cv_files')
        if not files or len(files) == 0:
            return jsonify({'error': 'No files uploaded'}), 400

        allowed_extensions = {'pdf', 'docx', 'doc', 'txt'}
        extracted_texts = []

        for file in files:
            if not file or file.filename == '':
                continue
            ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
            if ext not in allowed_extensions:
                return jsonify({'error': f'File "{file.filename}" is not a supported format. Use PDF, DOCX or TXT.'}), 400

            file_data = file.read()
            if len(file_data) > 5 * 1024 * 1024:
                return jsonify({'error': f'File "{file.filename}" exceeds 5MB limit'}), 400

            # Extract text
            if ext == 'pdf':
                text_content = extract_text_from_pdf(file_data)
            elif ext in ('docx', 'doc'):
                text_content = extract_text_from_docx(file_data)
            else:
                text_content = file_data.decode('utf-8', errors='replace')

            if text_content and text_content.strip():
                extracted_texts.append({'filename': file.filename, 'text': text_content})

        if not extracted_texts:
            return jsonify({'error': 'Could not extract text from any of the uploaded files'}), 400

        # If only one file, no consolidation needed — use it directly
        if len(extracted_texts) == 1:
            return jsonify({
                'success': True,
                'consolidated_text': extracted_texts[0]['text'],
                'stats': {'files_processed': 1, 'note': 'Single file uploaded — used directly as master template'}
            })

        # Build Claude consolidation prompt
        cv_sections = ""
        for i, item in enumerate(extracted_texts, 1):
            cv_sections += f"\n\n--- CV VERSION {i}: {item['filename']} ---\n{item['text']}\n"

        consolidation_prompt = f"""You are consolidating {len(extracted_texts)} versions of the same person's CV into one comprehensive master template.

RULES — follow these exactly:
1. Keep ALL unique bullet points and phrases VERBATIM — do not rewrite, paraphrase, or improve any wording
2. Remove ONLY true word-for-word duplicate sentences/bullets (where text is identical or near-identical)
3. Keep ALL career summary variations as separate clearly labeled sections: [VERSION 1 — focus], [VERSION 2 — focus], etc.
4. Keep ALL job title variations for each role (list them on one line separated by " / ")
5. Keep ALL unique context/intro lines for each role
6. Organise the output in this structure:
   - CAREER SUMMARY (all versions labeled)
   - CORE EXPERTISE (all unique skills from all CVs)
   - EMPLOYMENT HISTORY (reverse chronological, with all unique bullets per role)
   - PERMANENT ROLES SUMMARISED
   - EDUCATION
   - CERTIFICATIONS
   - TECHNOLOGY & TOOLS (if present)
7. Use plain text formatting with ===== section dividers
8. Do NOT add any commentary, explanations or notes — output only the master template text

Here are the {len(extracted_texts)} CV versions to consolidate:
{cv_sections}

Output the complete master template now:"""

        # Use Haiku for speed — this task is text organisation, not analysis
        anthropic_client = get_anthropic_client()
        message = anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=6000,
            messages=[{"role": "user", "content": consolidation_prompt}]
        )

        consolidated_text = message.content[0].text

        log_user_activity('cv_consolidation', f'Consolidated {len(extracted_texts)} CV files into master template')

        return jsonify({
            'success': True,
            'consolidated_text': consolidated_text,
            'stats': {
                'files_processed': len(extracted_texts),
                'filenames': [item['filename'] for item in extracted_texts]
            }
        })

    except Exception as e:
        print(f"Error consolidating CVs: {e}")
        return jsonify({'error': f'Consolidation failed: {str(e)}'}), 500


@app.route("/prompt-settings", methods=["GET"])
def prompt_settings():
    """Display prompt settings page"""
    if 'user_id' not in session:
        return redirect('/login')

    user_id = session['user_id']

    try:
        engine = get_db_connection()
        if not engine:
            return "Database error", 500

        with engine.connect() as conn:
            # Get all available templates
            templates_result = conn.execute(text("""
                SELECT id, name, description, target_profile, version
                FROM prompt_templates
                ORDER BY is_default DESC, name ASC
            """))

            templates = []
            for row in templates_result:
                templates.append({
                    'id': row[0],
                    'name': row[1],
                    'description': row[2],
                    'target_profile': row[3],
                    'version': row[4]
                })

            # Get user's current active template and prompt
            user_pref_result = conn.execute(text("""
                SELECT upp.template_id, upp.custom_prompt_text, pt.name, pt.prompt_text
                FROM user_prompt_preferences upp
                LEFT JOIN prompt_templates pt ON upp.template_id = pt.id
                WHERE upp.user_id = :user_id AND upp.is_active = TRUE
                ORDER BY upp.created_at DESC
                LIMIT 1
            """), {"user_id": user_id})

            user_pref_row = user_pref_result.fetchone()

            current_template_id = None
            current_template_name = "No template selected"
            current_prompt_text = ""

            if user_pref_row:
                current_template_id = user_pref_row[0]
                custom_prompt = user_pref_row[1]
                template_name = user_pref_row[2]
                template_prompt = user_pref_row[3]

                current_template_name = template_name or "Custom"
                current_prompt_text = custom_prompt if custom_prompt else template_prompt

            has_custom_prompt = bool(user_pref_row and user_pref_row[1])

            return render_template('prompt_settings.html',
                                 templates=templates,
                                 current_template_id=current_template_id,
                                 current_template_name=current_template_name,
                                 current_prompt_text=current_prompt_text,
                                 has_custom_prompt=has_custom_prompt)

    except Exception as e:
        print(f"Error loading prompt settings: {e}")
        import traceback
        traceback.print_exc()
        return "Error loading settings", 500


@app.route("/get-template-prompt/<int:template_id>", methods=["GET"])
def get_template_prompt(template_id):
    """Get full prompt text for a specific template (for preview)"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        engine = get_db_connection()
        if not engine:
            return jsonify({'error': 'Database error'}), 500

        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT prompt_text FROM prompt_templates WHERE id = :template_id
            """), {"template_id": template_id})

            row = result.fetchone()
            if not row:
                return jsonify({'error': 'Template not found'}), 404

            return jsonify({'prompt_text': row[0]})

    except Exception as e:
        print(f"Error getting template prompt: {e}")
        return jsonify({'error': 'Failed to get template'}), 500


@app.route("/switch-template", methods=["POST"])
def switch_template():
    """Switch user's active template"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        data = request.get_json()
        template_id = data.get('template_id')

        if not template_id:
            return jsonify({'error': 'template_id is required'}), 400

        engine = get_db_connection()
        if not engine:
            return jsonify({'error': 'Database error'}), 500

        with engine.connect() as conn:
            # Deactivate all current preferences for this user
            conn.execute(text("""
                UPDATE user_prompt_preferences
                SET is_active = FALSE
                WHERE user_id = :user_id
            """), {"user_id": user_id})

            # Create new active preference with selected template
            conn.execute(text("""
                INSERT INTO user_prompt_preferences (user_id, template_id, is_active)
                VALUES (:user_id, :template_id, TRUE)
            """), {"user_id": user_id, "template_id": template_id})

            conn.commit()

            # Get the template name for the response
            template_result = conn.execute(text("""
                SELECT name FROM prompt_templates WHERE id = :template_id
            """), {"template_id": template_id})

            template_row = template_result.fetchone()
            template_name = template_row[0] if template_row else "Unknown"

            log_user_activity('prompt_template_switch', f'Switched to template: {template_name}')

            return jsonify({
                'success': True,
                'message': f'Switched to {template_name} template',
                'template_name': template_name
            })

    except Exception as e:
        print(f"Error switching template: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to switch template'}), 500


@app.route("/save-custom-prompt", methods=["POST"])
def save_custom_prompt():
    """Save a custom prompt text for the user"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        data = request.get_json()
        custom_prompt = data.get('custom_prompt', '').strip()

        if not custom_prompt:
            return jsonify({'error': 'Prompt text cannot be empty'}), 400

        engine = get_db_connection()
        if not engine:
            return jsonify({'error': 'Database error'}), 500

        with engine.connect() as conn:
            # Check if user already has an active preference
            result = conn.execute(text("""
                SELECT id FROM user_prompt_preferences
                WHERE user_id = :user_id AND is_active = TRUE
                LIMIT 1
            """), {"user_id": user_id})

            existing = result.fetchone()

            if existing:
                # Update existing preference with custom prompt
                conn.execute(text("""
                    UPDATE user_prompt_preferences
                    SET custom_prompt_text = :custom_prompt
                    WHERE id = :pref_id
                """), {"custom_prompt": custom_prompt, "pref_id": existing[0]})
            else:
                # Create new preference with just custom prompt (no template)
                conn.execute(text("""
                    INSERT INTO user_prompt_preferences (user_id, custom_prompt_text, is_active)
                    VALUES (:user_id, :custom_prompt, TRUE)
                """), {"user_id": user_id, "custom_prompt": custom_prompt})

            conn.commit()

        log_user_activity('custom_prompt_saved', 'User saved a custom prompt')

        return jsonify({'success': True, 'message': 'Custom prompt saved successfully'})

    except Exception as e:
        print(f"Error saving custom prompt: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to save custom prompt'}), 500


@app.route("/clear-custom-prompt", methods=["POST"])
def clear_custom_prompt():
    """Clear user's custom prompt and revert to selected template"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        engine = get_db_connection()
        if not engine:
            return jsonify({'error': 'Database error'}), 500

        with engine.connect() as conn:
            conn.execute(text("""
                UPDATE user_prompt_preferences
                SET custom_prompt_text = NULL
                WHERE user_id = :user_id AND is_active = TRUE
            """), {"user_id": user_id})
            conn.commit()

        log_user_activity('custom_prompt_cleared', 'User reset to template prompt')

        return jsonify({'success': True, 'message': 'Reset to template prompt'})

    except Exception as e:
        print(f"Error clearing custom prompt: {e}")
        return jsonify({'error': 'Failed to reset prompt'}), 500


@app.route("/ai-usage-stats", methods=["GET"])
def ai_usage_stats():
    """Get AI usage statistics for the user"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    user_id = session['user_id']

    try:
        engine = get_db_connection()
        if not engine:
            return jsonify({'error': 'Database error'}), 500

        with engine.connect() as conn:
            # Get total usage this month
            result = conn.execute(text("""
                SELECT
                    COUNT(*) as total_analyses,
                    SUM(tokens_input) as total_input_tokens,
                    SUM(tokens_output) as total_output_tokens,
                    SUM(estimated_cost) as total_cost
                FROM ai_usage_tracking
                WHERE user_id = :user_id
                AND created_at >= DATE_TRUNC('month', CURRENT_DATE)
            """), {"user_id": user_id})

            stats = result.fetchone()

            return jsonify({
                'total_analyses': stats[0] or 0,
                'total_input_tokens': stats[1] or 0,
                'total_output_tokens': stats[2] or 0,
                'total_cost': float(stats[3] or 0)
            })

    except Exception as e:
        print(f"Error getting usage stats: {e}")
        return jsonify({'error': 'Failed to get stats'}), 500


@app.route("/ai-diagnostic", methods=["GET"])
def ai_diagnostic():
    """Diagnostic endpoint to check AI setup"""
    import os
    diagnostics = {}

    # Check API key
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    diagnostics['api_key_set'] = api_key is not None
    diagnostics['api_key_length'] = len(api_key) if api_key else 0

    # Check Anthropic library
    try:
        import anthropic
        diagnostics['anthropic_installed'] = True
        diagnostics['anthropic_version'] = anthropic.__version__ if hasattr(anthropic, '__version__') else 'unknown'
    except ImportError:
        diagnostics['anthropic_installed'] = False
        diagnostics['anthropic_version'] = None

    # Check PyPDF2
    try:
        import PyPDF2
        diagnostics['pypdf2_installed'] = True
    except ImportError:
        diagnostics['pypdf2_installed'] = False

    # Check python-docx
    try:
        import docx
        diagnostics['python_docx_installed'] = True
    except ImportError:
        diagnostics['python_docx_installed'] = False

    # Check database connection
    try:
        engine = get_db_connection()
        diagnostics['database_connected'] = engine is not None
        if engine:
            with engine.connect() as conn:
                result = conn.execute(text("SELECT COUNT(*) FROM user_cvs"))
                diagnostics['cv_table_accessible'] = True
                diagnostics['total_cvs'] = result.fetchone()[0]
    except Exception as e:
        diagnostics['database_connected'] = False
        diagnostics['database_error'] = str(e)

    return jsonify(diagnostics)


if __name__ == "__main__":
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=False)
























